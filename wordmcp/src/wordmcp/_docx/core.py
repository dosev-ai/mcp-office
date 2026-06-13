from __future__ import annotations

import contextlib
import json
import logging
import os
import re
import tempfile
from collections import OrderedDict
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from wordmcp._docx import _facade

logger = logging.getLogger(__name__)

_doc_cache: OrderedDict[str, tuple[Document, int | None]] = OrderedDict()


def _check_path(path: str, must_exist: bool = True) -> Path:
    facade = _facade()
    roots_env = os.getenv("WORD_ALLOWLIST_ROOTS", "").strip()
    if not roots_env:
        raise facade.ValidationError("No allowlist configured: WORD_ALLOWLIST_ROOTS is unset")
    if "\x00" in path:
        raise facade.ValidationError("Null byte in path")
    if path.startswith("\\\\") or path.startswith("//"):
        raise facade.ValidationError("UNC paths are not permitted")

    resolved = Path(path).resolve()
    ext = resolved.suffix.lower()
    if ext != ".docx":
        raise facade.ValidationError(
            f"Unsupported extension: {ext!r} (only .docx allowed)"
        )

    roots = [root.strip() for root in roots_env.split(",") if root.strip()]
    for root in roots:
        try:
            if resolved.is_relative_to(Path(root).resolve()):
                if must_exist and not resolved.exists():
                    raise facade.FileError(f"File not found: {resolved}")
                if must_exist:
                    file_size = resolved.stat().st_size
                    if file_size > facade._MAX_DOCX_BYTES:
                        raise facade.ValidationError(
                            "File exceeds maximum size "
                            f"({facade._MAX_DOCX_BYTES // (1024 * 1024)} MB)"
                        )
                return resolved
        except (ValueError, TypeError):
            continue

    raise facade.ValidationError("Path not in allowlist")


def _check_image_path(path: str) -> Path:
    facade = _facade()
    roots_env = os.getenv("WORD_ALLOWLIST_ROOTS", "").strip()
    if not roots_env:
        raise facade.ValidationError("No allowlist configured: WORD_ALLOWLIST_ROOTS is unset")
    if "\x00" in path:
        raise facade.ValidationError("Null byte in path")
    if path.startswith("\\\\") or path.startswith("//"):
        raise facade.ValidationError("UNC paths are not permitted")

    resolved = Path(path).resolve()
    ext = resolved.suffix.lower()
    if ext not in facade._ALLOWED_IMAGE_EXTENSIONS:
        raise facade.ValidationError(f"Unsupported image extension: {ext!r}")

    roots = [root.strip() for root in roots_env.split(",") if root.strip()]
    for root in roots:
        try:
            if resolved.is_relative_to(Path(root).resolve()):
                if not resolved.exists():
                    raise facade.FileError(f"Image not found: {resolved}")
                return resolved
        except (ValueError, TypeError):
            continue

    raise facade.ValidationError("Path not in allowlist")


def _check_evidence_path(path: str) -> Path:
    """Validate evidence_path: allowlist check + .jsonl suffix enforcement.

    Uses WORD_ALLOWLIST_ROOTS (same env var as _check_path).
    Raises ValidationError on failure.
    """
    facade = _facade()
    roots_env = os.environ.get("WORD_ALLOWLIST_ROOTS", "").strip()
    if not roots_env:
        raise facade.ValidationError("No allowlist configured: WORD_ALLOWLIST_ROOTS is unset")
    if "\x00" in path:
        raise facade.ValidationError("Null byte in path")
    if path.startswith("\\\\") or path.startswith("//"):
        raise facade.ValidationError("UNC paths are not permitted")
    resolved = Path(path).resolve()
    if resolved.suffix.lower() != ".jsonl":
        raise facade.ValidationError(
            f"evidence_path must have .jsonl suffix, got: {resolved.suffix!r}"
        )
    roots = [Path(r.strip()).resolve() for r in roots_env.split(",") if r.strip()]
    if not any(resolved.is_relative_to(root) for root in roots):
        raise facade.ValidationError(
            "evidence_path not within allowed roots"
        )
    return resolved


def _check_write() -> None:
    facade = _facade()
    if os.getenv("WORD_ENABLE_WRITE", "").strip().lower() != "true":
        raise facade.NotAllowedError("Write operations require WORD_ENABLE_WRITE=true")


def _check_confirm(confirm: bool) -> None:
    facade = _facade()
    if not confirm:
        raise facade.ValidationError("confirm=True required for mutating operations")


def _audit_log(op: str, path: Path | str, extra: dict | None = None) -> None:
    record: dict[str, Any] = {
        "ts": datetime.now(timezone.utc).isoformat(),
        "op": op,
        "path": str(path),
    }
    if extra:
        record.update(extra)
    logger.info(json.dumps(record))


def _atomic_save(doc: Document, target: Path) -> None:
    facade = _facade()
    fd, tmp_path = tempfile.mkstemp(dir=str(target.parent), suffix=".tmp")
    try:
        facade.os.close(fd)
        doc.save(tmp_path)
        facade.os.replace(tmp_path, str(target))
    except Exception:
        with contextlib.suppress(OSError):
            facade.os.unlink(tmp_path)
        raise


def _load_doc(path: Path | str) -> Document:
    path = Path(path)
    facade = _facade()
    key = str(path)
    try:
        current_mtime = path.stat().st_mtime_ns
    except OSError:
        current_mtime = None
    cached = facade._doc_cache.get(key)
    if cached is not None:
        doc, cached_mtime = cached
        if current_mtime is not None and current_mtime == cached_mtime:
            facade._doc_cache.move_to_end(key)
            return doc
        del facade._doc_cache[key]
    if len(facade._doc_cache) >= facade._DOC_CACHE_MAX:
        facade._doc_cache.popitem(last=False)
    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.error("Failed to load document %s: %s", path, exc)
        raise facade.FileError(f"Cannot open document ({type(exc).__name__})") from None
    facade._doc_cache[key] = (doc, current_mtime)
    return doc


def _evict_doc(path: Path) -> None:
    _facade()._doc_cache.pop(str(path), None)


def _outline_level(para: Any) -> int | None:
    sname = para.style.name
    if sname == "Title":
        return 0
    if sname.startswith("Heading "):
        parts = sname.split()
        if len(parts) == 2 and parts[-1].isdigit():
            return int(parts[-1])
    return None


def _emu_to_inches(value: Any) -> float | None:
    return round(value.inches, 4) if value is not None else None


def _run_to_dict(run: Any) -> dict:
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font_name": run.font.name,
        "font_size_pt": round(run.font.size / 12700, 2) if run.font.size else None,
    }


def _paragraph_to_dict(para: Any, index: int) -> dict:
    pf = para.paragraph_format
    return {
        "index": index,
        "style": para.style.name,
        "text": para.text,
        "outline_level": _outline_level(para),
        "runs": [_run_to_dict(run) for run in para.runs],
        "left_indent_inches": _emu_to_inches(pf.left_indent),
        "first_line_indent_inches": _emu_to_inches(pf.first_line_indent),
        "alignment": pf.alignment.name if pf.alignment else None,
    }


# ---------------------------------------------------------------------------
# Inline-markdown helpers (bold / italic / underline inside token values)
# ---------------------------------------------------------------------------

# Matches **bold**, __underline__, _italic_ — order matters: __ before _
_MD_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|_(.+?)_", re.DOTALL)


def _parse_inline_markdown(
    text: str,
) -> list[tuple[str, bool, bool, bool]]:
    """Split *text* into (segment, bold, italic, underline) tuples.

    Plain text between markers gets (False, False, False). Returns a single
    plain-text tuple when no markers are found.
    """
    segments: list[tuple[str, bool, bool, bool]] = []
    last_end = 0
    for m in _MD_INLINE_RE.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end : m.start()], False, False, False))
        if m.group(1) is not None:  # **bold**
            segments.append((m.group(1), True, False, False))
        elif m.group(2) is not None:  # __underline__
            segments.append((m.group(2), False, False, True))
        else:  # _italic_
            segments.append((m.group(3), False, True, False))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False, False, False))
    return segments or [(text, False, False, False)]


def _apply_inline_markdown_runs(
    para: Any,
    segments: list[tuple[str, bool, bool, bool]],
    base_run: Any,
) -> None:
    """Replace all paragraph runs with richly-formatted runs from *segments*.

    Each segment inherits the base run's rPr (font, size, colour) and then
    overlays the requested bold / italic / underline attribute. Plain segments
    (no markdown flags) are emitted with the base rPr unchanged.
    """
    p_elem = para._p
    base_r = base_run._r
    base_rpr = base_r.find(qn("w:rPr"))  # capture BEFORE removal loop

    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)

    for seg_text, seg_bold, seg_italic, seg_underline in segments:
        if not seg_text:
            continue

        r_elem = OxmlElement("w:r")
        rpr = deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")

        if seg_bold:
            b = rpr.find(qn("w:b"))
            if b is None:
                b = OxmlElement("w:b")
                rpr.append(b)
            # Remove w:val so the element activates bold (absence of val = True)
            b.attrib.pop(qn("w:val"), None)

        if seg_italic:
            i = rpr.find(qn("w:i"))
            if i is None:
                i = OxmlElement("w:i")
                rpr.append(i)
            i.attrib.pop(qn("w:val"), None)

        if seg_underline:
            u = rpr.find(qn("w:u"))
            if u is None:
                u = OxmlElement("w:u")
                rpr.append(u)
            u.set(qn("w:val"), "single")

        r_elem.append(rpr)

        t_elem = OxmlElement("w:t")
        if seg_text != seg_text.strip():
            t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_elem.text = seg_text
        r_elem.append(t_elem)

        p_elem.append(r_elem)


def _do_find_replace(
    doc: Document,
    find_text: str,
    replace_text: str,
    paragraph_index: int | None = None,
    occurrence: int | None = None,
) -> int:
    total_count = 0
    occ_seen = 0
    occ_done = False

    def _process_para(para: Any) -> None:
        nonlocal total_count, occ_seen, occ_done
        if occ_done:
            return

        full_text = "".join(run.text for run in para.runs)
        if find_text not in full_text or not para.runs:
            return

        if occurrence is None:
            new_text = full_text.replace(find_text, replace_text)
            replaced = full_text.count(find_text)
        else:
            out_parts: list[str] = []
            offset = 0
            replaced = 0
            while True:
                idx = full_text.find(find_text, offset)
                if idx == -1:
                    out_parts.append(full_text[offset:])
                    break
                occ_seen += 1
                if occ_seen == occurrence:
                    out_parts.append(full_text[offset:idx])
                    out_parts.append(replace_text)
                    out_parts.append(full_text[idx + len(find_text) :])
                    replaced = 1
                    occ_done = True
                    break
                out_parts.append(full_text[offset : idx + len(find_text)])
                offset = idx + len(find_text)
            if replaced == 0:
                return
            new_text = "".join(out_parts)

        first_run = para.runs[0]
        bold = first_run.bold
        italic = first_run.italic
        font_name = first_run.font.name
        font_size = first_run.font.size

        # Gate on replace_text (the value being injected), NOT new_text.
        # new_text is the full paragraph text after substitution, which may
        # contain underscores from other unfilled {{TOKEN_NAME}} patterns —
        # those would false-trigger the italic detector.
        has_markdown = _MD_INLINE_RE.search(replace_text) is not None

        if has_markdown:
            segments = _parse_inline_markdown(new_text)
            _apply_inline_markdown_runs(para, segments, first_run)
        else:
            for run in para.runs:
                run.text = ""
            first_run.text = new_text
            first_run.bold = bold
            first_run.italic = italic
            if font_name is not None:
                first_run.font.name = font_name
            if font_size is not None:
                first_run.font.size = font_size

        total_count += replaced

    if paragraph_index is not None:
        paragraphs = doc.paragraphs
        if 0 <= paragraph_index < len(paragraphs):
            _process_para(paragraphs[paragraph_index])
    else:
        for para in doc.paragraphs:
            _process_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_para(para)

    return total_count


def _get_list_num_id(document: Document, style_name: str) -> str | None:
    try:
        style_elem = document.styles[style_name].element
        style_ppr = style_elem.find(qn("w:pPr"))
        if style_ppr is None:
            return None
        style_numpr = style_ppr.find(qn("w:numPr"))
        if style_numpr is None:
            return None
        style_numid_el = style_numpr.find(qn("w:numId"))
        if style_numid_el is None:
            return None
        return style_numid_el.get(qn("w:val"))
    except (KeyError, AttributeError):
        return None
