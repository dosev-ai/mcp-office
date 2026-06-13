"""Word document review and outline extraction.

Layer: _docx (python-docx only — no MCP imports, no COM).
All functions accept an allowlist-validated path and return plain dicts.
"""
from __future__ import annotations

import logging
import zipfile
from typing import Any

from docx.opc.exceptions import PackageNotFoundError

from wordmcp._docx import _facade
from wordmcp._docx.core import _check_path, _load_doc

logger = logging.getLogger(__name__)

VALID_CHECKS: frozenset[str] = frozenset({
    "style_consistency",
    "heading_hierarchy",
    "word_count",
    "table_structure",
})

_DISPATCH: dict[str, Any] = {}  # populated after function defs below


def get_document_outline(path: str) -> dict:
    """Extract headings, paragraph summary, and table count from a .docx file.

    Returns:
        {
            "path": str,
            "headings": [{"level": int, "text": str}, ...],
            "paragraph_count": int,
            "table_count": int,
        }
    """
    facade = _facade()
    checked = _check_path(path)
    try:
        doc = _load_doc(str(checked))
    except (zipfile.BadZipFile, PackageNotFoundError, Exception) as exc:
        raise facade.FileError("Cannot open document: invalid format or file not found") from exc

    headings = []
    paragraph_count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 0
            headings.append({"level": level, "text": para.text})
        else:
            paragraph_count += 1

    return {
        "path": str(checked),
        "headings": headings,
        "paragraph_count": paragraph_count,
        "table_count": len(doc.tables),
    }


def review_document(path: str, check_names: list[str] | None = None) -> dict:
    """Run structural quality checks on a .docx document.

    Args:
        path: Path to the .docx file.
        check_names: Optional list of check names to run. If None, all checks run.
                     Valid values: "style_consistency", "heading_hierarchy",
                     "word_count", "table_structure"

    Returns:
        {
            "path": str,
            "checks_run": [str, ...],
            "results": {
                check_name: {"passed": bool, "detail": str}
            }
        }
    """
    facade = _facade()
    checked = _check_path(path)

    # Validate check_names
    if check_names is not None:
        unknown = set(check_names) - VALID_CHECKS
        if unknown:
            raise facade.ValidationError(
                f"Unknown check name(s): {sorted(unknown)}. "
                f"Valid: {sorted(VALID_CHECKS)}"
            )
        to_run = list(check_names)
    else:
        to_run = sorted(VALID_CHECKS)

    try:
        doc = _load_doc(str(checked))
    except (zipfile.BadZipFile, PackageNotFoundError, Exception) as exc:
        raise facade.FileError("Cannot open document: invalid format or file not found") from exc

    results = {}
    for name in to_run:
        results[name] = _DISPATCH[name](doc)

    return {
        "path": str(checked),
        "checks_run": to_run,
        "results": results,
    }


def _check_style_consistency(doc: Any) -> dict:
    """Check that all Normal paragraphs use the same font name."""
    fonts: set[str] = set()
    for para in doc.paragraphs:
        if para.style and para.style.name == "Normal" and para.text.strip():
            font_name = para.style.font.name if para.style.font else None
            if font_name:
                fonts.add(font_name)
    passed = len(fonts) <= 1
    detail = "Consistent" if passed else f"Multiple fonts detected: {sorted(fonts)}"
    return {"passed": passed, "detail": detail}


def _check_heading_hierarchy(doc: Any) -> dict:
    """Check that heading levels do not skip (e.g. H1 → H3 is invalid)."""
    levels: list[int] = []
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                levels.append(int(style_name.split()[-1]))
            except (ValueError, IndexError):
                pass

    violations: list[str] = []
    for i in range(1, len(levels)):
        if levels[i] > levels[i - 1] + 1:
            violations.append(f"H{levels[i-1]} → H{levels[i]} at heading {i+1}")

    passed = len(violations) == 0
    detail = "Hierarchy is valid" if passed else f"Violations: {violations}"
    return {"passed": passed, "detail": detail}


def _check_word_count(doc: Any) -> dict:
    """Count total words across all paragraphs."""
    total = sum(len(para.text.split()) for para in doc.paragraphs)
    return {"passed": True, "detail": f"Total word count: {total}", "word_count": total}


def _check_table_structure(doc: Any) -> dict:
    """Check that all rows in each table have the same column count."""
    violations: list[str] = []
    for t_idx, table in enumerate(doc.tables):
        col_counts = [len(row.cells) for row in table.rows]
        if len(set(col_counts)) > 1:
            violations.append(f"Table {t_idx+1}: inconsistent column counts {col_counts}")
    passed = len(violations) == 0
    detail = "All tables consistent" if passed else f"Violations: {violations}"
    return {"passed": passed, "detail": detail}


# Wire dispatch table
_DISPATCH["style_consistency"] = _check_style_consistency
_DISPATCH["heading_hierarchy"] = _check_heading_hierarchy
_DISPATCH["word_count"] = _check_word_count
_DISPATCH["table_structure"] = _check_table_structure
