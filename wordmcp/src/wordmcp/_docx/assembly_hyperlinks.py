"""manage_hyperlinks, _do_hyperlink_replace, _validate_url.

Extracted from assembly.py (structural decomposition, no logic changes).
"""
from __future__ import annotations

import re
import urllib.parse
from typing import Any

from docx.oxml.ns import qn

from wordmcp._docx import _facade

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_TOKEN_RE = re.compile(r"\{\{[^{}]+\}\}")
_ALLOWED_SCHEMES = {"http", "https", "mailto"}

_HYPERLINK_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)

_MAX_HYPERLINKS = 500


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _do_hyperlink_replace(doc: Any, part: Any, token: str, url: str) -> int:
    """Replace {{token}} in body paragraphs with a clickable w:hyperlink element.

    Scans body paragraphs and all table cells. For each paragraph whose full
    text contains the token, clears all runs, writes the "before" text into the
    first run, inserts a w:hyperlink element immediately after it, and appends
    a trailing run for any text after the token.

    Returns the count of replacements made.
    """
    from docx.oxml import OxmlElement

    rel_id = part.relate_to(url, _HYPERLINK_REL_TYPE, is_external=True)
    count = 0

    def _process_para(para: Any) -> None:
        nonlocal count
        full_text = "".join(r.text for r in para.runs)
        if token not in full_text or not para.runs:
            return
        before, _, after = full_text.partition(token)

        first_run = para.runs[0]

        # Clear all runs
        for r in para.runs:
            r.text = ""

        # "before" text in first run
        first_run.text = before

        # Build w:hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), rel_id)
        w_r = OxmlElement("w:r")
        w_t = OxmlElement("w:t")
        w_t.text = url
        w_t.set(qn("xml:space"), "preserve")
        w_r.append(w_t)
        hyperlink.append(w_r)
        first_run._r.addnext(hyperlink)

        # "after" text in a trailing run
        if after:
            a_r = OxmlElement("w:r")
            a_t = OxmlElement("w:t")
            a_t.text = after
            if after[0] == " ":
                a_t.set(qn("xml:space"), "preserve")
            a_r.append(a_t)
            hyperlink.addnext(a_r)

        count += 1

    for para in doc.paragraphs:
        _process_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)
    return count


def _validate_url(url: str, facade: Any) -> None:
    """Validate that url is a str with an allowed scheme.

    Allowed schemes: http, https, mailto.
    Raises ValidationError for javascript:, file://, data:, UNC paths,
    or any unrecognized scheme.
    """
    if not isinstance(url, str):
        raise facade.ValidationError(
            f"url must be a str, got {type(url).__name__}"
        )
    parsed = urllib.parse.urlparse(url)
    scheme = parsed.scheme.lower()
    if scheme not in _ALLOWED_SCHEMES:
        raise facade.ValidationError(
            f"URL scheme {scheme!r} is not allowed. "
            f"Allowed schemes: {sorted(_ALLOWED_SCHEMES)}"
        )


# ---------------------------------------------------------------------------
# Tool: manage_hyperlinks
# ---------------------------------------------------------------------------


def manage_hyperlinks(
    path: str,
    hyperlinks: dict[str, dict[str, Any]],
    confirm: bool = False,
) -> dict[str, Any]:
    """Update hyperlink URLs and labels identified by display text in a Word document.

    Each key in `hyperlinks` is the current display text of a hyperlink to update.
    Each value is a dict with required key 'url' (str) and optional key 'label' (str).

    Gate order:
      1. _check_write()
      2. _check_confirm(confirm)
      3. _check_path(path)
      4. PRE-FLIGHT: validate ALL entries before any document I/O
      5. Load document, scan hyperlink elements, update, save, evict, audit
    """
    facade = _facade()

    # Gate 1 — write enable (MUST be first)
    facade._check_write()
    # Gate 2 — explicit caller confirmation
    facade._check_confirm(confirm)
    # Gate 3 — path allowlist
    resolved = facade._check_path(path)

    # Gate 4 — PRE-FLIGHT validation of ALL entries before _load_doc
    if len(hyperlinks) > _MAX_HYPERLINKS:
        raise facade.ValidationError(
            f"Too many entries: {len(hyperlinks)} > {_MAX_HYPERLINKS}"
        )
    for display_text, entry in hyperlinks.items():
        if not isinstance(entry, dict):
            raise facade.ValidationError(
                f"Entry for {display_text!r} must be a dict, got {type(entry).__name__}"
            )
        if "url" not in entry:
            raise facade.ValidationError(
                f"Entry for {display_text!r} is missing required key 'url'"
            )
        _validate_url(entry["url"], facade)

    # Gate 5 — document operations
    doc = facade._load_doc(resolved)
    part = doc.part

    # Build display_text -> hyperlink element map.
    # Concatenate all w:t children within each w:hyperlink element to get display text.
    # This handles hyperlinks whose text is split across multiple w:r/w:t children.
    found_map: dict[str, Any] = {}
    for el in doc.element.body.iter(qn("w:hyperlink")):
        display_text = "".join(
            t.text or "" for t in el.iter(qn("w:t"))
        )
        if display_text in hyperlinks:
            found_map[display_text] = el

    results: dict[str, dict[str, Any]] = {}
    count_updated = 0

    for display_text, entry in hyperlinks.items():
        if display_text not in found_map:
            results[display_text] = {"status": "not_found"}
            continue

        el = found_map[display_text]
        new_url: str = entry["url"]
        new_label: str | None = entry.get("label")

        # Update URL via relationship: mutate _target on existing rel in-place.
        # _Relationship._target is a plain instance attribute (not a property).
        # Mutating it directly preserves the original rId and the w:hyperlink r:id.
        # python-docx version dependency: pin python-docx>=1.1.2 in pyproject.toml.
        r_id = el.get(qn("r:id"))
        if r_id and r_id in part.rels:
            part.rels[r_id]._target = new_url
        else:
            # No existing relationship: create a new one.
            new_r_id = part.relate_to(new_url, _HYPERLINK_REL_TYPE, is_external=True)
            el.set(qn("r:id"), new_r_id)

        # Update label: replace text in all w:t children of this hyperlink element.
        if new_label is not None:
            w_t_elements = list(el.iter(qn("w:t")))
            if w_t_elements:
                w_t_elements[0].text = new_label
                for extra_t in w_t_elements[1:]:
                    extra_t.text = ""

        results[display_text] = {"status": "updated", "url": new_url}
        count_updated += 1

    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log(
        "manage_hyperlinks",
        resolved,
        extra={"updated": count_updated},
    )

    return {
        "updated": count_updated,
        "results": results,
        "path": str(resolved),
    }
