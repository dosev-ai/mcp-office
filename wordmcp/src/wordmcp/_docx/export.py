from __future__ import annotations

import os
from pathlib import Path
from typing import Any

from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from wordmcp._docx import _facade


def export_as_text(path: str) -> dict:
    facade = _facade()
    resolved = facade._check_path(path)
    doc = facade._load_doc(resolved)

    lines: list[str] = []
    paragraph_count = 0
    table_count = 0

    for elem in doc.element.body:
        if elem.tag == qn("w:p"):
            paragraph = DocxParagraph(elem, doc)
            lines.append(paragraph.text)
            paragraph_count += 1
        elif elem.tag == qn("w:tbl"):
            table = DocxTable(elem, doc)
            table_count += 1
            for row in table.rows:
                row_text = " | ".join(
                    " ".join(paragraph.text for paragraph in cell.paragraphs).strip()
                    for cell in row.cells
                )
                lines.append(row_text)

    text = "\n".join(lines)
    max_chars = int(os.getenv("WORD_MAX_TEXT_CHARS", str(facade._DEFAULT_MAX_TEXT_CHARS)))
    truncated = False
    if max_chars > 0 and len(text) > max_chars:
        truncated = True
        text = text[:max_chars] + f"\n[TRUNCATED: output capped at {max_chars} characters]"

    return {
        "path": str(resolved),
        "text": text,
        "truncated": truncated,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
    }


def export_to_markdown(path: str) -> dict:
    facade = _facade()
    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    max_chars = int(os.getenv("WORD_MAX_TEXT_CHARS", str(facade._DEFAULT_MAX_TEXT_CHARS)))
    lines: list[str] = []
    total_chars = 0
    truncated = False
    list_number_counter = 0

    def _para_to_md(para: Any) -> str:
        nonlocal list_number_counter
        style = para.style.name if para.style else "Normal"
        inline = ""
        for run in para.runs:
            text = run.text
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            inline += text
        if not inline:
            inline = para.text

        if style == "Title":
            list_number_counter = 0
            return f"# {inline}"
        if style.startswith("Heading "):
            list_number_counter = 0
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            level = max(1, min(level, 6))
            return "#" * level + " " + inline
        if style == "List Bullet":
            list_number_counter = 0
            return f"- {inline}"
        if style == "List Number":
            list_number_counter += 1
            return f"{list_number_counter}. {inline}"
        list_number_counter = 0
        return inline

    for child in document.element.body:
        if total_chars >= max_chars:
            truncated = True
            break
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, document)
            md_line = _para_to_md(paragraph)
            lines.append(md_line)
            total_chars += len(md_line)
        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, document)
            if table.rows:
                header = " | ".join(cell.text.replace("|", "\\|") for cell in table.rows[0].cells)
                sep = " | ".join("---" for _ in table.rows[0].cells)
                lines.append(f"| {header} |")
                lines.append(f"| {sep} |")
                total_chars += len(header) + len(sep)
                for row in table.rows[1:]:
                    row_text = " | ".join(cell.text.replace("|", "\\|") for cell in row.cells)
                    lines.append(f"| {row_text} |")
                    total_chars += len(row_text)

    markdown = "\n".join(lines)
    if len(markdown) > max_chars:
        markdown = markdown[:max_chars]
        truncated = True
    return {"markdown": markdown, "truncated": truncated}


def _check_output_path(output_path: str, allowed_suffixes: tuple[str, ...]) -> Path:
    """Validate output_path: within allowlist roots, suffix in allowed_suffixes.

    Unlike _check_path, we do NOT require the file to exist (it will be created).
    We DO require the parent directory to exist, and the path to be within the
    WORD_ALLOWLIST_ROOTS.

    Raises facade.ValidationError on failure.
    """
    facade = _facade()
    roots_env = os.getenv("WORD_ALLOWLIST_ROOTS", "").strip()
    if not roots_env:
        raise facade.ValidationError("No allowlist configured: WORD_ALLOWLIST_ROOTS is unset")
    if "\x00" in output_path:
        raise facade.ValidationError("Null byte in path")
    if output_path.startswith("\\\\") or output_path.startswith("//"):
        raise facade.ValidationError("UNC paths are not permitted")

    resolved = Path(output_path).resolve()
    ext = resolved.suffix.lower()
    if ext not in allowed_suffixes:
        raise facade.ValidationError(
            f"Unsupported output extension: {ext!r} (allowed: {allowed_suffixes})"
        )
    if not resolved.parent.is_dir():
        raise facade.ValidationError(
            f"output_path parent directory does not exist: {resolved.parent}"
        )

    roots = [root.strip() for root in roots_env.split(",") if root.strip()]
    for root in roots:
        try:
            if resolved.is_relative_to(Path(root).resolve()):
                return resolved
        except (ValueError, TypeError):
            continue

    raise facade.ValidationError("output_path not in allowlist")


def export_as_file_txt(path: str, output_path: str) -> dict:
    """Export document as plain text to output_path.

    Requires WORD_ENABLE_WRITE=true and confirm=True (enforced by caller in server layer).
    Returns {"output_path": ..., "paragraphs_exported": N, "tables_exported": N}.
    """
    facade = _facade()
    resolved = facade._check_path(path)
    out = _check_output_path(output_path, (".txt",))
    doc = facade._load_doc(resolved)

    lines: list[str] = []
    paragraph_count = 0
    table_count = 0

    for elem in doc.element.body:
        if elem.tag == qn("w:p"):
            paragraph = DocxParagraph(elem, doc)
            lines.append(paragraph.text)
            paragraph_count += 1
        elif elem.tag == qn("w:tbl"):
            table = DocxTable(elem, doc)
            table_count += 1
            for row in table.rows:
                row_text = " | ".join(
                    " ".join(paragraph.text for paragraph in cell.paragraphs).strip()
                    for cell in row.cells
                )
                lines.append(row_text)

    text = "\n".join(lines)
    out.write_text(text, encoding="utf-8")
    return {
        "output_path": str(out),
        "paragraphs_exported": paragraph_count,
        "tables_exported": table_count,
    }


def export_markdown_to_file(path: str, output_path: str) -> dict:
    """Export document as Markdown to output_path (UTF-8 .md file).

    Heading 1-9 → # through #########.
    Bold runs → **text**, italic → *text*, bold+italic → ***text***.
    Normal paragraphs → plain text line + blank line separator.
    Tables → GFM pipe table with |---| separator after first row.
    Lists → - item (bullet) or 1. item (numbered).
    Page breaks → --- (horizontal rule).

    Requires WORD_ENABLE_WRITE=true and confirm=True (enforced by caller).
    Returns {"output_path": ..., "paragraphs_exported": N, "tables_exported": N}.
    """
    facade = _facade()
    resolved = facade._check_path(path)
    out = _check_output_path(output_path, (".md",))
    document = facade._load_doc(resolved)

    lines: list[str] = []
    paragraph_count = 0
    table_count = 0
    list_number_counter = 0

    def _run_inline(para: Any) -> str:
        """Render paragraph runs with bold/italic markers."""
        result = ""
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            result += text
        return result or para.text

    def _para_to_md(para: Any) -> list[str]:
        """Convert a paragraph to one or more markdown lines."""
        nonlocal list_number_counter
        style = para.style.name if para.style else "Normal"
        inline = _run_inline(para)

        if style == "Title":
            list_number_counter = 0
            return [f"# {inline}", ""]
        if style.startswith("Heading "):
            list_number_counter = 0
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            level = max(1, min(level, 9))
            return ["#" * level + " " + inline, ""]
        if style in ("List Bullet", "List Bullet 2", "List Bullet 3"):
            list_number_counter = 0
            return [f"- {inline}"]
        if style in ("List Number", "List Number 2", "List Number 3"):
            list_number_counter += 1
            return [f"{list_number_counter}. {inline}"]
        list_number_counter = 0
        # Check for page break
        for elem in para._element:
            if elem.tag == qn("w:r"):
                for sub in elem:
                    if sub.tag == qn("w:lastRenderedPageBreak") or sub.tag == qn("w:br"):
                        btype = sub.get(qn("w:type"), "")
                        if btype == "page":
                            return ["---", ""]
        return [inline, ""]

    for child in document.element.body:
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, document)
            md_lines = _para_to_md(paragraph)
            lines.extend(md_lines)
            paragraph_count += 1
        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, document)
            table_count += 1
            list_number_counter = 0
            if table.rows:
                # Header row
                header_cells = [cell.text.replace("|", "\\|") for cell in table.rows[0].cells]
                lines.append("| " + " | ".join(header_cells) + " |")
                lines.append("| " + " | ".join("---" for _ in header_cells) + " |")
                for row in table.rows[1:]:
                    row_cells = [cell.text.replace("|", "\\|") for cell in row.cells]
                    lines.append("| " + " | ".join(row_cells) + " |")
                lines.append("")

    markdown = "\n".join(lines)
    out.write_text(markdown, encoding="utf-8")
    return {
        "output_path": str(out),
        "paragraphs_exported": paragraph_count,
        "tables_exported": table_count,
    }


def get_headers_footers(path: str) -> list:
    facade = _facade()
    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    result = []
    for index, section in enumerate(document.sections):
        result.append(
            {
                "section_index": index,
                "header_text": "\n".join(paragraph.text for paragraph in section.header.paragraphs),
                "footer_text": "\n".join(paragraph.text for paragraph in section.footer.paragraphs),
                "header_linked_to_previous": section.header.is_linked_to_previous,
                "footer_linked_to_previous": section.footer.is_linked_to_previous,
            }
        )
    return result
