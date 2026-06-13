from __future__ import annotations

import re
from datetime import datetime, timezone

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from wordmcp._docx import _facade


def manage_comments(
    path: str,
    operation: str,
    text: str | None = None,
    author: str | None = None,
    paragraph_index: int | None = None,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    content_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    if operation == "list":
        resolved = facade._check_path(path)
        document = facade._load_doc(resolved)
        comments_part = None
        for rel in document.part.rels.values():
            if "comments" in rel.reltype.lower():
                comments_part = rel.target_part
                break
        if comments_part is None:
            return {"comments": []}

        from lxml import etree as lxml_etree

        comments_xml = lxml_etree.fromstring(comments_part.blob)
        return {
            "comments": [
                {
                    "id": int(comment.get(qn("w:id"), 0)),
                    "author": comment.get(qn("w:author"), ""),
                    "date": comment.get(qn("w:date"), ""),
                    "text": "".join(
                        elem.text or "" for elem in comment.iter() if elem.tag == qn("w:t")
                    ),
                }
                for comment in comments_xml.findall(qn("w:comment"))
            ]
        }

    if operation == "resolve":
        facade._check_write()
        facade._check_confirm(confirm)
        if paragraph_index is None:
            raise facade.ValidationError("resolve requires comment_id via paragraph_index param")
        comment_id = int(paragraph_index)
        resolved_path = facade._check_path(path)
        document = facade._load_doc(resolved_path)
        comments_part = None
        for rel in document.part.rels.values():
            if "comments" in rel.reltype.lower():
                comments_part = rel.target_part
                break
        if comments_part is None:
            raise facade.NotFoundError("Document has no comments part")
        from lxml import etree as lxml_etree
        comments_xml = lxml_etree.fromstring(comments_part.blob)
        target = None
        for comment in comments_xml.findall(qn("w:comment")):
            if int(comment.get(qn("w:id"), -1)) == comment_id:
                target = comment
                break
        if target is None:
            raise facade.NotFoundError(f"Comment id {comment_id} not found")
        target.set(qn("w:done"), "1")
        comments_part._blob = lxml_etree.tostring(comments_xml)
        facade._atomic_save(document, resolved_path)
        facade._evict_doc(resolved_path)
        facade._audit_log("manage_comments_resolve", resolved_path)
        return {"status": "resolved", "comment_id": comment_id}

    if operation == "delete":
        facade._check_write()
        facade._check_confirm(confirm)
        if paragraph_index is None:
            raise facade.ValidationError("delete requires comment_id via paragraph_index param")
        comment_id = int(paragraph_index)
        resolved_path = facade._check_path(path)
        document = facade._load_doc(resolved_path)
        comments_part = None
        for rel in document.part.rels.values():
            if "comments" in rel.reltype.lower():
                comments_part = rel.target_part
                break
        if comments_part is None:
            raise facade.NotFoundError("Document has no comments part")
        from lxml import etree as lxml_etree
        comments_xml = lxml_etree.fromstring(comments_part.blob)
        target = None
        for comment in comments_xml.findall(qn("w:comment")):
            if int(comment.get(qn("w:id"), -1)) == comment_id:
                target = comment
                break
        if target is None:
            raise facade.NotFoundError(f"Comment id {comment_id} not found")
        comments_xml.remove(target)
        cid_str = str(comment_id)
        body = document.element.body
        for tag_name in (qn("w:commentRangeStart"), qn("w:commentRangeEnd")):
            for elem in body.iter(tag_name):
                if elem.get(qn("w:id")) == cid_str:
                    elem.getparent().remove(elem)
        for ref in body.iter(qn("w:commentReference")):
            if ref.get(qn("w:id")) == cid_str:
                ref.getparent().remove(ref)
        comments_part._blob = lxml_etree.tostring(comments_xml)
        facade._atomic_save(document, resolved_path)
        facade._evict_doc(resolved_path)
        facade._audit_log("manage_comments_delete", resolved_path)
        return {"status": "deleted", "comment_id": comment_id}

    if operation != "add":
        raise facade.ValidationError(f"Unknown operation: {operation!r}. Valid: list, add, resolve, delete")

    facade._check_write()
    facade._check_confirm(confirm)
    if not text:
        raise facade.ValidationError("text is required for add operation")

    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    paragraphs = document.paragraphs
    if paragraph_index is not None and (paragraph_index < 0 or paragraph_index >= len(paragraphs)):
        raise facade.NotFoundError(f"Paragraph index {paragraph_index} out of range")

    comments_part = None
    comments_xml_is_new = False
    for rel in document.part.rels.values():
        if "comments" in rel.reltype.lower():
            comments_part = rel.target_part
            break

    if comments_part is None:
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part
        from lxml import etree as lxml_etree

        min_xml = (
            b'<w:comments xmlns:w="http://schemas.openxmlformats.org/'
            b'wordprocessingml/2006/main"/>'
        )
        comments_xml = lxml_etree.fromstring(min_xml)
        comments_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        comments_part = Part(
            PackURI("/word/comments.xml"),
            comments_ct,
            lxml_etree.tostring(comments_xml),
            document.part.package,
        )
        document.part.relate_to(comments_part, content_type)
        comments_xml_is_new = True
    else:
        from lxml import etree as lxml_etree

        comments_xml = lxml_etree.fromstring(comments_part.blob)
        comments_xml_is_new = True

    existing_ids = [int(comment.get(qn("w:id"), 0)) for comment in comments_xml.findall(qn("w:comment"))]
    new_id = max(existing_ids, default=-1) + 1

    comment = OxmlElement("w:comment")
    clean_text = facade._CTRL_RE.sub("", text) if text else ""
    clean_author = facade._CTRL_RE.sub("", author) if author else "MCP"
    comment.set(qn("w:id"), str(new_id))
    comment.set(qn("w:author"), clean_author)
    comment.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    p_elem = OxmlElement("w:p")
    r_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = clean_text
    r_elem.append(t_elem)
    p_elem.append(r_elem)
    comment.append(p_elem)
    comments_xml.append(comment)

    if paragraph_index is not None:
        para = paragraphs[paragraph_index]
        cs = OxmlElement("w:commentRangeStart")
        cs.set(qn("w:id"), str(new_id))
        ce = OxmlElement("w:commentRangeEnd")
        ce.set(qn("w:id"), str(new_id))
        ref_run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "CommentReference")
        rpr.append(rstyle)
        ref_run.append(rpr)
        ref_elem = OxmlElement("w:commentReference")
        ref_elem.set(qn("w:id"), str(new_id))
        ref_run.append(ref_elem)
        para._p.append(cs)
        para._p.append(ce)
        para._p.append(ref_run)

    if comments_xml_is_new:
        from lxml import etree as lxml_etree

        comments_part._blob = lxml_etree.tostring(comments_xml)

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("manage_comments_add", resolved)
    return {"status": "added", "added": True, "comment_id": new_id, "author": clean_author}


def add_hyperlink(
    path: str,
    paragraph_index: int,
    text: str,
    url: str,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    if not re.match(r"^https?://", url, re.IGNORECASE):
        raise facade.ValidationError("url must begin with http:// or https://")
    if len(url) > 2048:
        raise facade.ValidationError("url exceeds maximum length of 2048 characters")
    if "\x00" in url or facade._CTRL_RE.search(url):
        raise facade.ValidationError("URL contains invalid characters")

    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    paragraphs = document.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        raise facade.NotFoundError(
            f"Paragraph index {paragraph_index} out of range (0-{len(paragraphs) - 1})"
        )

    para = paragraphs[paragraph_index]
    hyperlink_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    rel_id = para.part.relate_to(url, hyperlink_rel, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_elem = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    run_elem.append(rpr)
    text_elem = OxmlElement("w:t")
    text_elem.text = facade._CTRL_RE.sub("", text) if text else url
    run_elem.append(text_elem)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("add_hyperlink", resolved, extra={"url_len": len(url)})
    return {"paragraph_index": paragraph_index, "url": url, "text": text or url}


def add_footnote(
    path: str,
    paragraph_index: int,
    text: str,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    paragraphs = document.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        raise facade.NotFoundError(
            f"Paragraph index {paragraph_index} out of range (0-{len(paragraphs) - 1})"
        )

    footnotes_part = None
    for rel in document.part.rels.values():
        if "footnotes" in rel.reltype.lower():
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        raise facade.WordMCPError(
            "Document has no footnotes part. "
            "Open the document in Word, add a manual footnote to create the footnotes part, then retry."
        )

    footnotes_xml = footnotes_part._element
    existing_ids = [
        int(footnote.get(qn("w:id"), 0))
        for footnote in footnotes_xml.findall(qn("w:footnote"))
        if footnote.get(qn("w:id"), "0").lstrip("-").isdigit()
    ]
    new_id = max((value for value in existing_ids if value >= 0), default=0) + 1

    footnote = OxmlElement("w:footnote")
    footnote.set(qn("w:id"), str(new_id))
    footnote.set(qn("w:type"), "normal")
    fn_p = OxmlElement("w:p")
    fn_r = OxmlElement("w:r")
    fn_t = OxmlElement("w:t")
    fn_t.text = facade._CTRL_RE.sub("", text) if text else ""
    fn_r.append(fn_t)
    fn_p.append(fn_r)
    footnote.append(fn_p)
    footnotes_xml.append(footnote)

    para = paragraphs[paragraph_index]
    ref_run = para.add_run()
    ref_rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "FootnoteReference")
    ref_rpr.append(rstyle)
    ref_run._r.insert(0, ref_rpr)
    ref_elem = OxmlElement("w:footnoteReference")
    ref_elem.set(qn("w:id"), str(new_id))
    ref_run._r.append(ref_elem)

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("add_footnote", resolved)
    return {"footnote_id": new_id, "paragraph_index": paragraph_index}
