from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from wordmcp._docx import _facade
from wordmcp._docx.core import _do_find_replace, _get_list_num_id


def add_list(
    path: str,
    items: list,
    list_type: str = "bullet",
    level: int = 0,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    if list_type not in ("bullet", "number"):
        raise facade.ValidationError(
            f"Unknown list_type {list_type!r}. Valid: bullet, number"
        )
    if not items:
        raise facade.ValidationError("items must not be empty")
    if len(items) > 500:
        raise facade.ValidationError("items exceeds max 500")
    if not (0 <= level <= 8):
        raise facade.ValidationError("level must be 0-8")

    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    style = "List Bullet" if list_type == "bullet" else "List Number"
    num_id = _get_list_num_id(document, style) if level > 0 else None

    for item in items:
        item_text = facade._CTRL_RE.sub("", str(item)) if item else ""
        para = document.add_paragraph(item_text, style=style)
        if level > 0 and num_id is not None:
            ppr = para._p.get_or_add_pPr()
            existing = ppr.find(qn("w:numPr"))
            if existing is not None:
                ppr.remove(existing)
            numpr = OxmlElement("w:numPr")
            ilvl_el = OxmlElement("w:ilvl")
            ilvl_el.set(qn("w:val"), str(level))
            numid_el = OxmlElement("w:numId")
            numid_el.set(qn("w:val"), num_id)
            numpr.append(ilvl_el)
            numpr.append(numid_el)
            ppr.append(numpr)

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("add_list", resolved)
    return {"added": len(items), "list_type": list_type, "level": level}


def find_replace(
    path: str,
    find_text: str,
    replace_text: str,
    paragraph_index: int | None = None,
    occurrence: int | None = None,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)

    if not find_text:
        raise facade.ValidationError("find_text must not be empty")
    if occurrence is not None and occurrence < 1:
        raise facade.ValidationError("occurrence must be >= 1 (1-based)")

    doc = facade._load_doc(resolved)
    count = _do_find_replace(
        doc,
        find_text,
        replace_text,
        paragraph_index=paragraph_index,
        occurrence=occurrence,
    )

    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log(
        "find_replace",
        resolved,
        extra={"find_len": len(find_text), "replacements": count},
    )
    return {
        "replacements_made": count,
        "find_len": len(find_text),
        "replace_len": len(replace_text),
    }


def insert_image(
    path: str,
    image_path: str,
    width_inches: float | None = None,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    image_resolved = facade._check_image_path(image_path)

    if width_inches is not None and width_inches <= 0:
        raise facade.ValidationError("width_inches must be > 0")

    doc = facade._load_doc(resolved)
    width = Inches(width_inches) if width_inches is not None else None
    try:
        doc.add_picture(str(image_resolved), width=width)
    except Exception as exc:
        raise facade.FileError(f"Failed to insert image ({type(exc).__name__})") from exc

    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log(
        "insert_image",
        resolved,
        extra={"image_path": str(image_resolved), "width_inches": width_inches},
    )
    return {"inserted": True, "image_path": str(image_resolved), "width_inches": width_inches}


def set_paragraph_format(
    path: str,
    paragraph_index: int,
    space_before: float | None = None,
    space_after: float | None = None,
    line_spacing: float | None = None,
    table_cell: dict | None = None,
    confirm: bool = False,
) -> dict:
    """Apply spacing/line-spacing to a paragraph's ParagraphFormat.

    paragraph_index: 0-based index into doc.paragraphs (body paragraphs only).
    space_before: space before paragraph in points (pt). None = no change.
    space_after: space after paragraph in points (pt). None = no change.
    line_spacing: line spacing multiplier (1.0 = single, 1.5 = one-and-a-half,
        2.0 = double). Sets WD_LINE_SPACING.MULTIPLE rule. None = no change.
    table_cell: if provided, must be a dict with 'row' (int) and 'col' (int) keys,
        targeting a paragraph in that cell of the first table (table_index 0).
        When specified, paragraph_index refers to the paragraph index within
        that cell (not the body paragraph list).
    """
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)

    if space_before is None and space_after is None and line_spacing is None:
        raise facade.ValidationError(
            "At least one of space_before, space_after, or line_spacing must be provided"
        )
    if space_before is not None and space_before < 0:
        raise facade.ValidationError("space_before must be >= 0")
    if space_after is not None and space_after < 0:
        raise facade.ValidationError("space_after must be >= 0")
    if line_spacing is not None and line_spacing <= 0:
        raise facade.ValidationError("line_spacing must be > 0")

    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)

    if table_cell is not None:
        row = int(table_cell.get("row", 0))
        col = int(table_cell.get("col", 0))
        tables = document.tables
        if not tables:
            raise facade.NotFoundError("Document has no tables")
        table = tables[0]
        if row < 0 or row >= len(table.rows):
            raise facade.ValidationError(f"row {row} out of range (0-{len(table.rows)-1})")
        if col < 0 or col >= len(table.columns):
            raise facade.ValidationError(f"col {col} out of range (0-{len(table.columns)-1})")
        cell_paragraphs = table.cell(row, col).paragraphs
        if paragraph_index < 0 or paragraph_index >= len(cell_paragraphs):
            raise facade.NotFoundError(
                f"paragraph_index {paragraph_index} out of range "
                f"(0-{len(cell_paragraphs)-1}) in cell ({row},{col})"
            )
        para = cell_paragraphs[paragraph_index]
    else:
        paragraphs = document.paragraphs
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            raise facade.NotFoundError(
                f"paragraph_index {paragraph_index} out of range (0-{len(paragraphs)-1})"
            )
        para = paragraphs[paragraph_index]

    pf = para.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log(
        "set_paragraph_format",
        resolved,
        extra={
            "paragraph_index": paragraph_index,
            "space_before": space_before,
            "space_after": space_after,
            "line_spacing": line_spacing,
        },
    )
    return {
        "paragraph_index": paragraph_index,
        "space_before_pt": space_before,
        "space_after_pt": space_after,
        "line_spacing": line_spacing,
        "table_cell": table_cell,
    }
