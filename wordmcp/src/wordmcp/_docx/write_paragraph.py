from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from wordmcp._docx import _facade


def add_paragraph(
    path: str,
    text: str,
    style: str | None = None,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    doc = facade._load_doc(resolved)

    if style is not None:
        try:
            doc.styles[style]
        except KeyError:
            raise facade.ValidationError(f"Unknown paragraph style: {style!r}")

    paragraph_count_before = len(doc.paragraphs)
    para = doc.add_paragraph(text, style=style)
    actual_style = para.style.name
    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("add_paragraph", resolved, extra={"style": style})
    return {"paragraph_index": paragraph_count_before, "style": actual_style, "text": text}


def add_heading(
    path: str,
    text: str,
    level: int = 1,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    if not (0 <= level <= 9):
        raise facade.ValidationError(f"level must be 0-9, got {level}")

    doc = facade._load_doc(resolved)
    paragraph_count_before = len(doc.paragraphs)
    para = doc.add_heading(text, level=level)
    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("add_heading", resolved, extra={"level": level})
    return {
        "paragraph_index": paragraph_count_before,
        "style": para.style.name,
        "level": level,
        "text": text,
    }


def add_page_break(path: str, confirm: bool = False) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    doc = facade._load_doc(resolved)

    paragraph_count_before = len(doc.paragraphs)
    doc.add_page_break()
    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("add_page_break", resolved)
    return {"inserted": True, "paragraph_index": paragraph_count_before}


def update_paragraph(path: str, paragraph_index: int, new_text: str, confirm: bool = False) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    doc = facade._load_doc(resolved)
    paragraphs = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        raise facade.NotFoundError(
            f"Paragraph index {paragraph_index} out of range (0-{len(paragraphs)-1})"
        )
    para = paragraphs[paragraph_index]
    clean_text = facade._CTRL_RE.sub("", new_text) if new_text else ""
    para.clear()
    para.add_run(clean_text)
    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("update_paragraph", resolved)
    return {"paragraph_index": paragraph_index, "text": clean_text}


def delete_paragraph(path: str, paragraph_index: int, confirm: bool = False) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    doc = facade._load_doc(resolved)
    paragraphs = doc.paragraphs
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        raise facade.NotFoundError(
            f"Paragraph index {paragraph_index} out of range (0-{len(paragraphs)-1})"
        )
    para = paragraphs[paragraph_index]
    text_preview = para.text[:50]
    para._element.getparent().remove(para._element)
    facade._atomic_save(doc, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("delete_paragraph", resolved)
    return {
        "deleted": True,
        "paragraph_index": paragraph_index,
        "text_preview": text_preview,
        "note": "Subsequent paragraph indices shift down by 1",
    }


def insert_paragraph(
    path: str,
    paragraph_index: int,
    text: str,
    style: str | None = None,
    confirm: bool = False,
) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    paragraphs = document.paragraphs

    if paragraph_index < 0 or paragraph_index > len(paragraphs):
        raise facade.ValidationError(
            f"paragraph_index {paragraph_index} out of range 0-{len(paragraphs)}"
        )

    style_obj = None
    if style is not None:
        try:
            style_obj = document.styles[style]
        except KeyError:
            raise facade.ValidationError(f"Unknown style: {style!r}")

    new_para = OxmlElement("w:p")
    new_run = OxmlElement("w:r")
    new_text_el = OxmlElement("w:t")
    clean_text = facade._CTRL_RE.sub("", text) if text else ""
    new_text_el.text = clean_text
    new_text_el.set(qn("xml:space"), "preserve")
    new_run.append(new_text_el)
    new_para.append(new_run)

    if style_obj is not None:
        ppr = OxmlElement("w:pPr")
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), style_obj.style_id)
        ppr.append(pstyle)
        new_para.insert(0, ppr)

    if paragraph_index == len(paragraphs):
        body = document.element.body
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(new_para)
        else:
            body.append(new_para)
    else:
        ref_elem = paragraphs[paragraph_index]._element
        ref_elem.getparent().insert(list(ref_elem.getparent()).index(ref_elem), new_para)

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("insert_paragraph", resolved, extra={"paragraph_index": paragraph_index})
    return {"inserted": True, "paragraph_index": paragraph_index, "text": clean_text}


def bulk_add_paragraphs(path: str, paragraphs: list, confirm: bool = False) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    if not paragraphs:
        raise facade.ValidationError("paragraphs must not be empty")
    if len(paragraphs) > 500:
        raise facade.ValidationError("max 500 paragraphs per call")

    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    try:
        for paragraph in paragraphs:
            text = paragraph.get("text", "")
            clean_text = facade._CTRL_RE.sub("", text) if text else ""
            style = paragraph.get("style")
            if style:
                try:
                    document.styles[style]
                except KeyError:
                    raise facade.ValidationError(f"Unknown style: {style!r}")
            document.add_paragraph(clean_text, style=style)
    except Exception:
        facade._evict_doc(resolved)
        raise

    paragraph_count = len(document.paragraphs)
    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("bulk_add_paragraphs", resolved)
    return {"added": len(paragraphs), "paragraph_count": paragraph_count}


def bulk_update_paragraphs(path: str, updates: list, confirm: bool = False) -> dict:
    facade = _facade()
    facade._check_write()
    facade._check_confirm(confirm)
    if not updates:
        raise facade.ValidationError("updates must not be empty")
    if len(updates) > 500:
        raise facade.ValidationError("max 500 updates per call")

    resolved = facade._check_path(path)
    document = facade._load_doc(resolved)
    paragraphs = document.paragraphs
    results = []
    errors = []

    for i, item in enumerate(updates):
        try:
            idx = item.get("paragraph_index")
            if idx is None:
                raise facade.ValidationError("paragraph_index is required")
            idx = int(idx)
            if idx < 0 or idx >= len(paragraphs):
                raise facade.NotFoundError(
                    f"paragraph_index {idx} out of range (0-{len(paragraphs) - 1})"
                )
            new_text = item.get("new_text", "")
            clean_text = facade._CTRL_RE.sub("", new_text) if new_text else ""
            para = paragraphs[idx]
            para.clear()
            para.add_run(clean_text)
            results.append({"index": i, "paragraph_index": idx, "status": "updated"})
        except Exception as exc:  # noqa: BLE001
            errors.append({"index": i, "error": f"Operation failed ({type(exc).__name__})"})

    facade._atomic_save(document, resolved)
    facade._evict_doc(resolved)
    facade._audit_log("bulk_update_paragraphs", resolved)
    return {
        "updated": len(results),
        "failed": len(errors),
        "results": results,
        "errors": errors,
    }
