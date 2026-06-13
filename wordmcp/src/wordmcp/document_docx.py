"""
python-docx data-access layer for Word MCP Phase 1.
No FastMCP / MCP imports — pure python-docx + stdlib only.

STDOUT DISCIPLINE (CRITICAL — MCP PROTOCOL INTEGRITY):
  - NO print() calls anywhere in this module.
  - NO logging to sys.stdout.
  - NO warnings.warn() calls.
  - All diagnostic output via: logger = logging.getLogger(__name__)
    (root logger configured for sys.stderr in server.py)
  - Violation corrupts JSON-RPC framing on stdout and breaks all MCP clients.
"""
from __future__ import annotations

import logging
import os
import re
from pathlib import Path

import docx
from docx.document import Document as DocxDocument

from wordmcp._docx import core as _core
from wordmcp._docx import context as _context
from wordmcp._docx import export as _export
from wordmcp._docx import read as _read
from wordmcp._docx import review as _review
from wordmcp._docx import style as _style
from wordmcp._docx import write as _write

try:
    from wordmcp import __version__ as _wordmcp_version
except ImportError:
    _wordmcp_version = "unknown"

logger = logging.getLogger(__name__)


IMAGE_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)

_ALLOWED_IMAGE_EXTENSIONS = frozenset(
    {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"}
)

_DEFAULT_MAX_TEXT_CHARS = 50000

_CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

try:
    _MAX_DOCX_BYTES: int = int(os.getenv("WORD_MAX_FILE_MB", "50")) * 1024 * 1024
except ValueError:
    _MAX_DOCX_BYTES = 50 * 1024 * 1024

try:
    _DOC_CACHE_MAX: int = max(1, int(os.getenv("WORD_MAX_CACHE_DOCS", "20")))
except ValueError:
    _DOC_CACHE_MAX = 20


class WordMCPError(Exception):
    """Base class for all Word MCP errors."""


class ValidationError(WordMCPError):
    """Input did not pass allowlist / type / format validation."""


class NotAllowedError(WordMCPError):
    """Operation blocked by WORD_ENABLE_WRITE gate or confirm=False."""


class NotFoundError(WordMCPError):
    """Requested resource (paragraph_index, table_index) does not exist."""


class FileError(WordMCPError):
    """Wraps python-docx FileNotFoundError or PackageNotFoundError."""


class OfficeCOMError(WordMCPError):
    """Raised when a Word COM automation call fails."""


_doc_cache = _core._doc_cache
_do_find_replace = _core._do_find_replace
_get_list_num_id = _core._get_list_num_id
_outline_level = _core._outline_level
_paragraph_to_dict = _core._paragraph_to_dict

read_document = _read.read_document
get_document_metadata = _read.get_document_metadata
list_paragraphs = _read.list_paragraphs
read_paragraph = _read.read_paragraph
list_tables = _read.list_tables
read_table = _read.read_table
list_headings = _read.list_headings
read_section = _read.read_section
search_text = _read.search_text
get_document_context = _context.get_document_context

export_as_text = _export.export_as_text
export_to_markdown = _export.export_to_markdown
get_headers_footers = _export.get_headers_footers

list_styles = _style.list_styles
apply_style = _style.apply_style

manage_comments = _review.manage_comments
add_hyperlink = _review.add_hyperlink
add_footnote = _review.add_footnote

create_document = _write.create_document
add_paragraph = _write.add_paragraph
add_heading = _write.add_heading
add_page_break = _write.add_page_break
add_table = _write.add_table
insert_image = _write.insert_image
find_replace = _write.find_replace
save = _write.save
update_paragraph = _write.update_paragraph
delete_paragraph = _write.delete_paragraph
update_table_cell = _write.update_table_cell
set_document_properties = _write.set_document_properties
add_list = _write.add_list
insert_paragraph = _write.insert_paragraph
bulk_add_paragraphs = _write.bulk_add_paragraphs
bulk_update_paragraphs = _write.bulk_update_paragraphs
bulk_update_table_cells = _write.bulk_update_table_cells
set_paragraph_format = _write.set_paragraph_format


def _check_path(path: str, must_exist: bool = True) -> Path:
    return _core._check_path(path, must_exist)


def _check_image_path(path: str) -> Path:
    return _core._check_image_path(path)


def _check_evidence_path(path: str) -> Path:
    return _core._check_evidence_path(path)


def _check_write() -> None:
    _core._check_write()


def _check_confirm(confirm: bool) -> None:
    _core._check_confirm(confirm)


def _audit_log(op: str, path: Path | str, extra: dict | None = None) -> None:
    _core._audit_log(op, path, extra)


def _atomic_save(doc: DocxDocument, target: Path) -> None:
    _core._atomic_save(doc, target)


def _load_doc(path: Path) -> DocxDocument:
    return _core._load_doc(path)


def _evict_doc(path: Path) -> None:
    _core._evict_doc(path)


def capabilities() -> dict:
    """Return metadata about this MCP server phase, backend, and available tools."""
    return {
        "phase": "1.0",
        "backend": "python-docx",
        "version": _wordmcp_version,
        "python_docx_version": docx.__version__,
        "tools": [
            "capabilities",
            "read_document",
            "get_document_metadata",
            "list_paragraphs",
            "read_paragraph",
            "list_tables",
            "read_table",
            "add_paragraph",
            "add_heading",
            "add_page_break",
            "add_table",
            "insert_image",
            "find_replace",
            "save",
        ],
        "governance": {
            "allowlist_roots_env": "WORD_ALLOWLIST_ROOTS",
            "enable_write_env": "WORD_ENABLE_WRITE",
            "max_text_chars_env": "WORD_MAX_TEXT_CHARS",
            "WORD_MAX_FILE_MB": os.getenv("WORD_MAX_FILE_MB", "50"),
            "WORD_MAX_CACHE_DOCS": os.getenv("WORD_MAX_CACHE_DOCS", "20"),
        },
    }
