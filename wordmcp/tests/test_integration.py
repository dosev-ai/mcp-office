"""
Integration tests for wordmcp — Sprint A.
All tests use @pytest.mark.integration.
All file I/O uses tmp_path pytest fixture.
Imports ONLY from wordmcp.document_docx (no server.py).
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from wordmcp.document_docx import (
    ValidationError,
    add_heading,
    add_paragraph,
    add_table,
    export_as_text,
    find_replace,
    get_document_metadata,
    list_paragraphs,
    list_tables,
    read_document,
    read_table,
    save,
)


# ---------------------------------------------------------------------------
# In-file fixtures (no conftest dependency for integration-specific state)
# ---------------------------------------------------------------------------

@pytest.fixture
def int_docx(tmp_path, monkeypatch):
    """Writable .docx with WORD_ALLOWLIST_ROOTS and WORD_ENABLE_WRITE set."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "int_test.docx"
    Document().save(str(path))
    return str(path)


@pytest.fixture
def int_docx_with_table(tmp_path, monkeypatch):
    """Writable .docx containing one 2x3 table."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "int_table.docx"
    d = Document()
    tbl = d.add_table(rows=2, cols=3, style="Table Grid")
    for c_idx, val in enumerate(["Name", "Age", "City"]):
        tbl.rows[0].cells[c_idx].text = val
    for c_idx, val in enumerate(["Alice", "30", "London"]):
        tbl.rows[1].cells[c_idx].text = val
    d.save(str(path))
    return str(path)


# ===========================================================================
# 12 integration tests
# ===========================================================================


@pytest.mark.integration
def test_full_roundtrip_create_read(int_docx):
    """Create a paragraph then read_document reflects it."""
    add_paragraph(int_docx, "Integration paragraph", confirm=True)
    result = read_document(int_docx)
    assert result["paragraph_count"] >= 1
    d = Document(int_docx)
    assert any("Integration paragraph" == p.text for p in d.paragraphs)


@pytest.mark.integration
def test_heading_roundtrip(int_docx):
    """add_heading levels 1-3 produce correct styles in list_paragraphs."""
    add_heading(int_docx, "Chapter One", level=1, confirm=True)
    add_heading(int_docx, "Section 1.1", level=2, confirm=True)
    add_heading(int_docx, "Sub-section 1.1.1", level=3, confirm=True)
    paras = list_paragraphs(int_docx)
    styles = [p["style"] for p in paras]
    assert "Heading 1" in styles
    assert "Heading 2" in styles
    assert "Heading 3" in styles


@pytest.mark.integration
def test_table_roundtrip(int_docx):
    """add_table 3x3 then read_table[0] returns correct cells."""
    data = [["A", "B", "C"], ["D", "E", "F"], ["G", "H", "I"]]
    add_table(int_docx, rows=3, cols=3, data=data, confirm=True)
    result = read_table(int_docx, 0)
    assert result["rows"] == 3
    assert result["cols"] == 3
    assert result["data"][0] == ["A", "B", "C"]
    assert result["data"][2] == ["G", "H", "I"]


@pytest.mark.integration
def test_find_replace_multi_occurrence(int_docx):
    """find_replace replaces all occurrences of a word across multiple paragraphs."""
    add_paragraph(int_docx, "foo bar foo", confirm=True)
    add_paragraph(int_docx, "another foo here", confirm=True)
    result = find_replace(int_docx, "foo", "baz", confirm=True)
    assert result["replacements_made"] == 3
    d = Document(int_docx)
    all_text = " ".join(p.text for p in d.paragraphs)
    assert "foo" not in all_text
    assert "baz" in all_text


@pytest.mark.integration
def test_save_to_new_path(int_docx, tmp_path, monkeypatch):
    """save() in-place after mutations — file still readable with correct content."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    add_paragraph(int_docx, "Content for save", confirm=True)
    result = save(int_docx, confirm=True)
    assert result["saved"] is True
    assert Path(int_docx).exists()
    # Reload from disk to confirm content survives the save
    d = Document(int_docx)
    assert any("Content for save" == p.text for p in d.paragraphs)


@pytest.mark.integration
def test_get_metadata_keys_present(int_docx):
    """get_document_metadata returns all expected keys for a fresh document."""
    result = get_document_metadata(int_docx)
    expected_keys = {
        "path", "title", "author", "subject", "keywords",
        "category", "last_modified_by", "revision", "version",
        "created", "modified",
    }
    assert expected_keys.issubset(result.keys())


@pytest.mark.integration
def test_export_as_text_content(int_docx):
    """export_as_text includes known text added to the document."""
    add_paragraph(int_docx, "Known integration text", confirm=True)
    result = export_as_text(int_docx)
    assert "Known integration text" in result["text"]
    assert isinstance(result["truncated"], bool)


@pytest.mark.integration
def test_empty_doc_reads_gracefully(int_docx):
    """Blank .docx: list_paragraphs returns a list, list_tables returns []."""
    paras = list_paragraphs(int_docx)
    tables = list_tables(int_docx)
    assert isinstance(paras, list)
    assert isinstance(tables, list)
    assert tables == []


@pytest.mark.integration
def test_large_doc_paragraph_count(int_docx):
    """Adding 100 paragraphs results in exactly 100 paragraphs visible on disk."""
    for i in range(100):
        add_paragraph(int_docx, f"Paragraph {i}", confirm=True)
    d = Document(int_docx)
    # Only paragraphs with text content (added by us)
    added = [p for p in d.paragraphs if p.text.startswith("Paragraph")]
    assert len(added) == 100


@pytest.mark.integration
def test_allowlist_enforcement_real_path(int_docx, tmp_path, monkeypatch):
    """Real path inside allowlist succeeds; path outside allowlist raises."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    # Inside allowlist — succeeds
    result = read_document(int_docx)
    assert "paragraph_count" in result
    # Outside allowlist — raises
    with pytest.raises(ValidationError, match="not in allowlist"):
        read_document("/absolute/outside/secret.docx")


@pytest.mark.integration
def test_write_guard_real_path(int_docx, tmp_path, monkeypatch):
    """WORD_ENABLE_WRITE=true + confirm=True succeeds; confirm=False rejects."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    # confirm=True succeeds
    result = add_paragraph(int_docx, "guarded write", confirm=True)
    assert result["text"] == "guarded write"
    # confirm=False rejects
    with pytest.raises(ValidationError, match="confirm=True required"):
        add_paragraph(int_docx, "should fail", confirm=False)


@pytest.mark.integration
def test_multiple_tables_index(int_docx):
    """Two tables added; read_table[0] and read_table[1] return distinct data."""
    add_table(int_docx, rows=2, cols=2, data=[["A", "B"], ["C", "D"]], confirm=True)
    add_table(int_docx, rows=2, cols=2, data=[["X", "Y"], ["Z", "W"]], confirm=True)
    r0 = read_table(int_docx, 0)
    r1 = read_table(int_docx, 1)
    assert r0["data"][0] == ["A", "B"]
    assert r1["data"][0] == ["X", "Y"]
    assert r0["index"] == 0
    assert r1["index"] == 1
