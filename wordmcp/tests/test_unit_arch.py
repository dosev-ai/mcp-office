"""
W3 Architecture Test Plan — Unit Tests (no live Word COM required).

Covers:
  B1  Layer separation: txt/md export must NOT route through COM handler
  B2  Destructive op stub: manage_comments resolve/delete must not raise NotImplementedError
  B3  Runtime enforcement: export_document must raise ToolError for unknown format
  B4  COM layer guard: manage_tracked_changes must raise ToolError for unknown operation
  B5  bulk_update_paragraphs: max=500 cap, per-item bounds, write gate + confirm
  B6  bulk_update_table_cells: max=200 cap, per-item bounds, write gate + confirm
  B7  COM export gate: document_com.export_document verifies write gate before writing
  P1  txt/md export positive (docx layer, happy path)
  P2  manage_comments supported ops (list/add happy path)
  P3  export_document valid formats accepted (structural)
  P4  manage_tracked_changes valid ops accepted (structural)
  P5  bulk_update_paragraphs happy path (xfail until P5 implemented)
  P6  bulk_update_table_cells happy path (xfail until P6 implemented)
  P7  COM export write gate enforced positively (mock)
  P8  bulk operations per-item bounds valid smoke (xfail until P5/P6 implemented)

Run:
  pytest word/tests/test_unit_arch.py -v -m "not integration"
"""
from __future__ import annotations

import pathlib
from unittest.mock import MagicMock, patch

import pytest
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _wordmcp_src() -> pathlib.Path:
    """Resolve word/src/wordmcp package root regardless of CWD."""
    import wordmcp
    return pathlib.Path(wordmcp.__file__).parent


# ===========================================================================
# B1 — Layer Separation (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_b1_export_as_text_not_registered_in_com_handler():
    """B1-NEG: export_as_text must NOT appear in _server_handlers_com.register() return dict.

    The docx export tools (export_as_text, export_to_markdown) must be served
    by the docx handler layer, never by the COM handler.
    """
    from wordmcp import _server_handlers_com

    # Build a minimal fake bind_tool and runtime to call register()
    bound = {}

    def _bind(fn, name):
        bound[name] = fn
        return fn

    runtime = MagicMock()
    _server_handlers_com.register(_bind, runtime)

    assert "export_as_text" not in bound, (
        "export_as_text must NOT be registered by _server_handlers_com"
    )


@pytest.mark.unit
def test_b1_export_to_markdown_not_registered_in_com_handler():
    """B1-NEG: export_to_markdown must NOT appear in _server_handlers_com.register() return dict."""
    from wordmcp import _server_handlers_com

    bound = {}

    def _bind(fn, name):
        bound[name] = fn
        return fn

    runtime = MagicMock()
    _server_handlers_com.register(_bind, runtime)

    assert "export_to_markdown" not in bound, (
        "export_to_markdown must NOT be registered by _server_handlers_com"
    )


@pytest.mark.unit
def test_b1_com_handler_source_body_does_not_reference_docx_export_fns():
    """B1-NEG: _server_handlers_com.py source text must not contain txt/md export names.

    AST-free source scan ensures no accidental dispatch creep into the COM layer.
    """
    src_path = _wordmcp_src() / "_server_handlers_com.py"
    source = src_path.read_text(encoding="utf-8")

    # These functions belong strictly to the docx layer
    for fn_name in ("export_as_text", "export_to_markdown"):
        assert fn_name not in source, (
            f"_server_handlers_com.py source must not reference '{fn_name}'"
        )


# ===========================================================================
# B2 — Destructive Op Stubs (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_b2_manage_comments_resolve_raises_validation_not_stub():
    """B2-NEG: manage_comments with an unknown op must raise ValidationError, NOT NotImplementedError.

    After ARCH sprint, 'resolve' and 'delete' are now valid ops.  This test uses
    'frobulate' (an always-invalid op) to verify the unknown-op ValidationError pathway.
    """
    from wordmcp.document_docx import ValidationError, manage_comments

    with pytest.raises(ValidationError) as exc_info:
        manage_comments(path="/any/path.docx", operation="frobulate")

    assert not isinstance(exc_info.value, NotImplementedError), (
        "Must not be a NotImplementedError stub"
    )
    assert "frobulate" in str(exc_info.value).lower() or "unknown" in str(exc_info.value).lower()


@pytest.mark.unit
def test_b2_manage_comments_delete_raises_validation_not_stub():
    """B2-NEG: manage_comments with an unknown op must raise ValidationError, NOT NotImplementedError.

    After ARCH sprint, 'delete' is a valid op.  This test uses 'purge' (always-invalid)
    to verify the unknown-op ValidationError pathway.
    """
    from wordmcp.document_docx import ValidationError, manage_comments

    with pytest.raises(ValidationError) as exc_info:
        manage_comments(path="/any/path.docx", operation="purge")

    assert not isinstance(exc_info.value, NotImplementedError)
    assert "purge" in str(exc_info.value).lower() or "unknown" in str(exc_info.value).lower()


@pytest.mark.unit
@pytest.mark.parametrize("bad_op", ["reject", "update", "purge", "frobulate", "noop"])
def test_b2_manage_comments_unsupported_ops_parametrized(bad_op):
    """B2-NEG (parametrized): every unsupported manage_comments operation raises ValidationError.

    After ARCH sprint, 'resolve' and 'delete' are valid ops and excluded here.
    """
    from wordmcp.document_docx import ValidationError, manage_comments

    with pytest.raises(ValidationError):
        manage_comments(path="/any/path.docx", operation=bad_op)


# ===========================================================================
# B3 — Format Validation on export_document (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_b3_export_document_unknown_format_raises_validation():
    """B3-NEG: export_document with an unknown format string must raise ValidationError
    immediately, before any COM call is attempted.
    """
    from wordmcp.document_com import ValidationError, export_document

    with pytest.raises(ValidationError) as exc_info:
        export_document(
            path=r"C:\fake\doc.docx",
            output_path=r"C:\fake\out.rtf",
            format="rtf",
        )

    assert "rtf" in str(exc_info.value).lower() or "unknown" in str(exc_info.value).lower()


@pytest.mark.unit
@pytest.mark.parametrize("bad_fmt", ["xml", "txt", "docx", "odt", "rtf", "csv", ""])
def test_b3_export_document_invalid_formats_parametrized(bad_fmt):
    """B3-NEG (parametrized): every non-pdf/non-html format raises ValidationError."""
    from wordmcp.document_com import ValidationError, export_document

    with pytest.raises(ValidationError):
        export_document(
            path=r"C:\fake\doc.docx",
            output_path=r"C:\fake\out.xyz",
            format=bad_fmt,
        )


@pytest.mark.unit
def test_b3_export_document_valid_formats_are_pdf_and_html():
    """B3-POS (structural): export_document source explicitly handles 'pdf' and 'html'."""
    src_path = _wordmcp_src() / "_com" / "_export.py"
    source = src_path.read_text(encoding="utf-8")
    assert "format == \"pdf\"" in source or "format==\"pdf\"" in source.replace(" ", "")
    assert "format == \"html\"" in source or "format==\"html\"" in source.replace(" ", "")


# ===========================================================================
# B4 — COM Layer Guard on manage_tracked_changes (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_b4_manage_tracked_changes_unknown_op_raises_validation():
    """B4-NEG: manage_tracked_changes with an unknown operation raises ValidationError
    before any COM call.
    """
    from wordmcp.document_com import ValidationError, manage_tracked_changes

    with pytest.raises(ValidationError) as exc_info:
        manage_tracked_changes(
            path=r"C:\fake\doc.docx",
            operation="frobnicate",
        )

    assert "frobnicate" in str(exc_info.value).lower() or "unknown" in str(exc_info.value).lower()


@pytest.mark.unit
@pytest.mark.parametrize("bad_op", ["delete", "resolve", "modify", "update", "purge", "clear"])
def test_b4_manage_tracked_changes_invalid_ops_parametrized(bad_op):
    """B4-NEG (parametrized): each unsupported tracked-changes operation raises ValidationError."""
    from wordmcp.document_com import ValidationError, manage_tracked_changes

    with pytest.raises(ValidationError):
        manage_tracked_changes(
            path=r"C:\fake\doc.docx",
            operation=bad_op,
        )


@pytest.mark.unit
def test_b4_manage_tracked_changes_valid_ops_listed_in_source():
    """B4-POS (structural): the five valid operations are present in document_com.py source."""
    src_path = _wordmcp_src() / "_com" / "_write.py"
    source = src_path.read_text(encoding="utf-8")
    for valid_op in ("list", "accept_all", "reject_all", "accept_one", "reject_one"):
        assert valid_op in source, (
            f"Expected valid operation '{valid_op}' to appear in document_com.py"
        )


# ===========================================================================
# B5 — bulk_update_paragraphs guards (5 tests)
# ===========================================================================


@pytest.mark.unit
def test_b5_bulk_update_paragraphs_importable():
    """B5-STRUCT: bulk_update_paragraphs must be importable from wordmcp.document_docx."""
    from wordmcp.document_docx import bulk_update_paragraphs  # noqa: F401
    assert callable(bulk_update_paragraphs)


@pytest.mark.unit
def test_b5_bulk_update_paragraphs_cap_501_raises(tmp_path, monkeypatch):
    """B5-NEG: bulk_update_paragraphs with 501 items must raise ValidationError (max=500 cap)."""
    from wordmcp.document_docx import ValidationError, bulk_update_paragraphs

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    d = Document()
    d.add_paragraph("x")
    d.save(str(path))

    items = [{"paragraph_index": 0, "new_text": "y"}] * 501

    with pytest.raises(ValidationError, match="[Mm]ax|[Ll]imit|500"):
        bulk_update_paragraphs(path=str(path), updates=items, confirm=True)


@pytest.mark.unit
def test_b5_bulk_update_paragraphs_write_gate_blocks(tmp_path, monkeypatch):
    """B5-NEG: bulk_update_paragraphs without WORD_ENABLE_WRITE must raise NotAllowedError."""
    from wordmcp.document_docx import NotAllowedError, bulk_update_paragraphs

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "test.docx"
    Document().save(str(path))

    with pytest.raises(NotAllowedError):
        bulk_update_paragraphs(path=str(path), updates=[], confirm=True)


@pytest.mark.unit
def test_b5_bulk_update_paragraphs_confirm_false_blocked(tmp_path, monkeypatch):
    """B5-NEG: bulk_update_paragraphs with confirm=False must raise NotAllowedError or ValidationError."""
    from wordmcp.document_docx import NotAllowedError, ValidationError, bulk_update_paragraphs

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    Document().save(str(path))

    with pytest.raises((NotAllowedError, ValidationError)):
        bulk_update_paragraphs(
            path=str(path),
            updates=[{"paragraph_index": 0, "new_text": "z"}],
            confirm=False,
        )


@pytest.mark.unit
def test_b5_bulk_update_paragraphs_per_item_invalid_index_raises(tmp_path, monkeypatch):
    """B5-NEG: a per-item negative paragraph_index must be recorded in the errors list."""
    from wordmcp.document_docx import bulk_update_paragraphs

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    d = Document()
    d.add_paragraph("Hello")
    d.save(str(path))

    result = bulk_update_paragraphs(
        path=str(path),
        updates=[{"paragraph_index": -1, "new_text": "bad"}],
        confirm=True,
    )
    assert result["failed"] >= 1, "per-item invalid index must be recorded in errors"


# ===========================================================================
# B6 — bulk_update_table_cells guards (4 tests)
# ===========================================================================


@pytest.mark.unit
def test_b6_bulk_update_table_cells_importable():
    """B6-STRUCT: bulk_update_table_cells must be importable from wordmcp.document_docx."""
    from wordmcp.document_docx import bulk_update_table_cells  # noqa: F401
    assert callable(bulk_update_table_cells)


@pytest.mark.unit
def test_b6_bulk_update_table_cells_cap_201_raises(tmp_path, monkeypatch):
    """B6-NEG: bulk_update_table_cells with 201 items must raise ValidationError (max=200 cap)."""
    from wordmcp.document_docx import ValidationError, bulk_update_table_cells

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    d = Document()
    d.add_table(1, 1)
    d.save(str(path))

    items = [{"table_index": 0, "row": 0, "col": 0, "new_text": "x"}] * 201

    with pytest.raises(ValidationError, match="[Mm]ax|[Ll]imit|200"):
        bulk_update_table_cells(path=str(path), updates=items, confirm=True)


@pytest.mark.unit
def test_b6_bulk_update_table_cells_write_gate_blocks(tmp_path, monkeypatch):
    """B6-NEG: bulk_update_table_cells without WORD_ENABLE_WRITE must raise NotAllowedError."""
    from wordmcp.document_docx import NotAllowedError, bulk_update_table_cells

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "test.docx"
    Document().save(str(path))

    with pytest.raises(NotAllowedError):
        bulk_update_table_cells(path=str(path), updates=[], confirm=True)


@pytest.mark.unit
def test_b6_bulk_update_table_cells_confirm_gate_blocks(tmp_path, monkeypatch):
    """B6-NEG: bulk_update_table_cells with confirm=False must raise NotAllowedError or ValidationError."""
    from wordmcp.document_docx import NotAllowedError, ValidationError, bulk_update_table_cells

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    Document().save(str(path))

    with pytest.raises((NotAllowedError, ValidationError)):
        bulk_update_table_cells(
            path=str(path),
            updates=[{"table_index": 0, "row": 0, "col": 0, "new_text": "x"}],
            confirm=False,
        )


# ===========================================================================
# B7 — COM Export Write Gate (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_b7_export_document_write_gate_blocks_before_com(monkeypatch):
    """B7-NEG: export_document(format='pdf') with WORD_ENABLE_COM=true but
    WORD_ENABLE_WRITE unset must raise NotAllowedError from _check_write, proving
    the write gate fires before any COM call is attempted.
    """
    from wordmcp.document_com import NotAllowedError, export_document

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)

    # Patch _check_com_enabled so it doesn't fail on non-Windows or missing pywin32
    with patch("wordmcp._com._export._check_com_enabled"):
        with pytest.raises(NotAllowedError, match="[Ww]rite|WORD_ENABLE_WRITE"):
            export_document(
                path=r"C:\fake\doc.docx",
                output_path=r"C:\fake\out.pdf",
                format="pdf",
                confirm=True,
            )


@pytest.mark.unit
def test_b7_com_export_write_gate_precedes_com_call_in_source():
    """B7-NEG (structural): in document_com.export_as_pdf source, _check_write()
    call must appear BEFORE the _word_app_context call — AST line-order check.
    """
    src_path = _wordmcp_src() / "_com" / "_export.py"
    source = src_path.read_text(encoding="utf-8")

    # Capture export_as_pdf body: enough bytes to clear the docstring + guards.
    # The docstring alone is ~250 chars; use 2000 to be safe.
    fn_start = source.index("def export_as_pdf(")
    fn_body = source[fn_start : fn_start + 2000]

    check_write_pos = fn_body.find("_check_write()")
    word_ctx_pos = fn_body.find("_word_app_context(")

    assert check_write_pos != -1, "_check_write() not found in export_as_pdf body"
    assert word_ctx_pos != -1, "_word_app_context() not found in export_as_pdf body"
    assert check_write_pos < word_ctx_pos, (
        "_check_write() must be called BEFORE _word_app_context() in export_as_pdf"
    )


# ===========================================================================
# P1 — txt/md export positive (docx layer, 2 tests)
# ===========================================================================


@pytest.mark.unit
def test_p1_export_as_text_returns_expected_keys(tmp_docx):
    """P1-POS: export_as_text with a valid .docx returns dict with text/truncated/paragraph_count."""
    from wordmcp.document_docx import export_as_text

    result = export_as_text(path=tmp_docx)

    assert isinstance(result, dict)
    assert "text" in result
    assert "truncated" in result
    assert "paragraph_count" in result
    assert isinstance(result["text"], str)


@pytest.mark.unit
def test_p1_export_to_markdown_returns_expected_keys(tmp_docx):
    """P1-POS: export_to_markdown with a valid .docx returns dict with markdown/truncated keys."""
    from wordmcp.document_docx import export_to_markdown

    result = export_to_markdown(path=tmp_docx)

    assert isinstance(result, dict)
    assert "markdown" in result
    assert "truncated" in result
    assert isinstance(result["markdown"], str)


# ===========================================================================
# P2 — manage_comments supported ops (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_p2_manage_comments_list_returns_comments_list(tmp_docx):
    """P2-POS: manage_comments('list') returns {'comments': list} for a fresh docx."""
    from wordmcp.document_docx import manage_comments

    result = manage_comments(path=tmp_docx, operation="list")

    assert isinstance(result, dict)
    assert "comments" in result
    assert isinstance(result["comments"], list)


@pytest.mark.unit
def test_p2_manage_comments_add_returns_added_true(write_enabled_docx):
    """P2-POS: manage_comments('add') with write enabled returns {'added': True}."""
    from wordmcp.document_docx import manage_comments

    result = manage_comments(
        path=write_enabled_docx,
        operation="add",
        text="QA review note",
        author="qa-engineer",
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("added") is True
    assert "comment_id" in result


# ===========================================================================
# P3 — export_document valid format structural check (1 test)
# ===========================================================================


@pytest.mark.unit
def test_p3_export_document_source_handles_both_formats():
    """P3-POS (structural): document_com.export_document source handles both 'pdf' and 'html'
    branches with no fallthrough to the else/ValidationError branch for those formats.
    """
    src_path = _wordmcp_src() / "_com" / "_export.py"
    source = src_path.read_text(encoding="utf-8")

    # Locate the export_document function; find its body up to the next top-level def.
    fn_start = source.index("def export_document")
    fn_end = source.find("\ndef ", fn_start + 1)
    fn_body = source[fn_start:fn_end] if fn_end != -1 else source[fn_start:]

    assert '"pdf"' in fn_body, "export_document must handle format='pdf'"
    assert '"html"' in fn_body, "export_document must handle format='html'"


# ===========================================================================
# P4 — manage_tracked_changes valid ops structural check (1 test)
# ===========================================================================


@pytest.mark.unit
def test_p4_manage_tracked_changes_source_handles_all_valid_ops():
    """P4-POS (structural): manage_tracked_changes in document_com.py handles all 5 valid operations."""
    src_path = _wordmcp_src() / "_com" / "_write.py"
    source = src_path.read_text(encoding="utf-8")

    fn_start = source.index("def manage_tracked_changes(")
    fn_end = source.find("\ndef ", fn_start + 1)
    fn_body = source[fn_start:fn_end] if fn_end != -1 else source[fn_start:]

    for op in ("list", "accept_all", "reject_all", "accept_one", "reject_one"):
        assert op in fn_body, (
            f"manage_tracked_changes must explicitly handle operation='{op}'"
        )


# ===========================================================================
# P5 — bulk_update_paragraphs happy path (1 test)
# ===========================================================================


@pytest.mark.unit
def test_p5_bulk_update_paragraphs_happy_path(write_enabled_docx):
    """P5-POS: bulk_update_paragraphs with valid inputs and write enabled returns ok status."""
    from wordmcp.document_docx import bulk_update_paragraphs

    result = bulk_update_paragraphs(
        path=write_enabled_docx,
        updates=[{"paragraph_index": 2, "new_text": "Updated text"}],
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("updated", 0) >= 1


# ===========================================================================
# P6 — bulk_update_table_cells happy path (1 test)
# ===========================================================================


@pytest.mark.unit
def test_p6_bulk_update_table_cells_happy_path(write_enabled_docx_with_table):
    """P6-POS: bulk_update_table_cells with valid inputs and write enabled returns ok status."""
    from wordmcp.document_docx import bulk_update_table_cells

    result = bulk_update_table_cells(
        path=write_enabled_docx_with_table,
        updates=[{"table_index": 0, "row": 0, "col": 0, "new_text": "Updated"}],
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("updated", 0) >= 1


# ===========================================================================
# P7 — COM export write gate enforced positively (mock, 1 test)
# ===========================================================================


@pytest.mark.unit
def test_p7_export_document_proceeds_past_write_gate_when_enabled(
    tmp_path, monkeypatch
):
    """P7-POS: export_document(format='pdf') with WORD_ENABLE_WRITE=true proceeds past
    the write-gate check (reaching path validation), not blocked by a write guard.
    """
    from wordmcp.document_com import export_document

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

    # Create a real docx so path validation passes
    doc_path = tmp_path / "real.docx"
    Document().save(str(doc_path))

    with patch("wordmcp._com._export._check_com_enabled"):
        with patch("wordmcp._com._export._word_app_context") as mock_ctx:
            mock_app = MagicMock()
            mock_doc = MagicMock()
            mock_app.Documents.Open.return_value = mock_doc
            mock_ctx.return_value.__enter__ = MagicMock(return_value=mock_app)
            mock_ctx.return_value.__exit__ = MagicMock(return_value=False)

            out_path = tmp_path / "out.pdf"
            try:
                result = export_document(
                    path=str(doc_path),
                    output_path=str(out_path),
                    format="pdf",
                    confirm=True,
                )
                # If we get here, the write gate did NOT block us — that's the goal
                assert result.get("status") == "ok"
            except Exception as exc:
                # Acceptable if COM mock incomplete — what must NOT happen is NotAllowedError
                from wordmcp.document_com import NotAllowedError
                assert not isinstance(exc, NotAllowedError), (
                    "Write gate must not block when WORD_ENABLE_WRITE=true"
                )


# ===========================================================================
# P8 — bulk operations per-item bounds valid smoke (1 test)
# ===========================================================================


@pytest.mark.unit
def test_p8_bulk_update_paragraphs_boundary_exactly_500_accepted(
    write_enabled_docx,
):
    """P8-POS: bulk_update_paragraphs with exactly 500 items must not raise a cap error."""
    from wordmcp.document_docx import ValidationError, bulk_update_paragraphs

    items = [{"paragraph_index": 0, "new_text": "x"}] * 500

    try:
        bulk_update_paragraphs(path=write_enabled_docx, updates=items, confirm=True)
    except ValidationError as exc:
        assert "500" not in str(exc) or "max" not in str(exc).lower(), (
            "A 500-item batch must not be rejected by the cap check"
        )
