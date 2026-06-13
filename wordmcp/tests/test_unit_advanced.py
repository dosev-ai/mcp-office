"""
Unit tests — advanced features: Sprint C (manage_comments/tracked_changes/export_document),
Batch A (list_headings/read_section/search_text/list_styles/create_document),
Batch B (apply_style/update_paragraph/delete_paragraph/update_table_cell/set_document_properties).
"""
from __future__ import annotations

import pytest
from docx import Document

from wordmcp.document_docx import (
    NotAllowedError,
    NotFoundError,
    ValidationError,
    apply_style,
    create_document,
    delete_paragraph,
    get_document_metadata,
    list_headings,
    list_styles,
    manage_comments,
    read_section,
    search_text,
    set_document_properties,
    update_paragraph,
    update_table_cell,
)


# ===========================================================================
# Sprint C — manage_comments, manage_tracked_changes, export_document
# ===========================================================================


@pytest.mark.unit
def test_manage_comments_list_returns_stub(tmp_docx):
    """manage_comments 'list' returns dict with 'comments' key."""

    result = manage_comments(path=tmp_docx, operation="list")
    assert isinstance(result, dict)
    assert "comments" in result


@pytest.mark.unit
def test_manage_comments_add_requires_confirm(write_enabled_docx):
    """manage_comments 'add' without confirm=True raises ValidationError."""

    with pytest.raises(ValidationError, match="confirm=True required"):
        manage_comments(path=write_enabled_docx, operation="add")


@pytest.mark.unit
def test_manage_comments_invalid_op():
    """manage_comments with unknown operation raises ValidationError."""

    with pytest.raises(ValidationError, match="Unknown operation"):
        manage_comments(path="fake.docx", operation="purge")


@pytest.mark.unit
def test_manage_tracked_changes_router_routes_list():
    """manage_tracked_changes 'list' op routes to list_tracked_changes."""
    from unittest.mock import patch
    from wordmcp import document_com
    from wordmcp._com import _write as _document_com_write

    with patch.object(
        _document_com_write, "list_tracked_changes", return_value={"revisions": [], "count": 0}
    ) as mock_fn:
        result = document_com.manage_tracked_changes(path="fake.docx", operation="list")
    mock_fn.assert_called_once_with("fake.docx")
    assert result["count"] == 0


@pytest.mark.unit
def test_manage_tracked_changes_router_routes_accept_all():
    """manage_tracked_changes 'accept_all' op routes to accept_all_track_changes."""
    from unittest.mock import patch
    from wordmcp import document_com
    from wordmcp._com import _write as _document_com_write

    with patch.object(
        _document_com_write,
        "accept_all_track_changes",
        return_value={"status": "ok", "accepted": True},
    ) as mock_fn:
        result = document_com.manage_tracked_changes(
            path="fake.docx", operation="accept_all", confirm=True
        )
    mock_fn.assert_called_once_with("fake.docx", True)
    assert result["accepted"] is True


@pytest.mark.unit
def test_manage_tracked_changes_invalid_operation_raises():
    """manage_tracked_changes rejects unrecognised operation values."""
    from wordmcp import document_com

    with pytest.raises(ValidationError, match="Unknown operation"):
        document_com.manage_tracked_changes(path="fake.docx", operation="delete_all")


@pytest.mark.unit
def test_export_document_routes_pdf():
    """export_document 'pdf' format routes to export_as_pdf after write gates."""
    from unittest.mock import patch
    from wordmcp import document_com
    from wordmcp._com import _export as _document_com_export

    with patch.object(_document_com_export, "_check_com_enabled"), \
         patch.object(_document_com_export, "_check_write"), \
         patch.object(_document_com_export, "_check_confirm"), \
         patch.object(
             _document_com_export,
             "export_as_pdf",
             return_value={"status": "ok", "output_path": "/fake.pdf"},
         ) as mock_fn:
        result = document_com.export_document(
            path="fake.docx", output_path="fake.pdf", format="pdf", confirm=True
        )
    mock_fn.assert_called_once_with("fake.docx", "fake.pdf", True)
    assert result["status"] == "ok"


@pytest.mark.unit
def test_export_document_html_requires_com(monkeypatch):
    """export_document 'html' format raises NotAllowedError when COM is not enabled."""
    from wordmcp import document_com

    monkeypatch.delenv("WORD_ENABLE_COM", raising=False)
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_COM"):
        document_com.export_document(
            path="fake.docx", output_path="fake.html", format="html"
        )


@pytest.mark.unit
def test_export_document_invalid_format():
    """export_document with unknown format raises ValidationError."""
    from wordmcp import document_com

    with pytest.raises(ValidationError):
        document_com.export_document(
            path="fake.docx", output_path="fake.xyz", format="rtf"
        )


# ===========================================================================
# Batch A — list_headings, read_section, search_text, list_styles, create_document
# ===========================================================================


@pytest.mark.unit
def test_list_headings_returns_only_headings(tmp_path, monkeypatch):
    """list_headings returns exactly the heading paragraphs with correct level/text."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "h.docx"
    d = Document()
    d.add_heading("Top Heading", level=1)
    d.add_paragraph("Body text 1")
    d.add_heading("Sub Heading", level=2)
    d.add_paragraph("Body text 2")
    d.save(str(path))
    result = list_headings(str(path))
    assert len(result) == 2
    assert result[0]["level"] == 1
    assert result[0]["text"] == "Top Heading"
    assert result[1]["level"] == 2
    assert result[1]["text"] == "Sub Heading"


@pytest.mark.unit
def test_list_headings_empty_doc(tmp_path, monkeypatch):
    """list_headings returns [] for a document with no headings."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "noheadings.docx"
    d = Document()
    d.add_paragraph("Just a normal paragraph")
    d.save(str(path))
    result = list_headings(str(path))
    assert result == []


@pytest.mark.unit
def test_read_section_returns_body_until_next_heading(tmp_path, monkeypatch):
    """read_section(0) returns heading + body paras stopping at next H1."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "sections.docx"
    d = Document()
    d.add_heading("First Section", level=1)   # index 0
    d.add_paragraph("Body of first section")  # index 1
    d.add_heading("Second Section", level=1)  # index 2 — stops iteration
    d.save(str(path))
    result = read_section(str(path), heading_index=0)
    assert result["heading"]["text"] == "First Section"
    assert len(result["body"]) == 1
    assert result["body"][0]["text"] == "Body of first section"


@pytest.mark.unit
def test_read_section_non_heading_raises(tmp_path, monkeypatch):
    """read_section raises ValidationError when the target paragraph is not a heading."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "nothead.docx"
    d = Document()
    d.add_paragraph("Normal paragraph")  # index 0 — not a heading
    d.save(str(path))
    with pytest.raises(ValidationError, match="not a heading"):
        read_section(str(path), heading_index=0)


@pytest.mark.unit
def test_read_section_out_of_range(tmp_path, monkeypatch):
    """read_section raises NotFoundError for a negative index."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "oor.docx"
    d = Document()
    d.add_paragraph("Only paragraph")
    d.save(str(path))
    with pytest.raises(NotFoundError):
        read_section(str(path), heading_index=-1)


@pytest.mark.unit
def test_search_text_finds_matches(tmp_path, monkeypatch):
    """search_text finds the paragraph containing the query string."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "search.docx"
    d = Document()
    d.add_paragraph("The quick brown fox")
    d.add_paragraph("Completely unrelated")
    d.save(str(path))
    result = search_text(str(path), query="quick")
    assert result["total_matches"] >= 1
    assert any(r["paragraph_index"] == 0 for r in result["results"])


@pytest.mark.unit
def test_search_text_case_insensitive(tmp_path, monkeypatch):
    """search_text with case_sensitive=False matches regardless of case."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "case.docx"
    d = Document()
    d.add_paragraph("HELLO world")
    d.save(str(path))
    result = search_text(str(path), query="hello", case_sensitive=False)
    assert result["total_matches"] >= 1


@pytest.mark.unit
def test_search_text_max_results_guard(tmp_path, monkeypatch):
    """search_text raises ValidationError when max_results is 0."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "guard.docx"
    Document().save(str(path))
    with pytest.raises(ValidationError, match="max_results"):
        search_text(str(path), query="x", max_results=0)


@pytest.mark.unit
def test_search_text_max_results_upper_bound(tmp_path, monkeypatch):
    """max_results above 1000 should raise ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    doc_file = tmp_path / "test.docx"
    d = Document()
    d.add_paragraph("hello")
    d.save(str(doc_file))
    with pytest.raises(ValidationError, match="max_results"):
        search_text(str(doc_file), query="hello", max_results=1001)


@pytest.mark.unit
def test_search_text_empty_query_raises(tmp_path, monkeypatch):
    """search_text raises ValidationError when query is empty string."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "empty.docx"
    Document().save(str(path))
    with pytest.raises(ValidationError, match="query must not be empty"):
        search_text(str(path), query="")


@pytest.mark.unit
def test_list_styles_returns_list(tmp_docx):
    """list_styles returns a non-empty list for a real document."""
    result = list_styles(tmp_docx)
    assert isinstance(result, list)
    assert len(result) > 0
    # Each entry has name, type, builtin keys
    assert all("name" in s and "type" in s and "builtin" in s for s in result)


@pytest.mark.unit
def test_list_styles_filters_by_type(tmp_docx):
    """list_styles with style_type='paragraph' returns only paragraph styles."""
    result = list_styles(tmp_docx, style_type="paragraph")
    assert all(s["type"] == "PARAGRAPH" for s in result)


@pytest.mark.unit
def test_list_styles_invalid_type(tmp_docx):
    """list_styles raises ValidationError for an unknown style_type."""
    with pytest.raises(ValidationError, match="Unknown style_type"):
        list_styles(tmp_docx, style_type="unknown_xyz")


@pytest.mark.unit
def test_create_document_creates_file(tmp_path, monkeypatch):
    """create_document with confirm=True creates a .docx file on disk."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "new.docx"
    result = create_document(str(path), title="My Doc", confirm=True)
    assert result["created"] is True
    assert path.exists()


@pytest.mark.unit
def test_create_document_no_confirm_raises(tmp_path, monkeypatch):
    """create_document without confirm=True raises ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "nope.docx"
    with pytest.raises(ValidationError, match="confirm=True required"):
        create_document(str(path))  # confirm defaults to False


@pytest.mark.unit
def test_create_document_already_exists_raises(tmp_path, monkeypatch):
    """create_document raises ValidationError when the file already exists."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "exists.docx"
    Document().save(str(path))
    with pytest.raises(ValidationError, match="already exists"):
        create_document(str(path), confirm=True)


# ===========================================================================
# Batch B — apply_style, update_paragraph, delete_paragraph,
#            update_table_cell, set_document_properties
# ===========================================================================


@pytest.mark.unit
def test_apply_style_applies_correctly(write_enabled_docx):
    """apply_style changes the paragraph's style to the requested name."""
    result = apply_style(write_enabled_docx, 2, "Normal", confirm=True)
    assert result["paragraph_index"] == 2
    assert result["style"] == "Normal"
    d = Document(write_enabled_docx)
    assert d.paragraphs[2].style.name == "Normal"


@pytest.mark.unit
def test_apply_style_unknown_style_raises(write_enabled_docx):
    """apply_style raises ValidationError for a non-existent style name."""
    with pytest.raises(ValidationError, match="Unknown style"):
        apply_style(write_enabled_docx, 0, "NonExistentStyle_XYZ999", confirm=True)


@pytest.mark.unit
def test_apply_style_out_of_range_raises(write_enabled_docx):
    """apply_style raises NotFoundError when paragraph_index is out of range."""
    with pytest.raises(NotFoundError):
        apply_style(write_enabled_docx, 99, "Normal", confirm=True)


@pytest.mark.unit
def test_update_paragraph_changes_text(write_enabled_docx):
    """update_paragraph replaces paragraph text; change is persisted to disk."""
    result = update_paragraph(write_enabled_docx, 2, "Updated text here", confirm=True)
    assert result["text"] == "Updated text here"
    assert result["paragraph_index"] == 2
    d = Document(write_enabled_docx)
    assert d.paragraphs[2].text == "Updated text here"


@pytest.mark.unit
def test_update_paragraph_preserves_style(write_enabled_docx):
    """update_paragraph does not change the paragraph's existing style."""
    # index 1 is "Heading 1" in write_enabled_docx (from tmp_docx fixture)
    style_before = Document(write_enabled_docx).paragraphs[1].style.name
    update_paragraph(write_enabled_docx, 1, "New heading text", confirm=True)
    d = Document(write_enabled_docx)
    assert d.paragraphs[1].style.name == style_before
    assert d.paragraphs[1].text == "New heading text"


@pytest.mark.unit
def test_delete_paragraph_removes_it(tmp_path, monkeypatch):
    """delete_paragraph removes the targeted paragraph; count reduces by 1."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "del_test.docx"
    d = Document()
    d.add_paragraph("Para A")
    d.add_paragraph("Para B")
    d.add_paragraph("Para C")
    d.save(str(path))
    result = delete_paragraph(str(path), 1, confirm=True)
    assert result["deleted"] is True
    assert result["paragraph_index"] == 1
    assert result["text_preview"] == "Para B"  # captures text before deletion
    assert "shift" in result["note"].lower()
    d2 = Document(str(path))
    assert len(d2.paragraphs) == 2
    assert d2.paragraphs[0].text == "Para A"
    assert d2.paragraphs[1].text == "Para C"


@pytest.mark.unit
def test_delete_paragraph_out_of_range(write_enabled_docx):
    """delete_paragraph raises NotFoundError for an out-of-range index."""
    with pytest.raises(NotFoundError):
        delete_paragraph(write_enabled_docx, 999, confirm=True)


@pytest.mark.unit
def test_update_table_cell_sets_text(tmp_path, monkeypatch):
    """update_table_cell sets cell text; change is persisted and readable."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "table_cell_test.docx"
    d = Document()
    tbl = d.add_table(2, 3, style="Table Grid")
    tbl.rows[0].cells[0].text = "Original"
    d.save(str(path))
    result = update_table_cell(str(path), 0, 0, 0, "Updated", confirm=True)
    assert result["text"] == "Updated"
    assert result["table_index"] == 0
    assert result["row"] == 0
    assert result["col"] == 0
    d2 = Document(str(path))
    assert d2.tables[0].rows[0].cells[0].text == "Updated"


@pytest.mark.unit
def test_update_table_cell_bad_row(tmp_path, monkeypatch):
    """update_table_cell raises ValidationError when row index is out of range."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "bad_row.docx"
    d = Document()
    d.add_table(2, 2)
    d.save(str(path))
    with pytest.raises(ValidationError, match="Row"):
        update_table_cell(str(path), 0, 99, 0, "text", confirm=True)


@pytest.mark.unit
def test_set_document_properties_updates_all(tmp_path, monkeypatch):
    """set_document_properties sets title and author; get_document_metadata reads them back."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "props_test.docx"
    Document().save(str(path))
    result = set_document_properties(
        str(path), title="My Title", author="Jane Doe", confirm=True
    )
    assert result["updated"] is True
    assert result["properties"]["title"] == "My Title"
    assert result["properties"]["author"] == "Jane Doe"
    meta = get_document_metadata(str(path))
    assert meta["title"] == "My Title"
    assert meta["author"] == "Jane Doe"


@pytest.mark.unit
def test_set_document_properties_all_none_raises(write_enabled_docx):
    """set_document_properties raises ValidationError when all params are None."""
    with pytest.raises(ValidationError, match="At least one property"):
        set_document_properties(write_enabled_docx, confirm=True)
