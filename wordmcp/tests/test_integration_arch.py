"""
W3 Architecture Test Plan — Integration Tests (live COM mocked via fixtures).

All tests in this file require either:
  - The `patch_word_com` fixture (monkeypatches _word_app_context — no real Word needed), OR
  - @pytest.mark.integration (requires real Word COM; skipped in CI)

Tests cover:
  P3  export_document pdf/html routes through COM layer (mock)
  P4  manage_tracked_changes list/accept_all/reject_all routes through COM (mock)
  P5  bulk_update_paragraphs happy path via docx layer (xfail until implemented)
  P6  bulk_update_table_cells happy path via docx layer (xfail until implemented)
  B7  COM export gate: write gate fired before _word_app_context is entered

Run (mock COM — no live Word):
  pytest wordmcp/tests/test_integration_arch.py -v -m "not integration and not com"

Run (live Word — Windows only):
  pytest wordmcp/tests/test_integration_arch.py -v -m integration
"""
from __future__ import annotations

from unittest.mock import MagicMock

import pytest
from docx import Document


# ===========================================================================
# P3 — export_document valid formats via COM mock (2 tests)
# ===========================================================================


@pytest.mark.com
def test_p3_export_document_pdf_calls_com_layer(com_env, patch_word_com):
    """P3-POS: export_document(format='pdf') with write env set reaches COM and returns ok.

    Uses patch_word_com fixture to replace _word_app_context with a mock so
    no real Word installation is required.
    """
    from wordmcp.document_com import export_document

    doc_path = com_env  # com_env fixture creates a .docx and sets all env vars

    out_path = doc_path.parent / "test_output.pdf"

    result = export_document(
        path=str(doc_path),
        output_path=str(out_path),
        format="pdf",
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("status") == "ok"
    assert "output_path" in result


@pytest.mark.com
def test_p3_export_document_html_calls_com_layer(com_env, monkeypatch):
    """P3-POS: export_document(format='html') with write env and COM mock returns ok.

    Patches _word_app_context directly to avoid needing real Word.
    """
    import contextlib

    from wordmcp import document_com as _dc
    from wordmcp._com import _base as _dc_base

    doc_path = com_env
    out_path = doc_path.parent / "test_output.html"

    mock_doc = MagicMock()
    mock_app = MagicMock()
    mock_app.Documents.Open.return_value = mock_doc

    @contextlib.contextmanager
    def _fake_app_ctx():
        yield mock_app

    monkeypatch.setattr(_dc_base, "_word_app_context", _fake_app_ctx)
    monkeypatch.setattr(_dc_base, "_COM_AVAILABLE", True)

    result = _dc.export_document(
        path=str(doc_path),
        output_path=str(out_path),
        format="html",
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("status") == "ok"
    assert "output_path" in result


# ===========================================================================
# P4 — manage_tracked_changes valid ops via COM mock (4 tests)
# ===========================================================================


@pytest.mark.com
def test_p4_manage_tracked_changes_list_via_mock(com_env, patch_word_com):
    """P4-POS: manage_tracked_changes('list') returns dict with 'revisions' key."""
    from wordmcp.document_com import manage_tracked_changes

    result = manage_tracked_changes(
        path=str(com_env),
        operation="list",
    )

    assert isinstance(result, dict)
    assert "revisions" in result
    assert isinstance(result["revisions"], list)


@pytest.mark.com
def test_p4_manage_tracked_changes_accept_all_via_mock(com_env, patch_word_com):
    """P4-POS: manage_tracked_changes('accept_all') returns {'accepted': True}."""
    from wordmcp.document_com import manage_tracked_changes

    result = manage_tracked_changes(
        path=str(com_env),
        operation="accept_all",
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("accepted") is True


@pytest.mark.com
def test_p4_manage_tracked_changes_reject_all_via_mock(com_env, patch_word_com):
    """P4-POS: manage_tracked_changes('reject_all') returns {'rejected': True}."""
    from wordmcp.document_com import manage_tracked_changes

    result = manage_tracked_changes(
        path=str(com_env),
        operation="reject_all",
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("rejected") is True


@pytest.mark.com
def test_p4_manage_tracked_changes_accept_all_calls_revisions_accept(
    com_env, patch_word_com
):
    """P4-POS (interaction): accept_all must call doc.Revisions.AcceptAll() exactly once."""
    from wordmcp.document_com import manage_tracked_changes

    # patch_word_com fixture gives us back the mock_word_app
    mock_app = patch_word_com
    mock_doc = mock_app.Documents.Open.return_value

    manage_tracked_changes(
        path=str(com_env),
        operation="accept_all",
        confirm=True,
    )

    mock_doc.Revisions.AcceptAll.assert_called_once()


# ===========================================================================
# B7 — COM export write gate fires before _word_app_context enters (1 test)
# ===========================================================================


@pytest.mark.unit
def test_b7_export_document_write_gate_before_word_app_context(
    com_env, monkeypatch
):
    """B7-NEG (integration): export_document with COM enabled but write disabled must
    raise NotAllowedError BEFORE _word_app_context is entered — the COM channel
    must never be opened for an unauthorised write.
    """
    import contextlib

    from wordmcp import document_com as _dc
    from wordmcp._com import _base as _dc_base
    from wordmcp._com import _export as _dc_export
    from wordmcp.document_com import NotAllowedError

    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)

    ctx_entered = []

    @contextlib.contextmanager
    def _spy_ctx():
        ctx_entered.append(True)
        yield MagicMock()

    monkeypatch.setattr(_dc_base, "_word_app_context", _spy_ctx)
    monkeypatch.setattr(_dc_export, "_word_app_context", _spy_ctx)

    with pytest.raises(NotAllowedError):
        _dc.export_as_pdf(
            path=str(com_env),
            output_path=str(com_env.parent / "spy_out.pdf"),
            confirm=True,
        )

    assert ctx_entered == [], (
        "_word_app_context must NOT be entered when _check_write() raises NotAllowedError"
    )


# ===========================================================================
# P5 — bulk_update_paragraphs happy path
# ===========================================================================


@pytest.mark.unit
def test_p5_bulk_update_paragraphs_integration_happy_path(write_enabled_docx):
    """P5-POS: bulk_update_paragraphs writes all items atomically and returns ok result."""
    from wordmcp.document_docx import bulk_update_paragraphs

    result = bulk_update_paragraphs(
        path=write_enabled_docx,
        updates=[
            {"paragraph_index": 0, "new_text": "Updated heading"},
            {"paragraph_index": 2, "new_text": "Updated body"},
        ],
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("status") == "ok" or result.get("updated", 0) > 0


# ===========================================================================
# P6 — bulk_update_table_cells happy path
# ===========================================================================


@pytest.mark.unit
def test_p6_bulk_update_table_cells_integration_happy_path(
    write_enabled_docx_with_table,
):
    """P6-POS: bulk_update_table_cells writes all cell updates atomically and returns ok."""
    from wordmcp.document_docx import bulk_update_table_cells

    result = bulk_update_table_cells(
        path=write_enabled_docx_with_table,
        updates=[
            {"table_index": 0, "row": 0, "col": 0, "new_text": "Name*"},
            {"table_index": 0, "row": 1, "col": 2, "new_text": "New York"},
        ],
        confirm=True,
    )

    assert isinstance(result, dict)
    assert result.get("status") == "ok" or result.get("updated", 0) > 0


# ===========================================================================
# Live COM tests (require real Word — skipped in CI)
# ===========================================================================


@pytest.mark.integration
def test_live_export_document_pdf_creates_file(tmp_path, monkeypatch):
    """LIVE-COM: export_document(format='pdf') actually writes a .pdf file.

    Requires: Windows + Word installed + WORD_ENABLE_COM=true env in the shell.
    Skip: pytest -m "not integration"
    """
    from wordmcp.document_com import export_document

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

    doc_path = tmp_path / "live_test.docx"
    d = Document()
    d.add_paragraph("Live COM export test")
    d.save(str(doc_path))

    out_path = tmp_path / "live_test.pdf"

    result = export_document(
        path=str(doc_path),
        output_path=str(out_path),
        format="pdf",
        confirm=True,
    )

    assert result.get("status") == "ok"
    assert out_path.exists(), "PDF file was not created on disk"


@pytest.mark.integration
def test_live_manage_tracked_changes_list(tmp_path, monkeypatch):
    """LIVE-COM: manage_tracked_changes('list') returns revisions list on a real docx.

    Requires: Windows + Word installed + WORD_ENABLE_COM=true.
    Skip: pytest -m "not integration"
    """
    from wordmcp.document_com import manage_tracked_changes

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

    doc_path = tmp_path / "live_tc.docx"
    d = Document()
    d.add_paragraph("Some content")
    d.save(str(doc_path))

    result = manage_tracked_changes(
        path=str(doc_path),
        operation="list",
    )

    assert "revisions" in result
    assert isinstance(result["revisions"], list)
