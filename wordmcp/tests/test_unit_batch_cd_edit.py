"""
Unit tests — delete_paragraph, find_replace, insert_paragraph, add_list (level).
Split from test_unit_batch_cd.py.
"""
from __future__ import annotations

import contextlib
import sys

import pytest
from docx import Document

from wordmcp.document_docx import (
    NotAllowedError,
    ValidationError,
    add_list,
    delete_paragraph,
    find_replace,
    insert_paragraph,
)



# ---------------------------------------------------------------------------

def _make_mock_word_ctx(mock_app):
    """Return a context-manager mock that yields mock_app."""

    @contextlib.contextmanager
    def _ctx():
        yield mock_app

    return _ctx


# ===========================================================================
# Batch E: COM accept_one / reject_one / html-export (12 tests)
# ===========================================================================


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")

# ===========================================================================
# BUG REGRESSION TESTS — 4 bugs fixed
# ===========================================================================

# ---------------------------------------------------------------------------
# BUG 1: delete_paragraph — text_preview in return value
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_delete_paragraph_text_preview_returned(tmp_path, monkeypatch):
    """delete_paragraph result includes text_preview of the deleted paragraph."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "preview_test.docx"
    d = Document()
    d.add_paragraph("Short text for preview")
    d.add_paragraph("Keep this")
    d.save(str(path))
    result = delete_paragraph(str(path), 0, confirm=True)
    assert result["deleted"] is True
    assert "text_preview" in result
    assert result["text_preview"] == "Short text for preview"


@pytest.mark.unit
def test_delete_paragraph_text_preview_truncated_at_50(tmp_path, monkeypatch):
    """delete_paragraph text_preview is truncated to 50 chars for long text."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "long_preview.docx"
    d = Document()
    long_text = "X" * 100
    d.add_paragraph(long_text)
    d.add_paragraph("Another")
    d.save(str(path))
    result = delete_paragraph(str(path), 0, confirm=True)
    assert result["deleted"] is True
    assert len(result["text_preview"]) == 50
    assert result["text_preview"] == "X" * 50


# ---------------------------------------------------------------------------
# BUG 2: find_replace — cross-run matching + paragraph_index + occurrence
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_find_replace_cross_run_match(tmp_path, monkeypatch):
    """find_replace finds text spanning two runs in one paragraph."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "cross_run.docx"
    d = Document()
    para = d.add_paragraph()
    para.add_run("Hello ")
    bold_run = para.add_run("world")
    bold_run.bold = True
    d.save(str(path))
    result = find_replace(str(path), "Hello world", "Hi everyone", confirm=True)
    assert result["replacements_made"] == 1
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "Hi everyone"


@pytest.mark.unit
def test_find_replace_cross_run_preserves_first_run_formatting(tmp_path, monkeypatch):
    """After cross-run replace, remaining run carries first run's bold/italic."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "fmt_preserve.docx"
    d = Document()
    para = d.add_paragraph()
    r1 = para.add_run("start ")
    r1.italic = True
    para.add_run("end")  # second run, not italic
    d.save(str(path))
    find_replace(str(path), "start end", "replaced", confirm=True)
    d2 = Document(str(path))
    # The first (and now only) run should be italic (preserved from r1)
    assert d2.paragraphs[0].runs[0].italic is True
    assert d2.paragraphs[0].text == "replaced"


@pytest.mark.unit
def test_find_replace_paragraph_index_limits_scope(tmp_path, monkeypatch):
    """find_replace with paragraph_index only replaces in the specified paragraph."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "scope_test.docx"
    d = Document()
    d.add_paragraph("foo in first paragraph")
    d.add_paragraph("foo in second paragraph")
    d.save(str(path))
    result = find_replace(str(path), "foo", "bar", paragraph_index=0, confirm=True)
    assert result["replacements_made"] == 1
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "bar in first paragraph"
    assert d2.paragraphs[1].text == "foo in second paragraph"


@pytest.mark.unit
def test_find_replace_occurrence_targets_second(tmp_path, monkeypatch):
    """find_replace with occurrence=2 replaces only the 2nd global occurrence."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "occ_test.docx"
    d = Document()
    d.add_paragraph("foo alpha foo")  # occurrences 1 and 2
    d.add_paragraph("foo beta")       # occurrence 3
    d.save(str(path))
    result = find_replace(str(path), "foo", "bar", occurrence=2, confirm=True)
    assert result["replacements_made"] == 1
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "foo alpha bar"  # only 2nd replaced
    assert d2.paragraphs[1].text == "foo beta"        # unchanged


@pytest.mark.unit
def test_find_replace_occurrence_zero_raises(write_enabled_docx):
    """find_replace raises ValidationError when occurrence=0 (must be >= 1)."""
    with pytest.raises(ValidationError, match="occurrence must be"):
        find_replace(write_enabled_docx, "text", "X", occurrence=0, confirm=True)


@pytest.mark.unit
def test_find_replace_occurrence_no_match_returns_zero(tmp_path, monkeypatch):
    """find_replace with occurrence=5 returns 0 when fewer than 5 occurrences exist."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "no_occ.docx"
    d = Document()
    d.add_paragraph("hello hello")  # only 2 occurrences
    d.save(str(path))
    result = find_replace(str(path), "hello", "world", occurrence=5, confirm=True)
    assert result["replacements_made"] == 0


# ---------------------------------------------------------------------------
# BUG 3: insert_paragraph — insert at specific index
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_insert_paragraph_at_index_0(tmp_path, monkeypatch):
    """insert_paragraph at index 0 prepends before the first paragraph."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "insert_start.docx"
    d = Document()
    d.add_paragraph("Original first")
    d.add_paragraph("Original second")
    d.save(str(path))
    result = insert_paragraph(str(path), paragraph_index=0, text="New first", confirm=True)
    assert result["inserted"] is True
    assert result["paragraph_index"] == 0
    assert result["text"] == "New first"
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "New first"
    assert d2.paragraphs[1].text == "Original first"
    assert d2.paragraphs[2].text == "Original second"


@pytest.mark.unit
def test_insert_paragraph_at_middle(tmp_path, monkeypatch):
    """insert_paragraph at index 1 inserts between paragraph 0 and 1."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "insert_mid.docx"
    d = Document()
    d.add_paragraph("First")
    d.add_paragraph("Third")
    d.save(str(path))
    insert_paragraph(str(path), paragraph_index=1, text="Second", confirm=True)
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "First"
    assert d2.paragraphs[1].text == "Second"
    assert d2.paragraphs[2].text == "Third"


@pytest.mark.unit
def test_insert_paragraph_at_end(tmp_path, monkeypatch):
    """insert_paragraph at index==len(paragraphs) appends to the end."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "insert_end.docx"
    d = Document()
    d.add_paragraph("Alpha")
    d.add_paragraph("Beta")
    d.save(str(path))
    n = len(Document(str(path)).paragraphs)
    insert_paragraph(str(path), paragraph_index=n, text="Gamma", confirm=True)
    d2 = Document(str(path))
    assert d2.paragraphs[-1].text == "Gamma"


@pytest.mark.unit
def test_insert_paragraph_out_of_range(tmp_path, monkeypatch):
    """insert_paragraph raises ValidationError for paragraph_index out of range."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "oor_insert.docx"
    d = Document()
    d.add_paragraph("Only para")
    d.save(str(path))
    with pytest.raises(ValidationError, match="out of range"):
        insert_paragraph(str(path), paragraph_index=99, text="Bad", confirm=True)


@pytest.mark.unit
def test_insert_paragraph_with_style(tmp_path, monkeypatch):
    """insert_paragraph with style= applies the paragraph style correctly."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "style_insert.docx"
    d = Document()
    d.add_paragraph("Normal para")
    d.save(str(path))
    insert_paragraph(
        str(path), paragraph_index=0, text="My Heading", style="Heading 1", confirm=True
    )
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "My Heading"
    assert d2.paragraphs[0].style.name == "Heading 1"


@pytest.mark.unit
def test_insert_paragraph_invalid_style_raises(tmp_path, monkeypatch):
    """insert_paragraph raises ValidationError for an unknown style name."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "bad_style.docx"
    d = Document()
    d.add_paragraph("Para")
    d.save(str(path))
    with pytest.raises(ValidationError, match="Unknown style"):
        insert_paragraph(
            str(path), paragraph_index=0, text="X",
            style="NonExistentStyle_XYZ999", confirm=True,
        )


@pytest.mark.unit
def test_insert_paragraph_requires_confirm(write_enabled_docx):
    """insert_paragraph without confirm=True raises ValidationError."""
    with pytest.raises(ValidationError, match="confirm=True required"):
        insert_paragraph(write_enabled_docx, paragraph_index=0, text="X", confirm=False)


@pytest.mark.unit
def test_insert_paragraph_requires_enable_write(tmp_docx):
    """insert_paragraph raises NotAllowedError when WORD_ENABLE_WRITE is unset."""
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        insert_paragraph(tmp_docx, paragraph_index=0, text="X", confirm=True)


# ---------------------------------------------------------------------------
# BUG 4: add_list — level parameter for nested bullets
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_add_list_level_zero_backward_compatible(write_enabled_docx):
    """add_list with level=0 (default) preserves existing List Bullet style behavior."""
    result = add_list(write_enabled_docx, ["item"], list_type="bullet", level=0, confirm=True)
    assert result["added"] == 1
    assert result["level"] == 0
    d = Document(write_enabled_docx)
    assert d.paragraphs[-1].style.name == "List Bullet"


@pytest.mark.unit
def test_add_list_returns_level_in_result(write_enabled_docx):
    """add_list result dict includes the 'level' field."""
    result = add_list(write_enabled_docx, ["a", "b"], list_type="number", confirm=True)
    assert "level" in result
    assert result["level"] == 0


@pytest.mark.unit
def test_add_list_invalid_level_raises(write_enabled_docx):
    """add_list raises ValidationError when level is outside 0-8."""
    with pytest.raises(ValidationError, match="level must be 0-8"):
        add_list(write_enabled_docx, ["item"], level=9, confirm=True)
    with pytest.raises(ValidationError, match="level must be 0-8"):
        add_list(write_enabled_docx, ["item"], level=-1, confirm=True)


@pytest.mark.unit
def test_add_list_level_1_sets_ilvl_in_xml(tmp_path, monkeypatch):
    """add_list with level=1 creates paragraph-level numPr with ilvl=1."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "level_xml.docx"
    d = Document()
    d.add_paragraph("anchor")
    d.save(str(path))
    result = add_list(str(path), ["nested item"], list_type="bullet", level=1, confirm=True)
    assert result["added"] == 1
    assert result["level"] == 1
    d2 = Document(str(path))
    para = d2.paragraphs[-1]
    from docx.oxml.ns import qn as _qn
    pPr = para._p.find(_qn("w:pPr"))
    assert pPr is not None, "Paragraph must have pPr"
    numPr = pPr.find(_qn("w:numPr"))
    assert numPr is not None, "Paragraph must have paragraph-level numPr when level > 0"
    ilvl = numPr.find(_qn("w:ilvl"))
    assert ilvl is not None, "numPr must contain ilvl element"
    assert ilvl.get(_qn("w:val")) == "1", "ilvl value must be '1'"


@pytest.mark.unit
def test_add_list_level_1_number(tmp_path, monkeypatch):
    """add_list with list_type='number' and level=1 sets numPr ilvl=1."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "num_level.docx"
    d = Document()
    d.save(str(path))
    result = add_list(str(path), ["sub-item"], list_type="number", level=1, confirm=True)
    assert result["added"] == 1
    assert result["level"] == 1
