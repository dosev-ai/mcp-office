"""Security guard tests for wordmcp — covers H-001 through H-004 and H-009.

All tests are unit-level (no COM required, no live Word).
"""
from __future__ import annotations

import pytest


# ---------------------------------------------------------------------------
# H-009 — _check_path: null byte rejection
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestCheckPathNullByte:
    """_check_path must reject paths containing null bytes."""

    def test_null_byte_raises(self, tmp_path, monkeypatch):
        from wordmcp.document_docx import ValidationError, _check_path

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        with pytest.raises(ValidationError, match="[Nn]ull"):
            _check_path(str(tmp_path / "fi\x00le.docx"), must_exist=False)

    def test_null_byte_in_middle_raises(self, tmp_path, monkeypatch):
        from wordmcp.document_docx import ValidationError, _check_path

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        with pytest.raises(ValidationError, match="[Nn]ull"):
            _check_path("/\x00etc/passwd.docx", must_exist=False)

    def test_clean_path_not_rejected(self, tmp_path, monkeypatch):
        """A valid path inside the allowlist (must_exist=False) is accepted."""
        from wordmcp.document_docx import _check_path

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        result = _check_path(str(tmp_path / "valid.docx"), must_exist=False)
        assert result.suffix == ".docx"


# ---------------------------------------------------------------------------
# H-001 — _doc_cache: bounded LRU eviction
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestDocCacheLRU:
    """_doc_cache must evict the LRU entry when at capacity."""

    def test_cache_bounded(self, tmp_path, monkeypatch):
        """Cache must not exceed _DOC_CACHE_MAX entries after loading more docs."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

        # Create _DOC_CACHE_MAX + 1 minimal docx files
        paths = []
        for i in range(dd._DOC_CACHE_MAX + 1):
            p = tmp_path / f"doc{i}.docx"
            DocxDoc().save(str(p))
            paths.append(p)

        # Work with a fresh cache so previous test state doesn't bleed in
        dd._doc_cache.clear()
        for p in paths:
            dd._load_doc(p)

        assert len(dd._doc_cache) <= dd._DOC_CACHE_MAX

    def test_lru_eviction_order(self, tmp_path, monkeypatch):
        """The least-recently-used entry should be evicted first."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setattr(dd, "_DOC_CACHE_MAX", 2)
        dd._doc_cache.clear()

        p0 = tmp_path / "e0.docx"
        p1 = tmp_path / "e1.docx"
        p2 = tmp_path / "e2.docx"
        for p in (p0, p1, p2):
            DocxDoc().save(str(p))

        dd._load_doc(p0)  # cache: [p0]
        dd._load_doc(p1)  # cache: [p0, p1]
        # Access p0 so it becomes MRU; p1 becomes LRU
        dd._load_doc(p0)
        # Loading p2 should evict p1 (LRU)
        dd._load_doc(p2)

        assert str(p1) not in dd._doc_cache, "p1 (LRU) should have been evicted"
        assert str(p0) in dd._doc_cache
        assert str(p2) in dd._doc_cache


# ---------------------------------------------------------------------------
# H-002 — _check_path: file size limit
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestCheckPathFileSize:
    """_check_path must reject files exceeding the configured size limit."""

    def test_oversized_file_rejected(self, tmp_path, monkeypatch):
        import wordmcp.document_docx as dd
        from wordmcp.document_docx import ValidationError, _check_path

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        # Pin the constant to 1 MB for this test
        monkeypatch.setattr(dd, "_MAX_DOCX_BYTES", 1024 * 1024)

        f = tmp_path / "big.docx"
        f.write_bytes(b"x" * (1024 * 1024 + 1))  # 1 byte over the limit

        with pytest.raises(ValidationError, match="[Ss]ize|[Mm]aximum"):
            _check_path(str(f), must_exist=True)

    def test_under_limit_accepted(self, tmp_path, monkeypatch):
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd
        from wordmcp.document_docx import _check_path

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setattr(dd, "_MAX_DOCX_BYTES", 10 * 1024 * 1024)  # 10 MB

        f = tmp_path / "small.docx"
        DocxDoc().save(str(f))  # minimal docx is well under 10 MB

        result = _check_path(str(f), must_exist=True)
        assert result.exists()


# ---------------------------------------------------------------------------
# H-003 — manage_comments 'add': control character sanitization
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestManageCommentsControlChars:
    """manage_comments 'add' must sanitize control characters in author/text."""

    def test_control_char_in_text_sanitized(self, tmp_path, monkeypatch):
        """Control chars in comment text must NOT cause a ValueError from lxml."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        p = tmp_path / "ctrl_text.docx"
        d = DocxDoc()
        d.add_paragraph("Hello world")
        d.save(str(p))

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")

        # Should NOT raise ValueError — control chars are stripped
        result = dd.manage_comments(
            str(p),
            operation="add",
            paragraph_index=0,
            text="review\x01here\x0ccheck",
            author="MCP",
            confirm=True,
        )
        assert result.get("added") is True or result.get("status") == "added"
        assert "comment_id" in result

    def test_control_char_in_author_sanitized(self, tmp_path, monkeypatch):
        """Control chars in the author field must NOT cause ValueError from lxml."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        p = tmp_path / "ctrl_author.docx"
        d = DocxDoc()
        d.add_paragraph("Hello world")
        d.save(str(p))

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")

        result = dd.manage_comments(
            str(p),
            operation="add",
            paragraph_index=0,
            text="Normal text",
            author="Evil\x01Author\x07",
            confirm=True,
        )
        assert result.get("added") is True or result.get("status") == "added"

    def test_null_byte_stripped_from_text(self, tmp_path, monkeypatch):
        """Null bytes in comment text must be stripped (XML 1.0 forbids them)."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        p = tmp_path / "null_text.docx"
        d = DocxDoc()
        d.add_paragraph("Test")
        d.save(str(p))

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")

        result = dd.manage_comments(
            str(p),
            operation="add",
            paragraph_index=0,
            text="bad\x00text",
            author="MCP",
            confirm=True,
        )
        assert result.get("added") is True or result.get("status") == "added"


# ---------------------------------------------------------------------------
# H-004 — add_hyperlink: null byte / control char in URL
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestAddHyperlinkNullByte:
    """add_hyperlink must reject URLs containing null bytes or control chars."""

    def test_null_byte_in_url_raises(self, tmp_path, monkeypatch):
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd
        from wordmcp.document_docx import ValidationError

        p = tmp_path / "hyperlink_null.docx"
        d = DocxDoc()
        d.add_paragraph("Click here")
        d.save(str(p))

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")

        with pytest.raises(ValidationError, match="[Ii]nvalid|[Nn]ull"):
            dd.add_hyperlink(str(p), 0, "click", "http://\x00evil.com", confirm=True)

    def test_control_char_in_url_raises(self, tmp_path, monkeypatch):
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd
        from wordmcp.document_docx import ValidationError

        p = tmp_path / "hyperlink_ctrl.docx"
        d = DocxDoc()
        d.add_paragraph("Click here")
        d.save(str(p))

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")

        with pytest.raises(ValidationError, match="[Ii]nvalid"):
            dd.add_hyperlink(str(p), 0, "click", "http://evil\x01.com", confirm=True)

    def test_valid_url_accepted(self, tmp_path, monkeypatch):
        """A clean https URL should not be blocked by the control-char guard."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        p = tmp_path / "hyperlink_ok.docx"
        d = DocxDoc()
        d.add_paragraph("Visit our site")
        d.save(str(p))

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")

        result = dd.add_hyperlink(
            str(p), 0, "here", "https://example.com/path?q=1", confirm=True
        )
        assert result["url"] == "https://example.com/path?q=1"


# ---------------------------------------------------------------------------
# Env-var hardening regressions — Notes 1 & 2 from peer review
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestEnvVarHardening:
    """Env var parsing must fall back to safe defaults on bad input."""

    def test_max_docx_bytes_is_positive(self):
        """_MAX_DOCX_BYTES must always be > 0 (guards against bad env at import)."""
        import wordmcp.document_docx as dd

        assert dd._MAX_DOCX_BYTES > 0

    def test_doc_cache_max_is_at_least_one(self):
        """_DOC_CACHE_MAX must always be >= 1; WORD_MAX_CACHE_DOCS=0 is clamped."""
        import wordmcp.document_docx as dd

        assert dd._DOC_CACHE_MAX >= 1

    def test_doc_cache_evicts_safely_at_limit(self, tmp_path, monkeypatch):
        """Cache must not crash when loading exactly _DOC_CACHE_MAX+1 documents.

        Regression for: WORD_MAX_CACHE_DOCS=0 → popitem() KeyError on empty dict.
        """
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        dd._doc_cache.clear()
        monkeypatch.setattr(dd, "_DOC_CACHE_MAX", 2)  # small to exercise eviction

        for i in range(3):
            p = tmp_path / f"evict{i}.docx"
            DocxDoc().save(str(p))
            dd._load_doc(p)  # must not raise KeyError

        assert len(dd._doc_cache) <= 2

    def test_capabilities_governance_includes_max_file_mb(self):
        """capabilities() governance dict must include WORD_MAX_FILE_MB."""
        import wordmcp.document_docx as dd

        result = dd.capabilities()
        assert "WORD_MAX_FILE_MB" in result["governance"]

    def test_capabilities_governance_includes_max_cache_docs(self):
        """capabilities() governance dict must include WORD_MAX_CACHE_DOCS."""
        import wordmcp.document_docx as dd

        result = dd.capabilities()
        assert "WORD_MAX_CACHE_DOCS" in result["governance"]


# ---------------------------------------------------------------------------
# H-010 — _load_doc: mtime-based cache invalidation
# ---------------------------------------------------------------------------


@pytest.mark.unit
class TestDocCacheInvalidation:
    """_load_doc must evict the cache entry when the file mtime changes."""

    def test_mtime_drift_evicts_cache(self, tmp_path, monkeypatch):
        """When mtime changes between two calls, Document() must be called twice."""
        from unittest.mock import MagicMock

        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        dd._doc_cache.clear()

        p = tmp_path / "drift.docx"
        DocxDoc().save(str(p))

        call_count = 0
        real_stat = p.stat

        def _stat_side_effect():
            # Simulate a changed mtime on the second call
            s = real_stat()
            if call_count == 0:
                return s
            # Return a stat-like object with a different mtime_ns
            result = MagicMock()
            result.st_mtime_ns = s.st_mtime_ns + 1_000_000_000  # +1 second
            return result

        # Patch Path.stat to return a higher mtime on the second stat() call
        real_stat_ns = p.stat().st_mtime_ns
        # First call returns original mtime; second call returns mtime+1s
        stat_ns_sequence = [real_stat_ns, real_stat_ns + 1_000_000_000]

        original_path_class = type(p)
        original_stat = original_path_class.stat

        call_order = [0]

        def patched_stat(self, **kwargs):
            result = original_stat(self, **kwargs)
            if str(self) == str(p):
                idx = call_order[0]
                call_order[0] += 1
                mtime = stat_ns_sequence[min(idx, len(stat_ns_sequence) - 1)]
                mock_result = MagicMock()
                mock_result.st_mtime_ns = mtime
                return mock_result
            return result

        monkeypatch.setattr(original_path_class, "stat", patched_stat)

        doc_calls = []

        import wordmcp._docx.core as core_mod

        original_document = core_mod.Document

        def tracking_document(path_str):
            doc_calls.append(path_str)
            return original_document(path_str)

        monkeypatch.setattr(core_mod, "Document", tracking_document)

        # First load — populates cache with mtime_ns[0]
        dd._load_doc(p)
        # Second load — mtime_ns[1] differs, must reload
        dd._load_doc(p)

        assert len(doc_calls) == 2, (
            f"Expected Document() called twice (mtime drift), got {len(doc_calls)}"
        )

    def test_mtime_unchanged_returns_cached(self, tmp_path, monkeypatch):
        """When mtime is identical on second call, Document() must be called only once."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd
        import wordmcp._docx.core as core_mod

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        dd._doc_cache.clear()

        p = tmp_path / "stable.docx"
        DocxDoc().save(str(p))

        doc_calls = []
        original_document = core_mod.Document

        def tracking_document(path_str):
            doc_calls.append(path_str)
            return original_document(path_str)

        monkeypatch.setattr(core_mod, "Document", tracking_document)

        # Both loads see the same real mtime — no file modification between them
        dd._load_doc(p)
        dd._load_doc(p)

        assert len(doc_calls) == 1, (
            f"Expected Document() called once (mtime stable), got {len(doc_calls)}"
        )

    def test_delete_recreate_evicts_cache(self, tmp_path, monkeypatch):
        """When stat() raises OSError (file transiently missing), the cache entry
        is evicted and Document() is called again on the next successful load."""
        from docx import Document as DocxDoc

        import wordmcp.document_docx as dd
        import wordmcp._docx.core as core_mod

        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        dd._doc_cache.clear()

        p = tmp_path / "recreated.docx"
        DocxDoc().save(str(p))

        doc_calls = []
        original_document = core_mod.Document

        def tracking_document(path_str):
            doc_calls.append(path_str)
            return original_document(path_str)

        monkeypatch.setattr(core_mod, "Document", tracking_document)

        original_path_class = type(p)
        original_stat = original_path_class.stat
        stat_call_count = [0]

        def patched_stat(self, **kwargs):
            if str(self) == str(p):
                idx = stat_call_count[0]
                stat_call_count[0] += 1
                if idx == 1:
                    # Simulate file missing on second stat (between first and second load)
                    raise OSError("No such file or directory")
            return original_stat(self, **kwargs)

        monkeypatch.setattr(original_path_class, "stat", patched_stat)

        # First load — succeeds, caches with real mtime
        dd._load_doc(p)
        assert len(doc_calls) == 1

        # Second load — stat() raises OSError: entry must be evicted, doc reloaded
        dd._load_doc(p)
        assert len(doc_calls) == 2, (
            f"Expected Document() called twice after OSError eviction, got {len(doc_calls)}"
        )
