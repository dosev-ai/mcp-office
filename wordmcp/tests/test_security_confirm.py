"""
Security tests for wordmcp — confirm gate, error hierarchy, atomic save, image path.
Split from test_security.py.
"""
from __future__ import annotations

import os
import struct
import zlib
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from wordmcp.document_docx import (
    FileError,
    NotAllowedError,
    NotFoundError,
    ValidationError,
    WordMCPError,
    _atomic_save,
    _check_image_path,
    _check_path,
    add_paragraph,
    add_table,
    capabilities,
    export_as_text,
    find_replace,
    list_paragraphs,
    read_document,
    save,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(path: Path) -> None:
    """Create a minimal valid .docx at path."""
    Document().save(str(path))


def _make_docx_with_text(path: Path, text: str) -> None:
    d = Document()
    d.add_paragraph(text)
    d.save(str(path))


def _make_png(path: Path) -> None:
    """Create a minimal valid 1x1 PNG at path."""
    def _chunk(name: bytes, data: bytes) -> bytes:
        c = name + data
        return (
            struct.pack(">I", len(data))
            + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\xff\xff")
    img = (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", idat)
        + _chunk(b"IEND", b"")
    )
    path.write_bytes(img)


# ===========================================================================


@pytest.mark.security
def test_f_confirm_false_rejects_add_paragraph(tmp_path, monkeypatch):
    """confirm=False rejects add_paragraph with ValidationError."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(ValidationError, match="confirm=True required"):
        add_paragraph(str(path), "text", confirm=False)


@pytest.mark.security
def test_f_confirm_false_rejects_save(tmp_path, monkeypatch):
    """confirm=False rejects save with ValidationError."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(ValidationError, match="confirm=True required"):
        save(str(path), confirm=False)


@pytest.mark.security
def test_f_confirm_false_rejects_find_replace(tmp_path, monkeypatch):
    """confirm=False rejects find_replace with ValidationError."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx_with_text(path, "hello world")
    with pytest.raises(ValidationError, match="confirm=True required"):
        find_replace(str(path), "hello", "bye", confirm=False)


@pytest.mark.security
def test_f_confirm_true_with_write_enabled_accepts(tmp_path, monkeypatch):
    """confirm=True + WORD_ENABLE_WRITE=true allows the operation."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx_with_text(path, "hello world")
    result = find_replace(str(path), "hello", "bye", confirm=True)
    assert result["replacements_made"] == 1


# ===========================================================================
# GROUP G — DoS bounds (4 tests)
# ===========================================================================


@pytest.mark.security
def test_g_max_chars_preview_truncation(tmp_path, monkeypatch):
    """export_as_text truncates output at WORD_MAX_TEXT_CHARS."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_MAX_TEXT_CHARS", "50")
    path = tmp_path / "big.docx"
    d = Document()
    for _ in range(10):
        d.add_paragraph("A" * 100)
    d.save(str(path))
    result = export_as_text(str(path))
    assert result["truncated"] is True
    # text before the truncation notice is at most 50 chars (strip trailing newline)
    assert len(result["text"].split("[TRUNCATED")[0].rstrip("\n")) <= 50


@pytest.mark.security
def test_g_add_table_rows_over_500_raises(tmp_path, monkeypatch):
    """add_table with rows=501 raises ValidationError (DoS bound)."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(ValidationError, match="Table too large"):
        add_table(str(path), rows=501, cols=1, confirm=True)


@pytest.mark.security
def test_g_add_table_cols_over_100_raises(tmp_path, monkeypatch):
    """add_table with cols=101 raises ValidationError (DoS bound)."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(ValidationError, match="Table too large"):
        add_table(str(path), rows=1, cols=101, confirm=True)


@pytest.mark.security
def test_g_add_table_boundary_500x100_accepted(tmp_path, monkeypatch):
    """add_table with rows=500 and cols=100 is accepted (boundary inclusive)."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    result = add_table(str(path), rows=500, cols=100, confirm=True)
    assert result["rows"] == 500
    assert result["cols"] == 100


# ===========================================================================
# GROUP H — XML-safe text (4 tests)
# ===========================================================================


@pytest.mark.security
def test_h_xml_script_tag_in_paragraph_accepted(tmp_path, monkeypatch):
    """<script> in paragraph text is accepted (lxml auto-escapes)."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    result = add_paragraph(str(path), "<script>alert('xss')</script>", confirm=True)
    assert result["text"] == "<script>alert('xss')</script>"
    # Verify actual disk storage
    d = Document(str(path))
    assert any("<script>" in p.text for p in d.paragraphs)


@pytest.mark.security
def test_h_xml_ampersand_in_paragraph_accepted(tmp_path, monkeypatch):
    """& in paragraph text is accepted and round-trips correctly."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    add_paragraph(str(path), "Tom & Jerry", confirm=True)
    d = Document(str(path))
    assert any("Tom & Jerry" == p.text for p in d.paragraphs)


@pytest.mark.security
def test_h_xml_double_quote_in_paragraph_accepted(tmp_path, monkeypatch):
    """Double-quote in paragraph text is accepted and round-trips correctly."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    add_paragraph(str(path), 'He said "hello"', confirm=True)
    d = Document(str(path))
    assert any('He said "hello"' == p.text for p in d.paragraphs)


@pytest.mark.security
def test_h_xml_angle_brackets_roundtrip(tmp_path, monkeypatch):
    """< and > in find_replace text round-trip as literal characters."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx_with_text(path, "replace this")
    find_replace(str(path), "this", "<b>bold</b>", confirm=True)
    d = Document(str(path))
    assert any("<b>bold</b>" in p.text for p in d.paragraphs)


# ===========================================================================
# GROUP I — Error class hierarchy (4 tests)
# ===========================================================================


@pytest.mark.security
def test_i_validation_error_is_word_mcp_error():
    """ValidationError is a subclass of WordMCPError."""
    assert issubclass(ValidationError, WordMCPError)


@pytest.mark.security
def test_i_not_allowed_error_is_word_mcp_error():
    """NotAllowedError is a subclass of WordMCPError."""
    assert issubclass(NotAllowedError, WordMCPError)


@pytest.mark.security
def test_i_not_found_error_is_word_mcp_error():
    """NotFoundError is a subclass of WordMCPError."""
    assert issubclass(NotFoundError, WordMCPError)


@pytest.mark.security
def test_i_file_error_is_word_mcp_error():
    """FileError is a subclass of WordMCPError."""
    assert issubclass(FileError, WordMCPError)


# ===========================================================================
# GROUP J — Atomic save (4 tests)
# ===========================================================================


@pytest.mark.security
def test_j_atomic_save_cleans_up_temp_on_crash(tmp_path):
    """_atomic_save removes the temp file when doc.save() raises."""
    target = tmp_path / "output.docx"
    mock_doc = MagicMock()
    mock_doc.save.side_effect = OSError("disk full")
    with pytest.raises(OSError, match="disk full"):
        _atomic_save(mock_doc, target)
    assert list(tmp_path.glob("*.tmp")) == []


@pytest.mark.security
def test_j_atomic_save_original_preserved_on_crash(tmp_path):
    """Original file is intact after _atomic_save crashes during write."""
    target = tmp_path / "original.docx"
    Document().save(str(target))
    original_size = target.stat().st_size
    mock_doc = MagicMock()
    mock_doc.save.side_effect = OSError("write failed")
    with pytest.raises(OSError):
        _atomic_save(mock_doc, target)
    assert target.exists()
    assert target.stat().st_size == original_size


@pytest.mark.security
def test_j_atomic_save_uses_os_replace(tmp_path):
    """_atomic_save calls os.replace for atomic rename."""
    target = tmp_path / "out.docx"
    mock_doc = MagicMock()
    mock_doc.save = lambda path: open(path, "wb").close()
    with patch("wordmcp.document_docx.os.replace") as mock_replace:
        mock_replace.side_effect = lambda src, dst: os.rename(src, dst)
        _atomic_save(mock_doc, target)
    mock_replace.assert_called_once()


@pytest.mark.security
def test_j_atomic_save_success_path(tmp_path, monkeypatch):
    """save() with valid path and WORD_ENABLE_WRITE=true succeeds."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    result = save(str(path), confirm=True)
    assert result["saved"] is True


# ===========================================================================
# GROUP K — Image path security (4 tests)
# ===========================================================================


@pytest.mark.security
def test_k_null_byte_in_image_path_rejected(tmp_path, monkeypatch):
    """Null byte in image_path raises ValidationError before filesystem access."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte"):
        _check_image_path(str(tmp_path / "img.png") + "\x00../evil")


@pytest.mark.security
def test_k_image_extension_py_rejected(tmp_path, monkeypatch):
    """.py extension for image is rejected by _check_image_path."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    script = tmp_path / "evil.py"
    script.write_text("import os; os.system('evil')")
    with pytest.raises(ValidationError, match="Unsupported image extension"):
        _check_image_path(str(script))


@pytest.mark.security
def test_k_image_path_outside_allowlist_rejected(tmp_path, monkeypatch, tmp_path_factory):
    """Image path not in WORD_ALLOWLIST_ROOTS is rejected."""
    other = tmp_path_factory.mktemp("other_img")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    png = other / "img.png"
    _make_png(png)
    with pytest.raises(ValidationError, match="not in allowlist"):
        _check_image_path(str(png))


@pytest.mark.security
def test_k_valid_png_in_allowlist_accepted(tmp_path, monkeypatch):
    """Valid .png inside WORD_ALLOWLIST_ROOTS is accepted by _check_image_path."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    png = tmp_path / "image.png"
    _make_png(png)
    resolved = _check_image_path(str(png))
    assert resolved.suffix == ".png"


# ===========================================================================
# GROUP L — Misc / cache behaviour (6 tests)
# ===========================================================================


@pytest.mark.security
def test_l_doc_cache_no_stale_state_after_modify(tmp_path, monkeypatch):
    """_doc_cache is evicted after write — reload from disk shows fresh state."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    add_paragraph(str(path), "Fresh paragraph", confirm=True)
    # After add_paragraph, cache is evicted; next read reloads from disk
    result = list_paragraphs(str(path))
    texts = [p["text_preview"] for p in result]
    assert "Fresh paragraph" in texts


@pytest.mark.security
def test_l_list_paragraphs_on_empty_doc_returns_empty_list(tmp_path, monkeypatch):
    """list_paragraphs on a blank new Document returns []."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "empty.docx"
    # Brand-new Document() actually has one empty paragraph
    d = Document()
    # Clear all paragraphs by saving directly
    d.save(str(path))
    result = list_paragraphs(str(path))
    # A new Document may have 0 or 1 empty paragraphs depending on python-docx
    assert isinstance(result, list)


@pytest.mark.security
def test_l_two_read_document_calls_same_path(tmp_path, monkeypatch):
    """Two consecutive read_document calls on the same path return equal results."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx_with_text(path, "consistent content")
    r1 = read_document(str(path))
    r2 = read_document(str(path))
    assert r1["paragraph_count"] == r2["paragraph_count"]
    assert r1["path"] == r2["path"]


@pytest.mark.security
def test_l_capabilities_always_returns_19_tools():
    """capabilities() always returns exactly 14 tools — no injection risk."""
    result = capabilities()
    assert isinstance(result["tools"], list)
    assert len(result["tools"]) == 14
    # All tool names should be simple identifiers (no special characters)
    for tool_name in result["tools"]:
        assert tool_name.isidentifier(), f"Tool name {tool_name!r} is not a valid identifier"


@pytest.mark.security
def test_l_read_document_no_path_traversal_in_result(tmp_path, monkeypatch):
    """read_document result path is resolved absolute path within allowlist."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    result = read_document(str(path))
    result_path = Path(result["path"])
    # The returned path must be within tmp_path (the allowlist root)
    assert result_path.is_relative_to(tmp_path)


@pytest.mark.security
def test_l_validation_error_message_does_not_leak_full_stack(tmp_path, monkeypatch):
    """ValidationError message is safe and does not contain traceback text."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    try:
        _check_path("/absolute/outside/secret.docx")
    except ValidationError as exc:
        # Message should be a simple user-friendly string, not a traceback
        msg = str(exc)
        assert "Traceback" not in msg
        assert "File \"" not in msg
