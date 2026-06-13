"""
Pytest fixtures for wordmcp unit and smoke tests.

All fixtures are self-contained: they create temporary files using tmp_path,
monkeypatch the required env vars, and clean up automatically via pytest's
tmp_path scope. No binary files are committed to the repository.
"""
from __future__ import annotations

import struct
import zlib

import pytest
from docx import Document
from unittest.mock import MagicMock


def pytest_configure(config: pytest.Config) -> None:
    config.addinivalue_line("markers", "unit: unit tests (no live Office required)")
    config.addinivalue_line("markers", "smoke: smoke tests")
    config.addinivalue_line("markers", "integration: requires live Word/Office")
    config.addinivalue_line("markers", "security: marks a security test")


@pytest.fixture(autouse=True)
def _clear_doc_cache():
    """Clear the document cache before and after each test to prevent state leakage."""
    from wordmcp import document_docx
    document_docx._doc_cache.clear()
    yield
    document_docx._doc_cache.clear()


@pytest.fixture
def tmp_docx(tmp_path, monkeypatch):
    """Minimal .docx with 3 body paragraphs; WORD_ALLOWLIST_ROOTS set to tmp_path.

    Structure:
      index 0 — style "Title", text "Test Document"
      index 1 — style "Heading 1", text "Section One"
      index 2 — style "Normal", text "Hello " + "world" (bold run)
    """
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    d = Document()
    d.add_heading("Test Document", level=0)   # index 0 -> "Title"
    d.add_heading("Section One", level=1)     # index 1 -> "Heading 1"
    para = d.add_paragraph()                  # index 2 -> "Normal"
    para.add_run("Hello ")
    bold_run = para.add_run("world")
    bold_run.bold = True
    d.save(str(path))
    return str(path)


@pytest.fixture
def word_backend(tmp_docx, monkeypatch):
    """Alias for tmp_docx for naming clarity in test_unit.py."""
    return tmp_docx


@pytest.fixture
def read_only_backend(tmp_docx):
    """tmp_docx without WORD_ENABLE_WRITE — read operations only."""
    return tmp_docx


@pytest.fixture
def sample_docx_with_table(tmp_path, monkeypatch):
    """Minimal .docx with 1 body paragraph and a 2x3 Table Grid table.

    Table data:
      row 0: ["Name", "Age", "City"]
      row 1: ["Alice", "30",  "London"]
    """
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test_table.docx"
    d = Document()
    d.add_paragraph("Before table")
    tbl = d.add_table(2, 3, style="Table Grid")
    rows_data = [["Name", "Age", "City"], ["Alice", "30", "London"]]
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            tbl.rows[r_idx].cells[c_idx].text = val
    d.save(str(path))
    return str(path)


@pytest.fixture
def write_enabled_docx(tmp_docx, monkeypatch):
    """tmp_docx with WORD_ENABLE_WRITE=true set.

    Use this fixture for tests that exercise write tool success paths.
    Pass confirm=True explicitly in each test — the fixture does NOT
    pass confirm on the caller's behalf.
    """
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    return tmp_docx


@pytest.fixture
def write_enabled_docx_with_table(sample_docx_with_table, monkeypatch):
    """sample_docx_with_table with WORD_ENABLE_WRITE=true.

    Required for write tests targeting table content (e.g. test_find_replace_in_table).
    """
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    return sample_docx_with_table


@pytest.fixture
def tmp_png(tmp_path):
    """Returns path (str) to a valid minimal 1x1 RGB PNG in tmp_path.

    Constructed entirely from stdlib (no Pillow dependency). The PNG is
    a 1x1 pixel image with a single white pixel.

    Byte layout:
      PNG signature
      IHDR chunk: width=1 height=1 bitdepth=8 colortype=2 (RGB)
      IDAT chunk: zlib-compressed scanline \\x00 (filter=None) + \\xff\\xff\\xff (RGB white)
      IEND chunk
    """

    def _chunk(name: bytes, data: bytes) -> bytes:
        c = name + data
        return (
            struct.pack(">I", len(data))
            + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )

    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat_raw = b"\x00\xff\xff\xff"   # filter byte 0x00 + RGB(255,255,255)
    img_bytes = (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr_data)
        + _chunk(b"IDAT", zlib.compress(idat_raw))
        + _chunk(b"IEND", b"")
    )
    p = tmp_path / "test_image.png"
    p.write_bytes(img_bytes)
    return str(p)


# ── COM fixtures (Sprint B) ──────────────────────────────────────────────────

@pytest.fixture
def mock_revision():
    rev = MagicMock()
    rev.Author = "Test Author"
    rev.Date = "2025-01-01"
    rev.Type = 1
    rev.Range.Text = "sample tracked text"
    return rev


@pytest.fixture
def mock_word_doc(mock_revision):
    doc = MagicMock()
    doc.Revisions.__iter__ = MagicMock(return_value=iter([mock_revision]))
    doc.Revisions.AcceptAll = MagicMock()
    doc.Revisions.RejectAll = MagicMock()
    doc.Save = MagicMock()
    doc.Close = MagicMock()
    doc.ExportAsFixedFormat = MagicMock()
    return doc


@pytest.fixture
def mock_word_app(mock_word_doc):
    app = MagicMock()
    app.Documents.Open.return_value = mock_word_doc
    app.Quit = MagicMock()
    return app


@pytest.fixture
def patch_word_com(monkeypatch, mock_word_app):
    import contextlib
    from wordmcp._com import _base as _dc
    from wordmcp._com import _read as _dc_read
    from wordmcp._com import _write as _dc_write
    from wordmcp._com import _export as _dc_export

    @contextlib.contextmanager
    def _fake_app_context():
        yield mock_word_app

    monkeypatch.setattr(_dc, "_word_app_context", _fake_app_context)
    monkeypatch.setattr(_dc, "_COM_AVAILABLE", True)
    monkeypatch.setattr(_dc_read, "_word_app_context", _fake_app_context)
    monkeypatch.setattr(_dc_write, "_word_app_context", _fake_app_context)
    monkeypatch.setattr(_dc_export, "_word_app_context", _fake_app_context)
    return mock_word_app


@pytest.fixture
def review_docx(tmp_path, monkeypatch):
    """A .docx with H1 heading, H2 heading, and 3 Normal paragraphs."""
    from docx import Document
    path = tmp_path / "review_test.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_heading("Background", level=2)
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    doc.save(str(path))
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    return path


@pytest.fixture
def skipped_h3_docx(tmp_path, monkeypatch):
    """A .docx with H1 → H3 (skips H2) — triggers heading_hierarchy check failure."""
    from docx import Document
    path = tmp_path / "skipped_h3.docx"
    doc = Document()
    doc.add_heading("Top Level", level=1)
    doc.add_heading("Skipped Level", level=3)  # H1 → H3, skips H2
    doc.add_paragraph("Some content.")
    doc.save(str(path))
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    return path


@pytest.fixture
def write_review_env(tmp_path, monkeypatch):
    """Sets WORD_ENABLE_WRITE=true and WORD_ALLOWLIST_ROOTS to tmp_path."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    return tmp_path


@pytest.fixture
def com_env(tmp_path, monkeypatch):
    from docx import Document as _Doc
    doc_path = tmp_path / "test_com.docx"
    d = _Doc()
    d.add_paragraph("COM test paragraph")
    d.save(str(doc_path))
    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    return doc_path


@pytest.fixture
def assembly_docx_with_placeholders(tmp_path, monkeypatch):
    """Docx with body paragraphs and table cell containing {{...}} placeholders."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Dear {{FIRST_NAME}} {{LAST_NAME}},")
    doc.add_paragraph("Date: {{DATE}}")
    doc.add_paragraph("No placeholders here.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Invoice: {{INVOICE_NUM}}"
    path = tmp_path / "placeholders.docx"
    doc.save(str(path))
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    return str(path)


@pytest.fixture
def assembly_docx_clean(tmp_path, monkeypatch):
    """Docx with no placeholders — verify should return pass."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("This document is complete and ready.")
    path = tmp_path / "clean.docx"
    doc.save(str(path))
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    return str(path)


def _make_large_replacements(n: int) -> dict[str, str]:
    return {f"{{{{TOKEN_{i}}}}}": f"value_{i}" for i in range(n)}
