from __future__ import annotations

import json
from typing import Any

import pytest
from docx import Document


def _call_tool(tool, *args, **kwargs):
    fn = getattr(tool, "fn", tool)
    return fn(*args, **kwargs)


def _get_server_tool_names(mcp_obj) -> set[str]:
    if hasattr(mcp_obj, "_tool_manager") and hasattr(mcp_obj._tool_manager, "_tools"):
        return set(mcp_obj._tool_manager._tools)
    if hasattr(mcp_obj, "_tools"):
        return set(mcp_obj._tools)
    if hasattr(mcp_obj, "_local_provider") and hasattr(mcp_obj._local_provider, "_components"):
        return {
            key.split(":")[1].split("@")[0]
            for key in mcp_obj._local_provider._components
            if key.startswith("tool:")
        }
    raise RuntimeError("Unable to discover FastMCP tool registry")


def _make_runtime(doc_backend: Any) -> Any:
    from wordmcp._server_runtime import ServerRuntime
    from wordmcp.document_docx import WordMCPError

    def _require_com_stub() -> None:
        raise AssertionError("COM should not be required in context tool tests")

    return ServerRuntime(
        get_doc=lambda: doc_backend,
        get_word_error=lambda: WordMCPError,
        get_com=lambda: None,
        get_com_loaded=lambda: False,
        require_com=_require_com_stub,
    )


@pytest.fixture
def context_docx(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

    path = tmp_path / "context-sample.docx"
    doc = Document()
    props = doc.core_properties
    props.title = "Context Smoke"
    props.author = "Copilot"
    props.subject = "Context Tool"
    props.category = "QA"
    props.keywords = "context,focused"
    props.last_modified_by = "Copilot"
    props.revision = 7

    doc.add_heading("Executive Summary", level=0)
    doc.add_heading("Risk Review", level=1)
    doc.add_paragraph("BODY SECRET SHOULD NOT LEAK")
    doc.add_paragraph("ANOTHER BODY PARAGRAPH")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "TABLE SECRET"
    table.cell(0, 1).text = "HIDDEN CELL"
    doc.save(path)
    return path


@pytest.mark.unit
def test_get_document_context_returns_focused_packet_for_allowlisted_docx(context_docx) -> None:
    from wordmcp.document_docx import get_document_context

    result = get_document_context(str(context_docx))

    assert result["packet_type"] == "focused"
    assert result["adapter_type"] == "wordmcp"
    assert result["artifact_type"] == "document"
    assert result["source_ref"] == str(context_docx.resolve())
    assert result["artifact_id"] == f"doc:{context_docx.stem}"

    summary = result["summary"]
    assert isinstance(summary, str)
    assert summary
    assert summary.startswith("context-sample.docx: 4 paragraphs")
    assert "1 tables" in summary
    assert "1 sections" in summary
    assert "2 headings" in summary
    assert "title=Context Smoke" in summary
    assert "author=Copilot" in summary
    assert "headings=Executive Summary, Risk Review" in summary

    content = result["content"]
    assert set(content) == {
        "title",
        "author",
        "subject",
        "category",
        "keywords",
        "last_modified_by",
        "revision",
        "created",
        "modified",
        "paragraph_count",
        "table_count",
        "section_count",
        "heading_count",
        "headings",
    }
    assert content["title"] == "Context Smoke"
    assert content["author"] == "Copilot"
    assert content["subject"] == "Context Tool"
    assert content["category"] == "QA"
    assert content["keywords"] == "context,focused"
    assert content["last_modified_by"] == "Copilot"
    assert content["revision"] == 7
    assert content["table_count"] == 1
    assert content["heading_count"] == 2
    assert content["headings"] == ["Executive Summary", "Risk Review"]

    serialized = json.dumps(content)
    assert "BODY SECRET SHOULD NOT LEAK" not in summary
    assert "ANOTHER BODY PARAGRAPH" not in summary
    assert "TABLE SECRET" not in summary
    assert "HIDDEN CELL" not in summary
    assert "BODY SECRET SHOULD NOT LEAK" not in serialized
    assert "ANOTHER BODY PARAGRAPH" not in serialized
    assert "TABLE SECRET" not in serialized
    assert "HIDDEN CELL" not in serialized
    assert "Executive Summary" in serialized
    assert "Risk Review" in serialized


@pytest.mark.unit
def test_get_document_context_denies_path_outside_allowlist(tmp_path, monkeypatch) -> None:
    from wordmcp.document_docx import ValidationError, get_document_context

    allowed = tmp_path / "allowed"
    blocked = tmp_path / "blocked"
    allowed.mkdir()
    blocked.mkdir()

    path = blocked / "outside.docx"
    Document().save(path)
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(allowed))

    with pytest.raises(ValidationError, match="Path not in allowlist"):
        get_document_context(str(path))


@pytest.mark.unit
def test_server_context_tool_wrapper_calls_runtime_doc_lookup() -> None:
    from wordmcp import _server_context_tools

    calls: list[tuple[str, str]] = []

    class StubDocBackend:
        def get_document_context(self, path: str) -> dict[str, object]:
            calls.append(("get_document_context", path))
            return {"ok": True, "path": path}

    runtime = _make_runtime(StubDocBackend())
    result = _server_context_tools.get_document_context(runtime, "sample.docx")

    assert result == {"ok": True, "path": "sample.docx"}
    assert calls == [("get_document_context", "sample.docx")]


@pytest.mark.unit
def test_manifest_includes_get_document_context() -> None:
    from wordmcp._server_manifest import build_capabilities

    class StubDocBackend:
        @staticmethod
        def capabilities() -> dict[str, object]:
            return {
                "phase": "1.0",
                "backend": "python-docx",
                "tools": ["capabilities"],
                "governance": {},
            }

    result = build_capabilities(_make_runtime(StubDocBackend()))

    assert "get_document_context" in result["context_tools"]["tools"]
    assert "get_document_context" in result["tools"]


@pytest.mark.smoke
def test_server_registers_and_executes_get_document_context(context_docx) -> None:
    from wordmcp import server as srv

    assert hasattr(srv, "get_document_context")
    assert "get_document_context" in _get_server_tool_names(srv.mcp)

    result = _call_tool(srv.get_document_context, str(context_docx))

    assert result["packet_type"] == "focused"
    assert result["adapter_type"] == "wordmcp"
    assert result["artifact_type"] == "document"
    assert result["source_ref"] == str(context_docx.resolve())
    assert isinstance(result["summary"], str)
    assert result["summary"].startswith("context-sample.docx: 4 paragraphs")
    assert "headings=Executive Summary, Risk Review" in result["summary"]
    assert result["content"]["headings"] == ["Executive Summary", "Risk Review"]

    serialized = json.dumps(result["content"])
    assert "BODY SECRET SHOULD NOT LEAK" not in result["summary"]
    assert "BODY SECRET SHOULD NOT LEAK" not in serialized
    assert "TABLE SECRET" not in result["summary"]
    assert "TABLE SECRET" not in serialized


@pytest.mark.unit
def test_artifact_id_starts_with_doc_prefix(context_docx) -> None:
    from wordmcp.document_docx import get_document_context

    result = get_document_context(str(context_docx))
    assert result["artifact_id"].startswith("doc:")


@pytest.mark.unit
def test_artifact_id_equals_doc_stem_format(context_docx) -> None:
    from wordmcp.document_docx import get_document_context

    result = get_document_context(str(context_docx))
    assert result["artifact_id"] == f"doc:{context_docx.stem}"
    # Explicit: == "doc:context-sample"


@pytest.mark.unit
def test_word_artifact_id_prefix_matches_workflowruntime_constant() -> None:
    """W5 parity: _ARTIFACT_ID_PREFIX_DOC (wordmcp) == ARTIFACT_ID_PREFIX_DOC (workflowruntime)."""
    _wfr = pytest.importorskip("workflowruntime.contracts.constants")
    from wordmcp._docx.context import _ARTIFACT_ID_PREFIX_DOC as wordmcp_prefix
    wfr_prefix = _wfr.ARTIFACT_ID_PREFIX_DOC

    assert wordmcp_prefix == wfr_prefix


@pytest.mark.unit
def test_artifact_id_sanitizes_formula_injection_stem(tmp_path, monkeypatch) -> None:
    """H-001: stems starting with =, +, -, @ must be stripped."""
    from docx import Document as DocxDocument
    from wordmcp.document_docx import get_document_context

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "=evil.docx"
    DocxDocument().save(path)
    result = get_document_context(str(path))
    assert not result["artifact_id"].startswith("doc:=")
    assert result["artifact_id"] == "doc:evil"
