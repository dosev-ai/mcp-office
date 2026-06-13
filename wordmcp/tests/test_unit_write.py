"""
Unit tests — write operations: write guards, add_paragraph/heading/page_break/table,
insert_image, find_replace, save, BLOCKER tests (BLOCKER-3, BLOCKER-4, BLOCKER-5),
W3 GAP-FILL supplements for write groups.
"""
from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from wordmcp.document_docx import (
    NotAllowedError,
    ValidationError,
    _atomic_save,
    add_heading,
    add_page_break,
    add_paragraph,
    add_table,
    find_replace,
    insert_image,
    read_document,
    save,
)


# ===========================================================================
# GROUP 9 — Write guard ENABLE_WRITE unset (4 tests)
# ===========================================================================


@pytest.mark.unit
def test_add_paragraph_requires_enable_write(tmp_docx):
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        add_paragraph(tmp_docx, "hello", confirm=True)


@pytest.mark.unit
def test_add_table_requires_enable_write(tmp_docx):
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        add_table(tmp_docx, 1, 1, confirm=True)


@pytest.mark.unit
def test_find_replace_requires_enable_write(tmp_docx):
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        find_replace(tmp_docx, "Test", "New", confirm=True)


@pytest.mark.unit
def test_save_requires_enable_write(tmp_docx):
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        save(tmp_docx, confirm=True)


# ===========================================================================
# GROUP 10 — confirm=False guard (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_add_paragraph_requires_confirm(write_enabled_docx):
    with pytest.raises(ValidationError, match="confirm=True required"):
        add_paragraph(write_enabled_docx, "hello", confirm=False)


@pytest.mark.unit
def test_add_heading_requires_confirm(write_enabled_docx):
    with pytest.raises(ValidationError, match="confirm=True required"):
        add_heading(write_enabled_docx, "My Heading", level=2, confirm=False)


@pytest.mark.unit
def test_save_requires_confirm(write_enabled_docx):
    with pytest.raises(ValidationError, match="confirm=True required"):
        save(write_enabled_docx, confirm=False)


# ===========================================================================
# GROUP 11 — add_paragraph (4 tests)
# ===========================================================================


@pytest.mark.unit
def test_add_paragraph_returns_index(write_enabled_docx):
    result = add_paragraph(write_enabled_docx, "New paragraph", confirm=True)
    assert result["paragraph_index"] == 3  # was 3 paragraphs before


@pytest.mark.unit
def test_add_paragraph_with_style(write_enabled_docx):
    result = add_paragraph(write_enabled_docx, "Styled", style="Normal", confirm=True)
    assert result["style"] == "Normal"
    assert result["text"] == "Styled"


@pytest.mark.unit
def test_add_paragraph_invalid_style(write_enabled_docx):
    with pytest.raises(ValidationError, match="Unknown paragraph style"):
        add_paragraph(write_enabled_docx, "text", style="NonExistentStyle999", confirm=True)


@pytest.mark.unit
def test_add_paragraph_persists_to_disk(write_enabled_docx):
    add_paragraph(write_enabled_docx, "Persisted text", confirm=True)
    # Reload from disk
    d = Document(write_enabled_docx)
    texts = [p.text for p in d.paragraphs]
    assert "Persisted text" in texts


# ===========================================================================
# GROUP 12 — add_heading (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_add_heading_returns_level(write_enabled_docx):
    result = add_heading(write_enabled_docx, "My Heading", level=2, confirm=True)
    assert result["level"] == 2
    assert result["text"] == "My Heading"


@pytest.mark.unit
def test_add_heading_invalid_level(write_enabled_docx):
    with pytest.raises(ValidationError, match="level must be 0-9"):
        add_heading(write_enabled_docx, "Bad", level=10, confirm=True)


@pytest.mark.unit
def test_add_heading_level_zero_title_style(write_enabled_docx):
    result = add_heading(write_enabled_docx, "Main Title", level=0, confirm=True)
    assert result["style"] == "Title"


# ===========================================================================
# GROUP 13 — add_page_break (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_add_page_break_returns_inserted_true(write_enabled_docx):
    result = add_page_break(write_enabled_docx, confirm=True)
    assert result["inserted"] is True


@pytest.mark.unit
def test_add_page_break_increments_paragraph_count(write_enabled_docx):
    before = read_document(write_enabled_docx)["paragraph_count"]
    add_page_break(write_enabled_docx, confirm=True)
    d = Document(write_enabled_docx)
    after = len(d.paragraphs)
    assert after > before


# ===========================================================================
# GROUP 14 — add_table (4 tests)
# ===========================================================================


@pytest.mark.unit
def test_add_table_returns_table_index(write_enabled_docx):
    result = add_table(write_enabled_docx, rows=2, cols=2, confirm=True)
    assert result["table_index"] == 0  # no tables before
    assert result["rows"] == 2
    assert result["cols"] == 2


@pytest.mark.unit
def test_add_table_with_data(write_enabled_docx):
    data = [["A", "B"], ["C", "D"]]
    result = add_table(write_enabled_docx, rows=2, cols=2, data=data, confirm=True)
    assert result["rows"] == 2
    # Verify data persisted
    d = Document(write_enabled_docx)
    tbl = d.tables[0]
    assert tbl.rows[0].cells[0].text == "A"
    assert tbl.rows[1].cells[1].text == "D"


@pytest.mark.unit
def test_add_table_data_wrong_row_count(write_enabled_docx):
    with pytest.raises(ValidationError, match="rows"):
        add_table(write_enabled_docx, rows=2, cols=2, data=[["A", "B"]], confirm=True)


@pytest.mark.unit
def test_add_table_data_wrong_col_count(write_enabled_docx):
    with pytest.raises(ValidationError, match="cols"):
        add_table(
            write_enabled_docx, rows=1, cols=3, data=[["A", "B"]], confirm=True
        )


# ===========================================================================
# GROUP 15 — insert_image (4 tests)
# ===========================================================================


@pytest.mark.unit
def test_insert_image_success(write_enabled_docx, tmp_png):
    result = insert_image(write_enabled_docx, tmp_png, confirm=True)
    assert result["inserted"] is True


@pytest.mark.unit
def test_insert_image_invalid_extension(write_enabled_docx, tmp_path):
    svg_path = tmp_path / "icon.svg"
    svg_path.write_text("<svg/>")
    with pytest.raises(ValidationError, match="Unsupported image extension"):
        insert_image(write_enabled_docx, str(svg_path), confirm=True)


@pytest.mark.unit
def test_insert_image_not_in_allowlist(write_enabled_docx, tmp_path, monkeypatch):
    other = tmp_path / "other"
    other.mkdir()
    png = other / "img.png"
    png.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)
    # WORD_ALLOWLIST_ROOTS is set to tmp_path (parent of other),
    # so this should actually be allowed. Use a path truly outside:
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path / "subdir_only"))
    with pytest.raises(ValidationError):
        insert_image(write_enabled_docx, str(png), confirm=True)


@pytest.mark.unit
def test_insert_image_invalid_width(write_enabled_docx, tmp_png):
    with pytest.raises(ValidationError, match="width_inches must be > 0"):
        insert_image(write_enabled_docx, tmp_png, width_inches=-1.0, confirm=True)


# ===========================================================================
# GROUP 16 — find_replace (4 tests incl. BLOCKER-6 XML injection test)
# ===========================================================================


@pytest.mark.unit
def test_find_replace_returns_count(write_enabled_docx):
    result = find_replace(write_enabled_docx, "Section", "Chapter", confirm=True)
    assert result["replacements_made"] >= 1


@pytest.mark.unit
def test_find_replace_empty_find_raises(write_enabled_docx):
    with pytest.raises(ValidationError, match="find_text must not be empty"):
        find_replace(write_enabled_docx, "", "something", confirm=True)


@pytest.mark.unit
def test_find_replace_in_table(write_enabled_docx_with_table):
    result = find_replace(write_enabled_docx_with_table, "Alice", "Bob", confirm=True)
    assert result["replacements_made"] == 1
    # Verify on disk
    d = Document(write_enabled_docx_with_table)
    cell_text = d.tables[0].rows[1].cells[0].text
    assert cell_text == "Bob"


@pytest.mark.unit
def test_find_replace_replace_text_xml_chars(tmp_path, monkeypatch):
    """replace_text containing XML special chars must be stored as literal text (BLOCKER-6)."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "xmltest.docx"
    d = Document()
    d.add_paragraph("Hello world")
    d.save(str(path))
    result = find_replace(str(path), "Hello", "<b>Hi</b>", confirm=True)
    assert result["replacements_made"] == 1
    # Reload from disk and verify literal angle brackets are preserved
    d2 = Document(str(path))
    assert d2.paragraphs[0].text == "<b>Hi</b> world"


# ===========================================================================
# GROUP 17 — save (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_save_returns_saved_true(write_enabled_docx):
    result = save(write_enabled_docx, confirm=True)
    assert result["saved"] is True


@pytest.mark.unit
def test_save_path_matches(write_enabled_docx):
    result = save(write_enabled_docx, confirm=True)
    assert result["path"] == str(Path(write_enabled_docx).resolve())


@pytest.mark.unit
def test_atomic_save_cleanup_on_error(tmp_path):
    """_atomic_save cleans up temp file when doc.save() raises OSError (BLOCKER-2)."""
    target = tmp_path / "output.docx"
    mock_doc = MagicMock()
    mock_doc.save.side_effect = OSError("disk full")
    with pytest.raises(OSError, match="disk full"):
        _atomic_save(mock_doc, target)
    # No .tmp files should remain after _atomic_save's except-block cleanup
    assert list(tmp_path.glob("*.tmp")) == []


# ===========================================================================
# GROUP 18 — BLOCKER-3 confirm default (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_document_docx_default_confirm_is_false(tmp_path, monkeypatch):
    """document_docx.add_paragraph called without confirm arg must raise (BLOCKER-3)."""
    import wordmcp.document_docx as dd

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    Document().save(str(path))
    with pytest.raises(dd.ValidationError, match="confirm=True required"):
        dd.add_paragraph(str(path), "x")  # confirm not passed -> default False


# ===========================================================================
# GROUP 19 — BLOCKER-5 allowlist whitespace (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_allowlist_root_with_surrounding_whitespace(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", f"  {str(tmp_path)}  ")
    path = tmp_path / "test.docx"
    Document().save(str(path))
    result = read_document(str(path))
    assert result["path"] == str(path.resolve())


@pytest.mark.unit
def test_allowlist_second_root_with_leading_space(tmp_path, monkeypatch, tmp_path_factory):
    second = tmp_path_factory.mktemp("second")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", f"{str(tmp_path)}, {str(second)}")
    path = second / "test2.docx"
    Document().save(str(path))
    result = read_document(str(path))
    assert result["path"] == str(path.resolve())


# ===========================================================================
# BLOCKER-4 — no print() calls in src/
# ===========================================================================


@pytest.mark.unit
def test_no_print_calls_in_src():
    """No bare print() calls should exist anywhere under src/wordmcp/ (BLOCKER-4)."""
    import ast
    import pathlib

    src_root = pathlib.Path(__file__).parent.parent / "src" / "wordmcp"
    for py_file in src_root.rglob("*.py"):
        tree = ast.parse(py_file.read_text(encoding="utf-8"), filename=str(py_file))
        for node in ast.walk(tree):
            if (
                isinstance(node, ast.Call)
                and isinstance(node.func, ast.Name)
                and node.func.id == "print"
            ):
                raise AssertionError(
                    f"Forbidden print() call found in {py_file}:{node.lineno}"
                )


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 10
# ===========================================================================


@pytest.mark.unit
def test_confirm_false_add_table_raises(write_enabled_docx):
    with pytest.raises(ValidationError, match="confirm=True required"):
        add_table(write_enabled_docx, rows=1, cols=1, confirm=False)


@pytest.mark.unit
def test_confirm_false_find_replace_raises(write_enabled_docx):
    with pytest.raises(ValidationError, match="confirm=True required"):
        find_replace(write_enabled_docx, "Section", "Chapter", confirm=False)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 11
# ===========================================================================


@pytest.mark.unit
def test_add_paragraph_path_not_in_allowlist_raises(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    outside = str(tmp_path.parent / "outside.docx")
    with pytest.raises(ValidationError, match="not in allowlist"):
        add_paragraph(outside, "text", confirm=True)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 13
# ===========================================================================


@pytest.mark.unit
def test_add_page_break_write_blocked_raises(tmp_docx):
    """add_page_break must raise NotAllowedError when WORD_ENABLE_WRITE is unset."""
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        add_page_break(tmp_docx, confirm=True)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 14
# ===========================================================================


@pytest.mark.unit
def test_add_table_invalid_style_graceful(write_enabled_docx):
    with pytest.raises(ValidationError):
        add_table(
            write_enabled_docx,
            rows=2,
            cols=2,
            style="NonExistentTableStyle999",
            confirm=True,
        )


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 15
# ===========================================================================


@pytest.mark.unit
def test_insert_image_width_applied(write_enabled_docx, tmp_png):
    result = insert_image(write_enabled_docx, tmp_png, width_inches=2.0, confirm=True)
    assert result["inserted"] is True
    assert result["width_inches"] == 2.0


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 16
# ===========================================================================


@pytest.mark.unit
def test_find_replace_zero_replacements_on_no_match(write_enabled_docx):
    result = find_replace(
        write_enabled_docx, "THIS_TEXT_DOES_NOT_EXIST_XYZ999", "X", confirm=True
    )
    assert result["replacements_made"] == 0


@pytest.mark.unit
def test_find_replace_text_actually_replaced_roundtrip(write_enabled_docx):
    find_replace(write_enabled_docx, "Section", "Chapter", confirm=True)
    d = Document(write_enabled_docx)
    texts = [p.text for p in d.paragraphs]
    assert any("Chapter" in t for t in texts)
    assert not any("Section" in t for t in texts)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 17
# ===========================================================================


@pytest.mark.unit
def test_save_path_not_in_allowlist_raises(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    outside = str(tmp_path.parent / "outside.docx")
    with pytest.raises(ValidationError, match="not in allowlist"):
        save(outside, confirm=True)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 18 (BLOCKER-3)
# ===========================================================================


@pytest.mark.unit
def test_find_replace_default_confirm_is_false(tmp_path, monkeypatch):
    """document_docx.find_replace called without confirm must raise (BLOCKER-3)."""
    import wordmcp.document_docx as dd

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    Document().save(str(path))
    with pytest.raises(dd.ValidationError, match="confirm=True required"):
        dd.find_replace(str(path), "x", "y")  # confirm not passed -> default False


# ===========================================================================
# DECOMP GUARD — shim re-export completeness (write.py split into sub-modules)
# ===========================================================================

_EXPECTED_WRITE_EXPORTS = [
    "add_heading",
    "add_list",
    "add_page_break",
    "add_paragraph",
    "add_table",
    "bulk_add_paragraphs",
    "bulk_update_paragraphs",
    "bulk_update_table_cells",
    "create_document",
    "delete_paragraph",
    "find_replace",
    "insert_image",
    "insert_paragraph",
    "save",
    "set_document_properties",
    "set_paragraph_format",
    "update_paragraph",
    "update_table_cell",
]


@pytest.mark.unit
def test_write_shim_exports_all_18_names():
    """All 18 public names must be importable from wordmcp._docx.write (decomp guard).

    This test prevents future regressions where a name is accidentally dropped
    from the re-export shim when sub-modules are edited.
    """
    import importlib

    shim = importlib.import_module("wordmcp._docx.write")
    missing = [name for name in _EXPECTED_WRITE_EXPORTS if not hasattr(shim, name)]
    assert missing == [], f"Names missing from write.py shim: {missing}"


@pytest.mark.unit
def test_write_shim_all_names_are_callable():
    """Every name in write.__all__ must be callable (not None, not a module)."""
    import importlib

    shim = importlib.import_module("wordmcp._docx.write")
    not_callable = [
        name for name in _EXPECTED_WRITE_EXPORTS
        if not callable(getattr(shim, name, None))
    ]
    assert not_callable == [], f"Non-callable exports in write.py shim: {not_callable}"


@pytest.mark.unit
def test_write_shim_dunder_all_matches_expected():
    """write.__all__ must contain exactly the 18 expected public names — no drift."""
    import importlib

    shim = importlib.import_module("wordmcp._docx.write")
    actual = set(shim.__all__)
    expected = set(_EXPECTED_WRITE_EXPORTS)
    extra = actual - expected
    missing = expected - actual
    assert not extra and not missing, (
        f"write.__all__ drift — extra: {extra}, missing: {missing}"
    )
