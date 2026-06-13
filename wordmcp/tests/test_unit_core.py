"""
Unit tests — core read operations: capabilities, path guards, read_document,
get_document_metadata, list_paragraphs, read_paragraph, list_tables, read_table,
export_as_text. (Groups 0-8b + W3 supplements for those groups.)
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from wordmcp.document_docx import (
    FileError,
    NotFoundError,
    ValidationError,
    _check_path,
    capabilities,
    export_as_text,
    get_document_metadata,
    list_paragraphs,
    list_tables,
    read_document,
    read_paragraph,
    read_table,
)


# ===========================================================================
# GROUP 0 — capabilities (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_capabilities_returns_phase():
    result = capabilities()
    assert result["phase"] == "1.0"


@pytest.mark.unit
def test_capabilities_returns_19_tools():
    result = capabilities()
    assert len(result["tools"]) == 14  # document_docx layer: 14 base Phase-1 tools (export/COM surface now at server level)


@pytest.mark.unit
def test_capabilities_returns_backend_python_docx():
    result = capabilities()
    assert result["backend"] == "python-docx"
    assert "python_docx_version" in result
    assert "governance" in result


# ===========================================================================
# GROUP 1 — _check_path allowlist (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_check_path_rejects_file_outside_allowlist(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    outside = Path("/not/in/allowlist/file.docx")
    with pytest.raises(ValidationError, match="not in allowlist"):
        _check_path(str(outside))


@pytest.mark.unit
def test_check_path_rejects_non_docx_extension(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.pdf"
    with pytest.raises(ValidationError, match="Unsupported extension"):
        _check_path(str(path))


@pytest.mark.unit
def test_check_path_raises_when_no_allowlist_configured(tmp_path, monkeypatch):
    monkeypatch.delenv("WORD_ALLOWLIST_ROOTS", raising=False)
    path = tmp_path / "test.docx"
    with pytest.raises(ValidationError, match="WORD_ALLOWLIST_ROOTS"):
        _check_path(str(path))


# ===========================================================================
# GROUP 2 — read_document (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_read_document_returns_paragraph_count(tmp_docx):
    result = read_document(tmp_docx)
    assert result["paragraph_count"] == 3


@pytest.mark.unit
def test_read_document_returns_section_count(tmp_docx):
    result = read_document(tmp_docx)
    assert result["section_count"] >= 1


@pytest.mark.unit
def test_read_document_raises_for_outside_allowlist(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError):
        read_document("/not/in/allowlist/file.docx")


# ===========================================================================
# GROUP 3 — get_document_metadata (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_get_document_metadata_returns_path(tmp_docx):
    result = get_document_metadata(tmp_docx)
    assert "path" in result
    assert result["path"].endswith("test.docx")


@pytest.mark.unit
def test_get_document_metadata_none_for_unset_fields(tmp_docx):
    result = get_document_metadata(tmp_docx)
    # title is not set -> None
    assert result["title"] is None
    # author may default to "python-docx" or similar; key must be present
    assert "author" in result
    assert "created" in result


# ===========================================================================
# GROUP 4 — list_paragraphs (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_list_paragraphs_length_matches_paragraph_count(tmp_docx):
    result = list_paragraphs(tmp_docx)
    assert len(result) == 3


@pytest.mark.unit
def test_list_paragraphs_index_field(tmp_docx):
    result = list_paragraphs(tmp_docx)
    for i, item in enumerate(result):
        assert item["index"] == i


@pytest.mark.unit
def test_list_paragraphs_outline_level_for_heading(tmp_docx):
    result = list_paragraphs(tmp_docx)
    # index 0 is Title -> outline_level 0
    assert result[0]["outline_level"] == 0
    # index 1 is Heading 1 -> outline_level 1
    assert result[1]["outline_level"] == 1
    # index 2 is Normal -> outline_level None
    assert result[2]["outline_level"] is None


# ===========================================================================
# GROUP 5 — read_paragraph (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_read_paragraph_returns_text(tmp_docx):
    result = read_paragraph(tmp_docx, 0)
    assert result["text"] == "Test Document"


@pytest.mark.unit
def test_read_paragraph_index_out_of_range(tmp_docx):
    with pytest.raises(NotFoundError):
        read_paragraph(tmp_docx, 99)


@pytest.mark.unit
def test_read_paragraph_returns_runs(tmp_docx):
    result = read_paragraph(tmp_docx, 2)
    assert len(result["runs"]) >= 2
    # Find the bold run
    bold_runs = [r for r in result["runs"] if r["bold"] is True]
    assert len(bold_runs) >= 1
    assert bold_runs[0]["text"] == "world"


# ===========================================================================
# GROUP 6 — list_tables (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_list_tables_empty_for_no_tables(tmp_docx):
    result = list_tables(tmp_docx)
    assert result == []


@pytest.mark.unit
def test_list_tables_returns_correct_row_col_count(sample_docx_with_table):
    result = list_tables(sample_docx_with_table)
    assert len(result) == 1
    assert result[0]["rows"] == 2
    assert result[0]["cols"] == 3


# ===========================================================================
# GROUP 7 — read_table (3 tests)
# ===========================================================================


@pytest.mark.unit
def test_read_table_data(sample_docx_with_table):
    result = read_table(sample_docx_with_table, 0)
    assert result["data"][0] == ["Name", "Age", "City"]
    assert result["data"][1] == ["Alice", "30", "London"]


@pytest.mark.unit
def test_read_table_index_out_of_range(sample_docx_with_table):
    with pytest.raises(NotFoundError):
        read_table(sample_docx_with_table, 99)


@pytest.mark.unit
def test_read_table_style_name(sample_docx_with_table):
    result = read_table(sample_docx_with_table, 0)
    assert result["style"] == "Table Grid"


# ===========================================================================
# GROUP 8 — export_as_text (2 tests)
# ===========================================================================


@pytest.mark.unit
def test_export_as_text_contains_paragraph_text(tmp_docx):
    result = export_as_text(tmp_docx)
    assert "Test Document" in result["text"]
    assert "Section One" in result["text"]


@pytest.mark.unit
def test_export_as_text_truncated_false_when_under_limit(tmp_docx):
    result = export_as_text(tmp_docx)
    assert result["truncated"] is False


# ===========================================================================
# GROUP 8b — export_as_text truncation BLOCKER-1 (1 test)
# ===========================================================================


@pytest.mark.unit
def test_export_text_truncated_at_limit(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_MAX_TEXT_CHARS", "100")
    path = tmp_path / "big.docx"
    d = Document()
    for _i in range(200):
        d.add_paragraph("A" * 100)  # 200 paragraphs x 100 chars >> 100 cap
    d.save(str(path))
    result = export_as_text(str(path))
    assert result["truncated"] is True
    # text is capped at 100 chars + the truncation notice string
    assert result["text"].startswith("A" * 100)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 0
# ===========================================================================


@pytest.mark.unit
def test_capabilities_all_tool_names():
    """Exact set of 14 tool names must match specification (document_docx layer)."""
    result = capabilities()
    expected = {
        "capabilities",
        "read_document",
        "get_document_metadata",
        "list_paragraphs",
        "read_paragraph",
        "list_tables",
        "read_table",
        "add_paragraph",
        "add_heading",
        "add_page_break",
        "add_table",
        "insert_image",
        "find_replace",
        "save",
    }
    assert set(result["tools"]) == expected


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 1
# ===========================================================================


@pytest.mark.unit
def test_check_path_valid_in_allowlist(tmp_path, monkeypatch):
    """_check_path returns a resolved Path for a valid .docx within allowlist."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    p = tmp_path / "valid.docx"
    Document().save(str(p))
    resolved = _check_path(str(p))
    assert isinstance(resolved, Path)
    assert resolved.suffix == ".docx"
    assert resolved.exists()


@pytest.mark.unit
def test_check_path_traversal_attempt_raises(tmp_path, monkeypatch):
    """Directory traversal via ../../ must not escape the configured root."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    traversal = str(tmp_path / ".." / ".." / "etc" / "passwd.docx")
    with pytest.raises(ValidationError, match="not in allowlist"):
        _check_path(traversal, must_exist=False)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 2
# ===========================================================================


@pytest.mark.unit
def test_read_document_table_count(sample_docx_with_table):
    result = read_document(sample_docx_with_table)
    assert result["table_count"] == 1


@pytest.mark.unit
def test_read_document_nonexistent_raises(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    missing = str(tmp_path / "missing.docx")
    with pytest.raises(FileError):
        read_document(missing)


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 3
# ===========================================================================


@pytest.mark.unit
def test_get_metadata_returns_expected_keys(tmp_docx):
    meta = get_document_metadata(tmp_docx)
    expected_keys = {
        "path",
        "title",
        "author",
        "subject",
        "keywords",
        "category",
        "last_modified_by",
        "revision",
        "version",
        "created",
        "modified",
    }
    assert expected_keys <= set(meta.keys())


@pytest.mark.unit
def test_get_metadata_nonexistent_file_raises(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(FileError):
        get_document_metadata(str(tmp_path / "missing.docx"))


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 4
# ===========================================================================


@pytest.mark.unit
def test_list_paragraphs_entry_has_required_keys(tmp_docx):
    result = list_paragraphs(tmp_docx)
    required = {"index", "style", "text_preview", "outline_level", "run_count"}
    for entry in result:
        assert required <= set(entry.keys())
    assert result[0]["index"] == 0
    assert len(result[0]["text_preview"]) <= 120


@pytest.mark.unit
def test_list_paragraphs_empty_doc_returns_empty_list(tmp_path, monkeypatch):
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "near_empty.docx"
    d = Document()
    # Remove any default paragraphs python-docx may have added
    for p in list(d.paragraphs):
        p._element.getparent().remove(p._element)
    d.save(str(path))
    result = list_paragraphs(str(path))
    assert result == []


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 7
# ===========================================================================


@pytest.mark.unit
def test_read_table_2d_array_size_correct(sample_docx_with_table):
    result = read_table(sample_docx_with_table, 0)
    assert result["rows"] == 2
    assert result["cols"] == 3
    assert len(result["data"]) == 2
    assert len(result["data"][0]) == 3


# ===========================================================================
# W3 GAP-FILL supplements — GROUP 8
# ===========================================================================


@pytest.mark.unit
def test_export_text_includes_table_cell_text(sample_docx_with_table):
    """Table cell text must appear in export_as_text output (body-element iteration)."""
    result = export_as_text(sample_docx_with_table)
    assert "Alice" in result["text"]
    assert "Name" in result["text"]
    assert result["table_count"] == 1
