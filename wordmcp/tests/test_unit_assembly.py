"""Unit tests for word/src/wordmcp/_docx/assembly.py.

Test plan: 57 tests across 5 classes.
  - TestBulkFindReplace        (21 tests)
  - TestManageHyperlinks       (17 tests)
  - TestVerifyNoPlaceholders   (10 tests)
  - TestSignatures              ( 2 tests)
  - TestBulkFindReplaceHyperlinks ( 7 tests)

All tests are CI-safe: no live Office/COM required.
"""
from __future__ import annotations

import inspect
from pathlib import Path
from unittest.mock import MagicMock

import pytest

from docx import Document
from docx.oxml.ns import qn

from wordmcp._docx.assembly import (
    bulk_find_replace,
    manage_hyperlinks,
    verify_no_placeholders,
)
from wordmcp.document_docx import NotAllowedError, ValidationError


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ASSEMBLY_MOD = "wordmcp._docx.assembly"
_CORE_MOD = "wordmcp._docx.core"


def _patch_write_stack(monkeypatch):
    """Patch _atomic_save, _evict_doc, _audit_log, _load_doc in the facade
    (document_docx) so that no real file I/O happens after gates pass."""
    monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
    monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
    monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())

    mock_doc = MagicMock()
    # _do_find_replace returns count per token
    monkeypatch.setattr(f"{_ASSEMBLY_MOD}._do_find_replace", MagicMock(return_value=1))
    return mock_doc


def _patch_load_doc(monkeypatch, mock_doc=None):
    if mock_doc is None:
        mock_doc = MagicMock()
        # Make doc.paragraphs and doc.tables iterable
        mock_doc.paragraphs = []
        mock_doc.tables = []
    monkeypatch.setattr("wordmcp.document_docx._load_doc", MagicMock(return_value=mock_doc))
    return mock_doc


# ===========================================================================
# Class 1: TestBulkFindReplace
# ===========================================================================


class TestBulkFindReplace:
    """21 tests for bulk_find_replace."""

    # ── Gate tests (5) ──────────────────────────────────────────────────────

    def test_bulk_find_replace_blocked_without_enable_write(self, monkeypatch):
        """NotAllowedError when WORD_ENABLE_WRITE is absent."""
        monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
        with pytest.raises(NotAllowedError):
            bulk_find_replace("some/path.docx", {"{{NAME}}": "Alice"}, confirm=True)

    def test_bulk_find_replace_write_gate_fires_before_path_check(self, monkeypatch):
        """Gate fires even for a non-existent path — the write gate is first."""
        monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
        with pytest.raises(NotAllowedError):
            # Path does not exist but gate fires first
            bulk_find_replace("/nonexistent/path/file.docx", {"{{X}}": "y"}, confirm=True)

    def test_bulk_find_replace_write_gate_fires_before_validation(self, monkeypatch):
        """Gate fires before empty-dict ValidationError is raised."""
        monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
        with pytest.raises(NotAllowedError):
            # Empty replacements would raise ValidationError, but gate fires first
            bulk_find_replace("some/path.docx", {}, confirm=True)

    def test_bulk_find_replace_blocked_without_confirm(self, monkeypatch, tmp_docx):
        """ValidationError when confirm=False (after WORD_ENABLE_WRITE=true)."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            bulk_find_replace(tmp_docx, {"{{NAME}}": "Alice"}, confirm=False)

    def test_bulk_find_replace_confirm_default_is_false(self, monkeypatch, tmp_docx):
        """Calling without confirm kwarg raises ValidationError (default=False)."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            bulk_find_replace(tmp_docx, {"{{NAME}}": "Alice"})

    # ── Path gate (2) ───────────────────────────────────────────────────────

    def test_bulk_find_replace_blocked_for_path_outside_allowlist(self, monkeypatch, tmp_path):
        """ValidationError for path outside WORD_ALLOWLIST_ROOTS."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path / "allowed"))
        outside_path = tmp_path / "outside" / "doc.docx"
        with pytest.raises((ValidationError, NotAllowedError)):
            bulk_find_replace(str(outside_path), {"{{X}}": "y"}, confirm=True)

    def test_bulk_find_replace_blocked_for_non_docx_extension(self, monkeypatch, tmp_path):
        """ValidationError for .pdf extension."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        pdf_path = tmp_path / "doc.pdf"
        pdf_path.touch()
        with pytest.raises(ValidationError):
            bulk_find_replace(str(pdf_path), {"{{X}}": "y"}, confirm=True)

    # ── Validation (5) ──────────────────────────────────────────────────────

    def test_bulk_find_replace_empty_replacements_raises(self, monkeypatch, tmp_docx):
        """ValidationError for empty replacements dict."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError, match="empty"):
            bulk_find_replace(tmp_docx, {}, confirm=True)

    def test_bulk_find_replace_key_without_braces_raises(self, monkeypatch, tmp_docx):
        """'NAME' (no braces) raises ValidationError."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            bulk_find_replace(tmp_docx, {"NAME": "Alice"}, confirm=True)

    def test_bulk_find_replace_key_with_single_brace_raises(self, monkeypatch, tmp_docx):
        """'{NAME}' (single braces) raises ValidationError."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            bulk_find_replace(tmp_docx, {"{NAME}": "Alice"}, confirm=True)

    def test_bulk_find_replace_non_string_value_raises(self, monkeypatch, tmp_docx):
        """value=123 (non-str) raises ValidationError."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            bulk_find_replace(tmp_docx, {"{{NAME}}": 123}, confirm=True)  # type: ignore[arg-type]

    def test_bulk_find_replace_oversized_dict_raises(self, monkeypatch, tmp_docx):
        """501 entries raises ValidationError (limit is 500)."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        big = {f"{{{{TOKEN_{i}}}}}": f"v{i}" for i in range(501)}
        with pytest.raises(ValidationError):
            bulk_find_replace(tmp_docx, big, confirm=True)

    # ── Happy path (4) ──────────────────────────────────────────────────────

    def test_bulk_find_replace_returns_expected_keys(self, monkeypatch, write_enabled_docx):
        """Return dict has keys: replaced, results, tokens_not_found, path."""
        _patch_write_stack(monkeypatch)
        _patch_load_doc(monkeypatch)
        result = bulk_find_replace(
            write_enabled_docx, {"{{NAME}}": "Alice"}, confirm=True
        )
        assert set(result.keys()) >= {"replaced", "results", "tokens_not_found", "path"}

    def test_bulk_find_replace_replaced_count_is_sum_of_token_counts(
        self, monkeypatch, write_enabled_docx
    ):
        """result['replaced'] equals sum of per-token counts."""
        monkeypatch.setattr(
            f"{_ASSEMBLY_MOD}._do_find_replace", MagicMock(return_value=3)
        )
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())
        _patch_load_doc(monkeypatch)

        result = bulk_find_replace(
            write_enabled_docx,
            {"{{A}}": "a", "{{B}}": "b"},
            confirm=True,
        )
        # 3 per token × 2 tokens = 6
        assert result["replaced"] == 6

    def test_bulk_find_replace_tokens_not_found_populated_when_count_zero(
        self, monkeypatch, write_enabled_docx
    ):
        """Token with count 0 appears in tokens_not_found."""
        call_count = {"n": 0}

        def _side_effect(doc, token, value):
            # First call returns 1, second returns 0
            call_count["n"] += 1
            return 1 if call_count["n"] == 1 else 0

        monkeypatch.setattr(f"{_ASSEMBLY_MOD}._do_find_replace", _side_effect)
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())
        _patch_load_doc(monkeypatch)

        result = bulk_find_replace(
            write_enabled_docx,
            {"{{FOUND}}": "yes", "{{MISSING}}": "no"},
            confirm=True,
        )
        assert "{{MISSING}}" in result["tokens_not_found"]
        assert "{{FOUND}}" not in result["tokens_not_found"]

    def test_bulk_find_replace_path_in_result_matches_resolved_path(
        self, monkeypatch, write_enabled_docx
    ):
        """result['path'] equals str(resolved path)."""
        _patch_write_stack(monkeypatch)
        _patch_load_doc(monkeypatch)
        result = bulk_find_replace(
            write_enabled_docx, {"{{NAME}}": "Alice"}, confirm=True
        )
        assert result["path"] == str(Path(write_enabled_docx).resolve())

    # ── Side effects (3) ────────────────────────────────────────────────────

    def test_bulk_find_replace_calls_atomic_save_exactly_once(
        self, monkeypatch, write_enabled_docx
    ):
        """_atomic_save is called exactly once per invocation."""
        _patch_write_stack(monkeypatch)
        mock_save = MagicMock()
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", mock_save)
        _patch_load_doc(monkeypatch)

        bulk_find_replace(write_enabled_docx, {"{{X}}": "y"}, confirm=True)
        assert mock_save.call_count == 1

    def test_bulk_find_replace_calls_audit_log_with_op_name(
        self, monkeypatch, write_enabled_docx
    ):
        """_audit_log is called with op='bulk_find_replace'."""
        _patch_write_stack(monkeypatch)
        mock_audit = MagicMock()
        monkeypatch.setattr("wordmcp.document_docx._audit_log", mock_audit)
        _patch_load_doc(monkeypatch)

        bulk_find_replace(write_enabled_docx, {"{{X}}": "y"}, confirm=True)
        assert mock_audit.call_count == 1
        assert mock_audit.call_args[0][0] == "bulk_find_replace"

    def test_bulk_find_replace_calls_evict_doc(
        self, monkeypatch, write_enabled_docx
    ):
        """_evict_doc is called exactly once."""
        _patch_write_stack(monkeypatch)
        mock_evict = MagicMock()
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", mock_evict)
        _patch_load_doc(monkeypatch)

        bulk_find_replace(write_enabled_docx, {"{{X}}": "y"}, confirm=True)
        assert mock_evict.call_count == 1

    # ── Boundary (2) ────────────────────────────────────────────────────────

    def test_bulk_find_replace_single_valid_token_succeeds(
        self, monkeypatch, write_enabled_docx
    ):
        """Single valid token {"{{NAME}}": "Alice"} succeeds without raising."""
        _patch_write_stack(monkeypatch)
        _patch_load_doc(monkeypatch)
        result = bulk_find_replace(
            write_enabled_docx, {"{{NAME}}": "Alice"}, confirm=True
        )
        assert isinstance(result, dict)

    def test_bulk_find_replace_exactly_500_entries_is_valid(
        self, monkeypatch, write_enabled_docx
    ):
        """Exactly 500 entries does NOT raise (boundary of allowed limit)."""
        _patch_write_stack(monkeypatch)
        _patch_load_doc(monkeypatch)
        big = {f"{{{{TOKEN_{i}}}}}": f"v{i}" for i in range(500)}
        result = bulk_find_replace(write_enabled_docx, big, confirm=True)
        assert isinstance(result, dict)


# ===========================================================================
# Class 2: TestManageHyperlinks
# ===========================================================================


class TestManageHyperlinks:
    """17 tests for manage_hyperlinks."""

    # ── Write gate (2) ──────────────────────────────────────────────────────

    def test_manage_hyperlinks_blocked_without_enable_write(self, monkeypatch):
        """NotAllowedError when WORD_ENABLE_WRITE is absent."""
        monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
        with pytest.raises(NotAllowedError):
            manage_hyperlinks(
                "some/path.docx",
                {"Click here": {"url": "https://example.com"}},
                confirm=True,
            )

    def test_manage_hyperlinks_write_gate_fires_before_path_check(self, monkeypatch):
        """Write gate fires even for non-existent path."""
        monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
        with pytest.raises(NotAllowedError):
            manage_hyperlinks(
                "/nonexistent/file.docx",
                {"Click": {"url": "https://example.com"}},
                confirm=True,
            )

    # ── Confirm gate (2) ────────────────────────────────────────────────────

    def test_manage_hyperlinks_blocked_without_confirm(self, monkeypatch, tmp_docx):
        """ValidationError when confirm=False."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            manage_hyperlinks(
                tmp_docx,
                {"Click": {"url": "https://example.com"}},
                confirm=False,
            )

    def test_manage_hyperlinks_confirm_default_is_false(self, monkeypatch, tmp_docx):
        """Calling without confirm kwarg raises ValidationError."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        with pytest.raises(ValidationError):
            manage_hyperlinks(tmp_docx, {"Click": {"url": "https://example.com"}})

    # ── Path gate (1) ───────────────────────────────────────────────────────

    def test_manage_hyperlinks_blocked_for_path_outside_allowlist(
        self, monkeypatch, tmp_path
    ):
        """ValidationError for path outside WORD_ALLOWLIST_ROOTS."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path / "allowed"))
        outside = tmp_path / "outside" / "doc.docx"
        with pytest.raises((ValidationError, NotAllowedError)):
            manage_hyperlinks(
                str(outside),
                {"Click": {"url": "https://example.com"}},
                confirm=True,
            )

    # ── URL scheme validation (6) ────────────────────────────────────────────

    def test_manage_hyperlinks_rejects_javascript_scheme(
        self, monkeypatch, write_enabled_docx
    ):
        """ValidationError for javascript: scheme (B1)."""
        with pytest.raises(ValidationError):
            manage_hyperlinks(
                write_enabled_docx,
                {"Click": {"url": "javascript:alert(1)"}},
                confirm=True,
            )

    def test_manage_hyperlinks_rejects_file_scheme(
        self, monkeypatch, write_enabled_docx
    ):
        """ValidationError for file:// scheme (B1)."""
        with pytest.raises(ValidationError):
            manage_hyperlinks(
                write_enabled_docx,
                {"Click": {"url": "file:///etc/passwd"}},
                confirm=True,
            )

    def test_manage_hyperlinks_rejects_data_scheme(
        self, monkeypatch, write_enabled_docx
    ):
        """ValidationError for data: scheme (B1)."""
        with pytest.raises(ValidationError):
            manage_hyperlinks(
                write_enabled_docx,
                {"Click": {"url": "data:text/html,<h1>x</h1>"}},
                confirm=True,
            )

    def test_manage_hyperlinks_rejects_unc_path_as_url(
        self, monkeypatch, write_enabled_docx
    ):
        r"""ValidationError for UNC path (\\server\share) — no allowed scheme (B1)."""
        with pytest.raises(ValidationError):
            manage_hyperlinks(
                write_enabled_docx,
                {"Click": {"url": "\\\\server\\share\\file.docx"}},
                confirm=True,
            )

    def test_manage_hyperlinks_accepts_https_scheme(
        self, monkeypatch, write_enabled_docx
    ):
        """https:// URL passes validation without raising (B1 complement)."""
        mock_doc = MagicMock()
        mock_doc.element.body.iter.return_value = iter([])
        monkeypatch.setattr("wordmcp.document_docx._load_doc", MagicMock(return_value=mock_doc))
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())

        result = manage_hyperlinks(
            write_enabled_docx,
            {"Click": {"url": "https://example.com"}},
            confirm=True,
        )
        assert isinstance(result, dict)

    def test_manage_hyperlinks_accepts_mailto_scheme(
        self, monkeypatch, write_enabled_docx
    ):
        """mailto: URL passes validation without raising (B1 complement)."""
        mock_doc = MagicMock()
        mock_doc.element.body.iter.return_value = iter([])
        monkeypatch.setattr("wordmcp.document_docx._load_doc", MagicMock(return_value=mock_doc))
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())

        result = manage_hyperlinks(
            write_enabled_docx,
            {"Contact": {"url": "mailto:user@example.com"}},
            confirm=True,
        )
        assert isinstance(result, dict)

    # ── B6 pre-flight atomicity (1) ──────────────────────────────────────────

    def test_manage_hyperlinks_preflight_rejects_partial_bad_dict(
        self, monkeypatch, write_enabled_docx
    ):
        """B6: bad entry (javascript: scheme) causes ValidationError BEFORE _atomic_save."""
        mock_save = MagicMock()
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", mock_save)
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())

        with pytest.raises(ValidationError):
            manage_hyperlinks(
                write_enabled_docx,
                {
                    "Good link": {"url": "https://good.example.com"},
                    "Bad link": {"url": "javascript:evil()"},
                },
                confirm=True,
            )
        # _atomic_save must never have been called
        assert mock_save.call_count == 0

    # ── Input structure (2) ──────────────────────────────────────────────────

    def test_manage_hyperlinks_entry_missing_url_key_raises(
        self, monkeypatch, write_enabled_docx
    ):
        """ValidationError when entry dict is missing 'url' key."""
        with pytest.raises(ValidationError, match="url"):
            manage_hyperlinks(
                write_enabled_docx,
                {"Click": {"label": "no url here"}},
                confirm=True,
            )

    def test_manage_hyperlinks_url_not_string_raises(
        self, monkeypatch, write_enabled_docx
    ):
        """ValidationError when url value is not a string."""
        with pytest.raises(ValidationError):
            manage_hyperlinks(
                write_enabled_docx,
                {"Click": {"url": 12345}},
                confirm=True,
            )

    # ── Happy path + return shape (2) ────────────────────────────────────────

    def test_manage_hyperlinks_returns_expected_keys(
        self, monkeypatch, write_enabled_docx
    ):
        """Return dict has keys: updated, results, path."""
        mock_doc = MagicMock()
        mock_doc.element.body.iter.return_value = iter([])
        monkeypatch.setattr("wordmcp.document_docx._load_doc", MagicMock(return_value=mock_doc))
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())

        result = manage_hyperlinks(
            write_enabled_docx,
            {"Click": {"url": "https://example.com"}},
            confirm=True,
        )
        assert set(result.keys()) >= {"updated", "results", "path"}

    def test_manage_hyperlinks_not_found_when_display_text_absent(
        self, monkeypatch, write_enabled_docx
    ):
        """status=='not_found' and updated==0 when display text is absent in doc."""
        mock_doc = MagicMock()
        # No hyperlink elements in the document body
        mock_doc.element.body.iter.return_value = iter([])
        monkeypatch.setattr("wordmcp.document_docx._load_doc", MagicMock(return_value=mock_doc))
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._audit_log", MagicMock())

        result = manage_hyperlinks(
            write_enabled_docx,
            {"NonExistentText": {"url": "https://example.com"}},
            confirm=True,
        )
        assert result["updated"] == 0
        assert result["results"]["NonExistentText"]["status"] == "not_found"

    # ── Audit (1) ────────────────────────────────────────────────────────────

    def test_manage_hyperlinks_calls_audit_log_with_op_name(
        self, monkeypatch, write_enabled_docx
    ):
        """_audit_log is called with op='manage_hyperlinks'."""
        mock_doc = MagicMock()
        mock_doc.element.body.iter.return_value = iter([])
        monkeypatch.setattr("wordmcp.document_docx._load_doc", MagicMock(return_value=mock_doc))
        monkeypatch.setattr("wordmcp.document_docx._atomic_save", MagicMock())
        monkeypatch.setattr("wordmcp.document_docx._evict_doc", MagicMock())
        mock_audit = MagicMock()
        monkeypatch.setattr("wordmcp.document_docx._audit_log", mock_audit)

        manage_hyperlinks(
            write_enabled_docx,
            {"Click": {"url": "https://example.com"}},
            confirm=True,
        )
        assert mock_audit.call_count == 1
        assert mock_audit.call_args[0][0] == "manage_hyperlinks"


# ===========================================================================
# Class 3: TestVerifyNoPlaceholders
# ===========================================================================


class TestVerifyNoPlaceholders:
    """10 tests for verify_no_placeholders."""

    # ── Path gate (2) ───────────────────────────────────────────────────────

    def test_verify_no_placeholders_blocked_for_path_outside_allowlist(
        self, monkeypatch, tmp_path
    ):
        """ValidationError/NotAllowedError for path outside WORD_ALLOWLIST_ROOTS."""
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path / "allowed"))
        outside = tmp_path / "outside" / "doc.docx"
        with pytest.raises((ValidationError, NotAllowedError)):
            verify_no_placeholders(str(outside))

    def test_verify_no_placeholders_does_not_require_enable_write(
        self, monkeypatch, assembly_docx_clean
    ):
        """No NotAllowedError raised even when WORD_ENABLE_WRITE is absent (read-only op)."""
        monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
        # Should succeed without write env var
        result = verify_no_placeholders(assembly_docx_clean)
        assert result["status"] == "pass"

    # ── Pass status (2) ─────────────────────────────────────────────────────

    def test_verify_no_placeholders_returns_pass_when_clean(
        self, assembly_docx_clean
    ):
        """Returns status='pass' for a document with no placeholders."""
        result = verify_no_placeholders(assembly_docx_clean)
        assert result["status"] == "pass"
        assert result["residuals"] == []

    def test_verify_no_placeholders_returns_expected_keys(
        self, assembly_docx_clean
    ):
        """Return dict has: status, residuals, scanned_paragraphs, scanned_cells, path."""
        result = verify_no_placeholders(assembly_docx_clean)
        assert set(result.keys()) >= {
            "status",
            "residuals",
            "scanned_paragraphs",
            "scanned_cells",
            "path",
        }

    # ── Fail status (3) ─────────────────────────────────────────────────────

    def test_verify_no_placeholders_returns_fail_when_placeholder_in_paragraph(
        self, assembly_docx_with_placeholders
    ):
        """Returns status='fail' and residuals non-empty for doc with body placeholders."""
        result = verify_no_placeholders(assembly_docx_with_placeholders)
        assert result["status"] == "fail"
        assert len(result["residuals"]) > 0

    def test_verify_no_placeholders_returns_fail_when_placeholder_in_table_cell(
        self, assembly_docx_with_placeholders
    ):
        """Returns status='fail' when placeholder exists in a table cell."""
        result = verify_no_placeholders(assembly_docx_with_placeholders)
        # The fixture includes {{INVOICE_NUM}} in a table cell
        assert result["status"] == "fail"
        assert "{{INVOICE_NUM}}" in result["residuals"]

    def test_verify_no_placeholders_scanned_counts_are_nonnegative_integers(
        self, assembly_docx_with_placeholders
    ):
        """scanned_paragraphs and scanned_cells are non-negative integers."""
        result = verify_no_placeholders(assembly_docx_with_placeholders)
        assert isinstance(result["scanned_paragraphs"], int)
        assert isinstance(result["scanned_cells"], int)
        assert result["scanned_paragraphs"] >= 0
        assert result["scanned_cells"] >= 0

    # ── Ignore filter (2) ────────────────────────────────────────────────────

    def test_verify_no_placeholders_ignores_listed_tokens(
        self, assembly_docx_with_placeholders
    ):
        """Tokens listed in ignore are excluded from residuals."""
        result = verify_no_placeholders(
            assembly_docx_with_placeholders,
            ignore=["{{FIRST_NAME}}", "{{LAST_NAME}}", "{{DATE}}", "{{INVOICE_NUM}}"],
        )
        # All placeholders in the doc are now in the ignore list
        assert result["status"] == "pass"
        assert result["residuals"] == []

    def test_verify_no_placeholders_ignore_none_does_not_crash(
        self, assembly_docx_clean
    ):
        """Passing ignore=None does not raise."""
        result = verify_no_placeholders(assembly_docx_clean, ignore=None)
        assert isinstance(result, dict)

    # ── Return shape (1) ─────────────────────────────────────────────────────

    def test_verify_no_placeholders_path_in_result_is_string(
        self, assembly_docx_clean
    ):
        """result['path'] is a string."""
        result = verify_no_placeholders(assembly_docx_clean)
        assert isinstance(result["path"], str)


# ===========================================================================
# Class 4: TestSignatures
# ===========================================================================


class TestSignatures:
    """2 tests verifying function signatures match the contract (B3)."""

    def test_bulk_find_replace_signature_has_correct_parameter_names(self):
        """bulk_find_replace must have params: path, replacements, confirm; confirm default=False."""
        sig = inspect.signature(bulk_find_replace)
        params = list(sig.parameters.keys())
        assert "path" in params
        assert "replacements" in params
        assert "confirm" in params
        assert sig.parameters["confirm"].default is False

    def test_manage_hyperlinks_signature_has_correct_parameter_names(self):
        """manage_hyperlinks must have params: path, hyperlinks, confirm; confirm default=False."""
        sig = inspect.signature(manage_hyperlinks)
        params = list(sig.parameters.keys())
        assert "path" in params
        assert "hyperlinks" in params
        assert "confirm" in params
        assert sig.parameters["confirm"].default is False

    def test_bulk_find_replace_has_hyperlink_fields_parameter(self):
        """bulk_find_replace must accept hyperlink_fields kwarg with default None."""
        sig = inspect.signature(bulk_find_replace)
        assert "hyperlink_fields" in sig.parameters
        assert sig.parameters["hyperlink_fields"].default is None


# ===========================================================================
# Class 5: TestBulkFindReplaceHyperlinks
# ===========================================================================


class TestBulkFindReplaceHyperlinks:
    """7 tests for the hyperlink_fields parameter of bulk_find_replace."""

    # ── Validation (3) ──────────────────────────────────────────────────────

    def test_hyperlink_fields_overlap_with_replacements_raises(
        self, monkeypatch, write_enabled_docx
    ):
        """Token in both replacements and hyperlink_fields raises ValidationError."""
        with pytest.raises(ValidationError, match="both"):
            bulk_find_replace(
                write_enabled_docx,
                {"{{URL}}": "plain text"},
                confirm=True,
                hyperlink_fields={"{{URL}}": "https://example.com"},
            )

    def test_hyperlink_fields_invalid_token_pattern_raises(
        self, monkeypatch, write_enabled_docx
    ):
        """Token without {{}} pattern in hyperlink_fields raises ValidationError."""
        with pytest.raises(ValidationError):
            bulk_find_replace(
                write_enabled_docx,
                {"{{NAME}}": "Alice"},
                confirm=True,
                hyperlink_fields={"BAD_TOKEN": "https://example.com"},
            )

    def test_hyperlink_fields_invalid_url_scheme_raises(
        self, monkeypatch, write_enabled_docx
    ):
        """javascript: URL in hyperlink_fields raises ValidationError."""
        with pytest.raises(ValidationError):
            bulk_find_replace(
                write_enabled_docx,
                {"{{NAME}}": "Alice"},
                confirm=True,
                hyperlink_fields={"{{LINK}}": "javascript:alert(1)"},
            )

    # ── Happy path — real docx I/O (3) ──────────────────────────────────────

    def test_hyperlink_fields_creates_hyperlink_element(
        self, tmp_path, monkeypatch
    ):
        """hyperlink_fields replaces {{WEBSITE_URL}} with a w:hyperlink element."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

        doc_path = tmp_path / "hyperlink_test.docx"
        doc = Document()
        doc.add_paragraph("Visit {{WEBSITE_URL}} for more info.")
        doc.save(str(doc_path))

        result = bulk_find_replace(
            str(doc_path),
            {"{{PLACEHOLDER}}": "unused"},  # non-overlapping plain replacement
            confirm=True,
            hyperlink_fields={"{{WEBSITE_URL}}": "https://example.com"},
        )

        assert isinstance(result, dict)
        assert "{{WEBSITE_URL}}" in result["results"]

        # Reload and verify w:hyperlink element is present
        saved = Document(str(doc_path))
        hyperlinks = list(saved.element.body.iter(qn("w:hyperlink")))
        assert len(hyperlinks) >= 1, "Expected at least one w:hyperlink element"

    def test_hyperlink_fields_linkedin_url_regression(
        self, tmp_path, monkeypatch
    ):
        """Regression: {{LINKEDIN_URL}} in hyperlink_fields creates a clickable link."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

        doc_path = tmp_path / "linkedin_test.docx"
        doc = Document()
        doc.add_paragraph("LinkedIn: {{LINKEDIN_URL}}")
        doc.save(str(doc_path))

        result = bulk_find_replace(
            str(doc_path),
            {"{{NAME}}": "Alice"},
            confirm=True,
            hyperlink_fields={"{{LINKEDIN_URL}}": "https://linkedin.com/in/alice"},
        )

        assert result["results"]["{{LINKEDIN_URL}}"] == 1

        saved = Document(str(doc_path))
        hyperlinks = list(saved.element.body.iter(qn("w:hyperlink")))
        assert len(hyperlinks) >= 1, "Expected w:hyperlink element for LINKEDIN_URL"

    def test_hyperlink_fields_url_stored_in_relationship(
        self, tmp_path, monkeypatch
    ):
        """The relationship target URL matches the value in hyperlink_fields."""
        monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))

        target_url = "https://example.com"
        doc_path = tmp_path / "rel_test.docx"
        doc = Document()
        doc.add_paragraph("See {{SITE_URL}} here.")
        doc.save(str(doc_path))

        bulk_find_replace(
            str(doc_path),
            {"{{NAME}}": "Bob"},
            confirm=True,
            hyperlink_fields={"{{SITE_URL}}": target_url},
        )

        saved = Document(str(doc_path))
        # Collect all hyperlink relationship targets in the main document part
        rel_targets = [
            rel.target_ref
            for rel in saved.part.rels.values()
            if "hyperlink" in rel.reltype
        ]
        assert target_url in rel_targets, (
            f"Expected {target_url!r} in part relationships, got {rel_targets}"
        )

    # ── None / omitted (1) ──────────────────────────────────────────────────

    def test_hyperlink_fields_none_is_backward_compatible(
        self, monkeypatch, write_enabled_docx
    ):
        """Omitting hyperlink_fields (None) does not change existing behaviour."""
        _patch_write_stack(monkeypatch)
        _patch_load_doc(monkeypatch)

        result = bulk_find_replace(
            write_enabled_docx,
            {"{{NAME}}": "Alice"},
            confirm=True,
            # hyperlink_fields omitted — defaults to None
        )
        assert set(result.keys()) >= {"replaced", "results", "tokens_not_found", "path"}
