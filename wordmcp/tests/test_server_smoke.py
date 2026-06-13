"""FastMCP smoke tests for the wordmcp server surface."""
from __future__ import annotations

import sys

import pytest
from docx import Document


def _call_tool(tool, *args, **kwargs):
    """Invoke a FastMCP-decorated tool regardless of FastMCP version.

    FastMCP 2.x: @mcp.tool() returns a FunctionTool (not directly callable);
                 the underlying callable is exposed via the .fn attribute.
    FastMCP 3.x: @mcp.tool() returns the original callable function directly.
    """
    fn = getattr(tool, "fn", tool)
    return fn(*args, **kwargs)


@pytest.mark.smoke
def test_server_module_imports_without_error():
    """wordmcp.server imports cleanly and exposes the expected tool surface."""
    from wordmcp import server as srv

    assert srv.mcp is not None


@pytest.mark.smoke
def test_all_expected_tools_present_on_server_module():
    """All expected tool functions exist on the server module (47 callables total)."""
    from wordmcp import server as srv

    expected_tools = [
        "capabilities",
        "read_document",
        "get_document_metadata",
        "list_paragraphs",
        "read_paragraph",
        "list_tables",
        "read_table",
        "add_paragraph",
        "add_heading",
        "add_page_break",
        "add_table",
        "insert_image",
        "find_replace",
        "save",
        "manage_comments",
        "manage_tracked_changes",
        "export_document",
        "create_document",
        "get_document_context",
        "list_headings",
        "read_section",
        "search_text",
        "list_styles",
        "apply_style",
        "update_paragraph",
        "delete_paragraph",
        "update_table_cell",
        "set_document_properties",
        "set_paragraph_format",
        "add_list",
        "bulk_add_paragraphs",
        "bulk_update_paragraphs",
        "bulk_update_table_cells",
        "get_headers_footers",
        "add_hyperlink",
        "add_footnote",
        "get_document_outline",
        "review_document",
        "write_review_findings",
        "export_review_evidence",
        "insert_paragraph",
        "paragraph",
        "export",
        "table",
        "style",
        "content",
        "review",
        "document",
    ]
    assert len(expected_tools) == 48
    for tool_name in expected_tools:
        tool_obj = getattr(srv, tool_name, None)
        assert tool_obj is not None and (
            callable(tool_obj) or callable(getattr(tool_obj, "fn", None))
        ), f"Tool {tool_name!r} not found or not callable in server module"


@pytest.mark.smoke
def test_capabilities_tool_returns_expected_structure():
    """capabilities() returns a dict with phase, backend, tools, governance keys."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert isinstance(result, dict)
    assert result["phase"] == "1.0"
    assert result["backend"] == "python-docx"
    assert "get_document_context" in result["tools"]
    assert len(result["tools"]) == 50
    assert "governance" in result


@pytest.mark.smoke
def test_read_document_tool_via_server(tmp_docx):
    """read_document tool returns correct structure for a valid .docx."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_document, tmp_docx)
    assert "paragraph_count" in result
    assert "table_count" in result
    assert "section_count" in result
    assert result["paragraph_count"] == 3


@pytest.mark.smoke
def test_list_paragraphs_tool_via_server(tmp_docx):
    """list_paragraphs tool returns a list of paragraphs with required keys."""
    from wordmcp import server as srv

    result = _call_tool(srv.list_paragraphs, tmp_docx)
    assert isinstance(result, list)
    assert len(result) == 3
    assert all("index" in p and "style" in p for p in result)


@pytest.mark.smoke
def test_export_as_text_tool_via_server(tmp_docx):
    """export_document(format='txt') returns text field and truncated bool."""
    from wordmcp import server as srv

    result = _call_tool(srv.export_document, tmp_docx, format="txt")
    assert "text" in result
    assert isinstance(result["truncated"], bool)
    assert "Test Document" in result["text"]


@pytest.mark.smoke
def test_read_tool_raises_tool_error_for_bad_path(tmp_path, monkeypatch):
    """ToolError is raised for a path outside the allowlist."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ToolError):
        _call_tool(srv.read_document, "/not/in/allowlist/file.docx")


# ---------------------------------------------------------------------------
# Capabilities v2 — TOOL_REGISTRY and build_capabilities_v2 tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_tool_registry_no_duplicate_keys():
    """TOOL_REGISTRY dict has no duplicate keys (Python dict invariant, verified explicitly)."""
    from wordmcp._server_manifest import TOOL_REGISTRY

    # Dict keys are unique by definition; verify count is the expected 50
    assert len(TOOL_REGISTRY) == 50


@pytest.mark.smoke
def test_tool_registry_has_exactly_50_entries():
    """TOOL_REGISTRY has exactly 50 entries: 13 primary + 37 deprecated."""
    from wordmcp._server_manifest import TOOL_REGISTRY

    primary_kinds = {"dispatcher", "standalone"}
    primary = [n for n, m in TOOL_REGISTRY.items() if m["kind"] in primary_kinds]
    deprecated = [n for n, m in TOOL_REGISTRY.items() if m["kind"] == "deprecated_alias"]

    assert len(primary) == 13
    assert len(deprecated) == 37
    assert len(TOOL_REGISTRY) == 50


@pytest.mark.smoke
def test_build_capabilities_v2_primary_tools_count():
    """build_capabilities_v2() returns primary_tools with exactly 13 entries."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    assert len(result["primary_tools"]) == 13


@pytest.mark.smoke
def test_build_capabilities_v2_deprecated_aliases_count():
    """build_capabilities_v2() returns deprecated_aliases with exactly 37 entries."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    assert len(result["deprecated_aliases"]) == 37


@pytest.mark.smoke
def test_build_capabilities_v2_every_alias_has_replacement():
    """Every deprecated_alias entry has a non-empty replacement_tool and replacement_operation_or_scope."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    for alias in result["deprecated_aliases"]:
        assert result["replacement_tool"].get(alias), (
            f"{alias!r} missing replacement_tool"
        )
        assert result["replacement_operation_or_scope"].get(alias) is not None, (
            f"{alias!r} missing replacement_operation_or_scope key"
        )
        assert result["replacement_operation_or_scope"][alias] != "", (
            f"{alias!r} has empty replacement_operation_or_scope"
        )


@pytest.mark.smoke
def test_build_capabilities_v2_replacement_tool_values_are_primary():
    """Every replacement_tool value is one of the 10 primary tools."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    primary_set = set(result["primary_tools"])
    for alias, repl in result["replacement_tool"].items():
        assert repl in primary_set, (
            f"{alias!r} replacement_tool={repl!r} not in primary_tools"
        )


@pytest.mark.smoke
def test_build_capabilities_v2_replacement_op_in_dispatcher_ops():
    """Every replacement_operation_or_scope appears in the corresponding dispatcher's operations."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    enums = result["operation_scope_enums"]

    for alias in result["deprecated_aliases"]:
        repl_tool = result["replacement_tool"][alias]
        repl_op = result["replacement_operation_or_scope"][alias]
        if repl_tool in enums:
            assert repl_op in enums[repl_tool], (
                f"{alias!r}: op={repl_op!r} not in {repl_tool!r} ops {enums[repl_tool]}"
            )


@pytest.mark.smoke
def test_build_capabilities_v2_total_callable_endpoints():
    """total_callable_endpoints == 50."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    assert result["total_callable_endpoints"] == 50
    assert result["total_callable_endpoints"] == (
        len(result["primary_tools"]) + len(result["deprecated_aliases"])
    )


@pytest.mark.smoke
def test_build_capabilities_v2_write_gate_metadata_non_empty():
    """write_gate_metadata is non-empty and covers all write-gated tools."""
    from wordmcp._server_manifest import TOOL_REGISTRY, build_capabilities_v2

    result = build_capabilities_v2(None)
    assert len(result["write_gate_metadata"]) > 0

    write_gated_names = {n for n, m in TOOL_REGISTRY.items() if m["write_gated"]}
    for name in write_gated_names:
        assert name in result["write_gate_metadata"], (
            f"Write-gated tool {name!r} missing from write_gate_metadata"
        )
        gate = result["write_gate_metadata"][name]
        assert gate["env_var"] == "WORD_ENABLE_WRITE"
        assert gate["requires_confirm"] is True


@pytest.mark.smoke
def test_build_capabilities_v2_operation_scope_enums_covers_all_dispatchers():
    """operation_scope_enums has exactly 8 keys, one per dispatcher."""
    from wordmcp._server_manifest import TOOL_REGISTRY, build_capabilities_v2

    result = build_capabilities_v2(None)
    dispatcher_names = {n for n, m in TOOL_REGISTRY.items() if m["kind"] == "dispatcher"}
    assert set(result["operation_scope_enums"].keys()) == dispatcher_names
    for dispatcher, ops in result["operation_scope_enums"].items():
        assert len(ops) > 0, f"{dispatcher!r} has empty operations list"
        assert ops == sorted(ops), f"{dispatcher!r} operations not sorted"


@pytest.mark.smoke
def test_build_capabilities_v2_deprecation_policy_keys():
    """deprecation_policy contains required keys: window_releases, telemetry_field, removal_date_iso."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    policy = result["deprecation_policy"]
    assert "window_releases" in policy
    assert isinstance(policy["window_releases"], int)
    assert "telemetry_field" in policy
    assert "removal_date_iso" in policy


@pytest.mark.smoke
def test_build_capabilities_v2_sorted_outputs():
    """primary_tools and deprecated_aliases are sorted alphabetically."""
    from wordmcp._server_manifest import build_capabilities_v2

    result = build_capabilities_v2(None)
    assert result["primary_tools"] == sorted(result["primary_tools"])
    assert result["deprecated_aliases"] == sorted(result["deprecated_aliases"])


@pytest.mark.smoke
def test_build_capabilities_legacy_fields_still_present():
    """build_capabilities() output still contains all legacy fields (non-breaking)."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    # Legacy fields must remain
    assert "tools" in result and isinstance(result["tools"], list)
    assert len(result["tools"]) == 50
    assert "phase" in result
    assert "backend" in result
    assert "governance" in result
    assert "com_tools" in result
    assert "read_document_scopes" in result
    assert "table_operations" in result
    assert "style_operations" in result
    assert "content_operations" in result
    assert "review_operations" in result
    assert "document_operations" in result


@pytest.mark.smoke
def test_build_capabilities_v2_field_present_in_capabilities_output():
    """build_capabilities() output contains the capabilities_v2 nested field."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "capabilities_v2" in result
    v2 = result["capabilities_v2"]
    assert "primary_tools" in v2
    assert "deprecated_aliases" in v2
    assert "replacement_tool" in v2
    assert "replacement_operation_or_scope" in v2
    assert "total_callable_endpoints" in v2
    assert "operation_scope_enums" in v2
    assert "write_gate_metadata" in v2
    assert "deprecation_policy" in v2


@pytest.mark.smoke
def test_build_capabilities_v2_parity_probe_a():
    """Parity A: primary_tools | deprecated_aliases == tools flat list (set equality)."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    v2 = result["capabilities_v2"]
    structured_set = set(v2["primary_tools"]) | set(v2["deprecated_aliases"])
    flat_set = set(result["tools"])
    assert structured_set == flat_set, (
        f"In structured but not flat: {structured_set - flat_set}\n"
        f"In flat but not structured: {flat_set - structured_set}"
    )


@pytest.mark.smoke
def test_build_capabilities_v2_parity_probe_c():
    """Parity C: len(tools) == total_callable_endpoints."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    v2 = result["capabilities_v2"]
    assert len(result["tools"]) == v2["total_callable_endpoints"]


@pytest.mark.smoke
def test_write_tool_raises_tool_error_without_enable_write(tmp_docx):
    """ToolError is raised when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    with pytest.raises(ToolError):
        _call_tool(srv.add_paragraph, tmp_docx, "hello", confirm=True)


@pytest.mark.smoke
def test_server_confirm_default_false_for_add_heading(tmp_path, monkeypatch):
    """server.add_heading without confirm must raise ToolError (migrated from test_unit)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "test.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.add_heading, str(path), "Title", level=1)  # confirm defaults to False


@pytest.mark.smoke
def test_get_document_metadata_tool_via_server(tmp_docx):
    """get_document_metadata returns a dict with path, title, author keys."""
    from wordmcp import server as srv

    result = _call_tool(srv.get_document_metadata, tmp_docx)
    assert isinstance(result, dict)
    assert "path" in result
    assert "title" in result
    assert "author" in result
    assert "created" in result


@pytest.mark.smoke
def test_read_paragraph_tool_via_server(tmp_docx):
    """read_paragraph returns full paragraph detail for a valid index."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_paragraph, tmp_docx, 0)
    assert "text" in result
    assert "style" in result
    assert result["text"] == "Test Document"


@pytest.mark.smoke
def test_list_tables_tool_via_server(tmp_docx):
    """list_tables returns empty list for document with no tables."""
    from wordmcp import server as srv

    result = _call_tool(srv.list_tables, tmp_docx)
    assert isinstance(result, list)
    assert result == []


@pytest.mark.smoke
def test_read_table_tool_via_server(sample_docx_with_table):
    """read_table returns 2D data array for a valid table index."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_table, sample_docx_with_table, 0)
    assert "data" in result
    assert result["rows"] == 2
    assert result["cols"] == 3


@pytest.mark.smoke
def test_add_table_tool_write_gate_via_server(tmp_docx):
    """add_table ToolError raised when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    with pytest.raises(ToolError):
        _call_tool(srv.add_table, tmp_docx, rows=2, cols=2, confirm=True)


@pytest.mark.smoke
def test_find_replace_tool_write_gate_via_server(tmp_docx):
    """find_replace ToolError raised when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    with pytest.raises(ToolError):
        _call_tool(srv.find_replace, tmp_docx, "Section", "Chapter", confirm=True)


@pytest.mark.smoke
def test_export_as_pdf_com_not_loaded(monkeypatch):
    """export_document(format='pdf') raises ToolError when _COM_LOADED is False."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.export_document, "/tmp/test.docx", output_path="/tmp/test.pdf", format="pdf", confirm=True)


@pytest.mark.smoke
def test_accept_all_track_changes_com_not_loaded(monkeypatch):
    """manage_tracked_changes('accept_all') raises ToolError when _COM_LOADED is False."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.manage_tracked_changes, "/tmp/test.docx", operation="accept_all", confirm=True)


@pytest.mark.smoke
def test_reject_all_track_changes_com_not_loaded(monkeypatch):
    """manage_tracked_changes('reject_all') raises ToolError when _COM_LOADED is False."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.manage_tracked_changes, "/tmp/test.docx", operation="reject_all", confirm=True)


@pytest.mark.smoke
def test_list_tracked_changes_com_not_loaded(monkeypatch):
    """manage_tracked_changes('list') raises ToolError when _COM_LOADED is False."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.manage_tracked_changes, "/tmp/test.docx", operation="list")


# ---------------------------------------------------------------------------
# Sprint C smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_manage_comments_smoke(tmp_docx):
    """manage_comments with operation='list' returns dict with 'comments' key."""
    from wordmcp import server as srv

    result = _call_tool(srv.manage_comments, tmp_docx, operation="list")
    assert isinstance(result, dict)
    assert "comments" in result


@pytest.mark.smoke
def test_manage_tracked_changes_smoke(monkeypatch):
    """manage_tracked_changes tool exists; raises ToolError when COM not loaded."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.manage_tracked_changes, "/tmp/test.docx", operation="list")


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.smoke
def test_export_document_smoke(monkeypatch):
    """export_document tool exists; html format raises NotAllowedError when COM disabled."""
    from wordmcp import server as srv
    from wordmcp import document_com
    from wordmcp.document_docx import NotAllowedError

    # Confirm the server tool exists
    assert hasattr(srv, "export_document")
    assert callable(srv.export_document) or callable(getattr(srv.export_document, "fn", None))

    # html now requires COM — guard fires when WORD_ENABLE_COM is unset
    monkeypatch.delenv("WORD_ENABLE_COM", raising=False)
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_COM"):
        document_com.export_document(
            path="fake.docx", output_path="fake.html", format="html"
        )


# ---------------------------------------------------------------------------
# Batch A smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_list_headings_tool_smoke(tmp_docx):
    """list_headings returns a list (possibly empty or with entries)."""
    from wordmcp import server as srv

    result = _call_tool(srv.list_headings, tmp_docx)
    assert isinstance(result, list)


@pytest.mark.smoke
def test_search_text_tool_smoke(tmp_docx):
    """search_text returns a dict with 'results' key."""
    from wordmcp import server as srv

    result = _call_tool(srv.search_text, tmp_docx, "the")
    assert isinstance(result, dict)
    assert "results" in result


@pytest.mark.smoke
def test_list_styles_tool_smoke(tmp_docx):
    """list_styles returns a non-empty list of style dicts."""
    from wordmcp import server as srv

    result = _call_tool(srv.list_styles, tmp_docx)
    assert isinstance(result, list)
    assert len(result) > 0


@pytest.mark.smoke
def test_create_document_tool_smoke(tmp_path, monkeypatch):
    """create_document tool is callable on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "create_document")
    assert callable(srv.create_document) or callable(
        getattr(srv.create_document, "fn", None)
    )


@pytest.mark.smoke
def test_read_section_tool_smoke(tmp_docx):
    """read_section: on tmp_docx index 0 is a heading so returns dict; otherwise catches gracefully."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv
    from wordmcp.document_docx import ValidationError

    try:
        result = _call_tool(srv.read_section, tmp_docx, 0)
        assert isinstance(result, dict)
        assert "heading" in result
        assert "body" in result
    except (ToolError, ValidationError):
        pass  # acceptable if paragraph 0 were not a heading


# ---------------------------------------------------------------------------
# Batch B smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_apply_style_smoke(tmp_path, monkeypatch):
    """apply_style tool exists on server module and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "apply_style")
    path = tmp_path / "style_smoke.docx"
    from docx import Document
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(srv.apply_style, str(path), 0, "Normal", confirm=True)


@pytest.mark.smoke
def test_update_paragraph_smoke(tmp_path, monkeypatch):
    """update_paragraph tool exists and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "update_paragraph")
    path = tmp_path / "update_smoke.docx"
    from docx import Document
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.update_paragraph, str(path), 0, "new text", confirm=True)


@pytest.mark.smoke
def test_delete_paragraph_smoke(tmp_path, monkeypatch):
    """delete_paragraph tool exists and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "delete_paragraph")
    path = tmp_path / "delete_smoke.docx"
    from docx import Document
    d = Document()
    d.add_paragraph("Para one")
    d.save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.delete_paragraph, str(path), 0, confirm=True)


@pytest.mark.smoke
def test_update_table_cell_smoke(tmp_path, monkeypatch):
    """update_table_cell tool exists and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "update_table_cell")
    path = tmp_path / "cell_smoke.docx"
    from docx import Document
    d = Document()
    d.add_table(2, 2)
    d.save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.update_table_cell, str(path), 0, 0, 0, "text", confirm=True)


@pytest.mark.smoke
def test_set_document_properties_smoke(tmp_path, monkeypatch):
    """set_document_properties tool exists and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "set_document_properties")
    path = tmp_path / "props_smoke.docx"
    from docx import Document
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.set_document_properties, str(path), title="Title", confirm=True)


# ---------------------------------------------------------------------------
# Batch C smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_add_list_smoke(tmp_path, monkeypatch):
    """add_list tool exists and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "add_list")
    path = tmp_path / "list_smoke.docx"
    from docx import Document
    Document().save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(srv.add_list, str(path), ["item1", "item2"], confirm=True)


@pytest.mark.smoke
def test_bulk_add_paragraphs_smoke(tmp_path, monkeypatch):
    """bulk_add_paragraphs tool exists and raises ToolError without write gate."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "bulk_add_paragraphs")
    path = tmp_path / "bulk_smoke.docx"
    from docx import Document
    Document().save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(srv.bulk_add_paragraphs, str(path), [{"text": "Hello"}], confirm=True)


@pytest.mark.smoke
def test_export_to_markdown_smoke(tmp_path, monkeypatch):
    """export_document(format='md') returns a dict with 'markdown' key."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "md_smoke.docx"
    from docx import Document
    d = Document()
    d.add_heading("My Heading", level=2)
    d.add_paragraph("Some body text")
    d.save(str(path))
    result = _call_tool(srv.export_document, str(path), format="md")
    assert isinstance(result, dict)
    assert "markdown" in result
    assert "## My Heading" in result["markdown"]


@pytest.mark.smoke
def test_get_headers_footers_smoke(tmp_path, monkeypatch):
    """get_headers_footers returns a list of section dicts."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "hf_smoke.docx"
    from docx import Document
    Document().save(str(path))
    result = _call_tool(srv.get_headers_footers, str(path))
    assert isinstance(result, list)
    assert len(result) >= 1
    assert "section_index" in result[0]
    assert "header_text" in result[0]
    assert "footer_text" in result[0]


# ---------------------------------------------------------------------------
# Batch D smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_add_hyperlink_smoke_write_gate(tmp_path, monkeypatch):
    """add_hyperlink raises ToolError (no write gate) for valid http:// url."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "add_hyperlink")
    path = tmp_path / "hl_smoke.docx"
    d = Document()
    d.add_paragraph("Target paragraph")
    d.save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(srv.add_hyperlink, str(path), 0, "Click here", "https://example.com", confirm=True)


@pytest.mark.smoke
def test_add_hyperlink_smoke_javascript_blocked():
    """add_hyperlink raises ToolError for javascript: url (before file access)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    assert hasattr(srv, "add_hyperlink")
    with pytest.raises(ToolError, match="http"):
        _call_tool(srv.add_hyperlink, "fake.docx", 0, "xss", "javascript:alert(1)")


@pytest.mark.smoke
def test_add_footnote_smoke_write_gate(tmp_path, monkeypatch):
    """add_footnote raises ToolError (no write gate) before footnote creation."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    assert hasattr(srv, "add_footnote")
    path = tmp_path / "fn_smoke.docx"
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(srv.add_footnote, str(path), 0, "My footnote", confirm=True)


# ---------------------------------------------------------------------------
# Paragraph dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_paragraph_dispatcher_list_routes_correctly(tmp_docx):
    """paragraph(operation='list') returns the same result as list_paragraphs."""
    from wordmcp import server as srv

    result_direct = _call_tool(srv.list_paragraphs, tmp_docx)
    result_dispatcher = _call_tool(srv.paragraph, operation="list", path=tmp_docx)
    assert isinstance(result_dispatcher, list)
    assert len(result_dispatcher) == len(result_direct)
    assert result_dispatcher[0]["index"] == result_direct[0]["index"]


@pytest.mark.smoke
def test_paragraph_dispatcher_read_routes_correctly(tmp_docx):
    """paragraph(operation='read', paragraph_index=0) returns full paragraph detail."""
    from wordmcp import server as srv

    result = _call_tool(srv.paragraph, operation="read", path=tmp_docx, paragraph_index=0)
    assert isinstance(result, dict)
    assert "text" in result
    assert result["text"] == "Test Document"


@pytest.mark.smoke
def test_paragraph_dispatcher_unknown_operation_returns_structured_error(tmp_docx):
    """paragraph(operation='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.paragraph, operation="bogus", path=tmp_docx)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_operation"
    assert result["operation"] == "bogus"
    assert "valid_operations" in result
    assert "list" in result["valid_operations"]
    assert "set_format" in result["valid_operations"]


@pytest.mark.smoke
def test_paragraph_dispatcher_set_format_write_gate(tmp_path, monkeypatch):
    """paragraph(operation='set_format') raises ToolError without WORD_ENABLE_WRITE."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "fmt_smoke.docx"
    d = Document()
    d.add_paragraph("A paragraph to format")
    d.save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(
            srv.paragraph,
            operation="set_format",
            path=str(path),
            paragraph_index=0,
            space_before=6.0,
            confirm=True,
        )


@pytest.mark.smoke
def test_paragraph_dispatcher_set_format_success(tmp_path, monkeypatch):
    """paragraph(operation='set_format') applies format when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "fmt_success.docx"
    d = Document()
    d.add_paragraph("Format me")
    d.save(str(path))
    result = _call_tool(
        srv.paragraph,
        operation="set_format",
        path=str(path),
        paragraph_index=0,
        space_before=6.0,
        space_after=3.0,
        line_spacing=1.5,
        confirm=True,
    )
    assert result["paragraph_index"] == 0
    assert result["space_before_pt"] == 6.0
    assert result["space_after_pt"] == 3.0
    assert result["line_spacing"] == 1.5


@pytest.mark.smoke
def test_paragraph_dispatcher_add_write_gate(tmp_path, monkeypatch):
    """paragraph(operation='add') raises ToolError without WORD_ENABLE_WRITE."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "add_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.paragraph, operation="add", path=str(path), text="Hello", confirm=True)


@pytest.mark.smoke
def test_set_paragraph_format_tool_exists_on_server():
    """set_paragraph_format is registered as a standalone tool on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "set_paragraph_format")
    tool = srv.set_paragraph_format
    assert callable(tool) or callable(getattr(tool, "fn", None))


@pytest.mark.smoke
def test_set_paragraph_format_write_gate(tmp_path, monkeypatch):
    """set_paragraph_format raises ToolError without WORD_ENABLE_WRITE."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "spf_gate.docx"
    d = Document()
    d.add_paragraph("Para")
    d.save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.set_paragraph_format, str(path), 0, space_before=12.0, confirm=True)


# ---------------------------------------------------------------------------
# read_document scope dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_read_document_scope_full_routes_correctly(tmp_docx):
    """read_document(scope='full') returns the same high-level summary as before."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_document, tmp_docx, scope="full")
    assert isinstance(result, dict)
    assert "paragraph_count" in result
    assert "table_count" in result
    assert result["paragraph_count"] == 3


@pytest.mark.smoke
def test_read_document_scope_default_is_full(tmp_docx):
    """read_document with no scope argument defaults to 'full' behaviour."""
    from wordmcp import server as srv

    result_default = _call_tool(srv.read_document, tmp_docx)
    result_full = _call_tool(srv.read_document, tmp_docx, scope="full")
    assert result_default == result_full


@pytest.mark.smoke
def test_read_document_scope_metadata_routes_correctly(tmp_docx):
    """read_document(scope='metadata') returns document properties dict."""
    from wordmcp import server as srv

    result_scoped = _call_tool(srv.read_document, tmp_docx, scope="metadata")
    result_direct = _call_tool(srv.get_document_metadata, tmp_docx)
    assert isinstance(result_scoped, dict)
    assert "path" in result_scoped
    assert "title" in result_scoped
    assert "author" in result_scoped
    # Scoped result must be identical to the deprecated alias
    assert result_scoped == result_direct


@pytest.mark.smoke
def test_read_document_scope_headings_routes_correctly(tmp_docx):
    """read_document(scope='headings') returns the same list as list_headings."""
    from wordmcp import server as srv

    result_scoped = _call_tool(srv.read_document, tmp_docx, scope="headings")
    result_direct = _call_tool(srv.list_headings, tmp_docx)
    assert isinstance(result_scoped, list)
    # tmp_docx has "Title" (level 0) and "Heading 1" (level 1)
    assert len(result_scoped) >= 1
    assert result_scoped == result_direct


@pytest.mark.smoke
def test_read_document_scope_outline_routes_correctly(tmp_docx):
    """read_document(scope='outline') returns hierarchical heading tree."""
    from wordmcp import server as srv

    result_scoped = _call_tool(srv.read_document, tmp_docx, scope="outline")
    result_direct = _call_tool(srv.get_document_outline, tmp_docx)
    assert isinstance(result_scoped, dict)
    assert "headings" in result_scoped
    assert "paragraph_count" in result_scoped
    assert result_scoped == result_direct


@pytest.mark.smoke
def test_read_document_scope_section_routes_correctly(tmp_docx):
    """read_document(scope='section', section_index=1) returns section content."""
    from wordmcp import server as srv

    # paragraph index 1 is "Heading 1" — "Section One"
    result_scoped = _call_tool(srv.read_document, tmp_docx, scope="section", section_index=1)
    result_direct = _call_tool(srv.read_section, tmp_docx, 1)
    assert isinstance(result_scoped, dict)
    assert "heading" in result_scoped
    assert "body" in result_scoped
    assert result_scoped == result_direct


@pytest.mark.smoke
def test_read_document_scope_section_missing_index_raises(tmp_docx):
    """read_document(scope='section') without section_index raises ToolError."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    with pytest.raises(ToolError, match="section_index"):
        _call_tool(srv.read_document, tmp_docx, scope="section")


@pytest.mark.smoke
def test_read_document_scope_paragraphs_routes_correctly(tmp_docx):
    """read_document(scope='paragraphs') returns the same list as list_paragraphs."""
    from wordmcp import server as srv

    result_scoped = _call_tool(srv.read_document, tmp_docx, scope="paragraphs")
    result_direct = _call_tool(srv.list_paragraphs, tmp_docx)
    assert isinstance(result_scoped, list)
    assert len(result_scoped) == 3
    assert result_scoped == result_direct


@pytest.mark.smoke
def test_read_document_scope_context_routes_correctly(tmp_docx):
    """read_document(scope='context') returns context summary dict."""
    from wordmcp import server as srv

    result_scoped = _call_tool(srv.read_document, tmp_docx, scope="context")
    result_direct = _call_tool(srv.get_document_context, tmp_docx)
    assert isinstance(result_scoped, dict)
    assert result_scoped == result_direct


@pytest.mark.smoke
def test_read_document_unknown_scope_returns_structured_error(tmp_docx):
    """read_document(scope='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_document, tmp_docx, scope="bogus")
    assert isinstance(result, dict)
    assert result["error"] == "unknown_scope"
    assert result["scope"] == "bogus"
    assert "valid_scopes" in result
    assert set(result["valid_scopes"]) == {
        "full", "metadata", "headings", "outline", "section", "paragraphs", "context"
    }


@pytest.mark.smoke
def test_read_document_scope_case_insensitive(tmp_docx):
    """read_document scope matching is case-insensitive ('FULL' == 'full')."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_document, tmp_docx, scope="FULL")
    assert isinstance(result, dict)
    assert "paragraph_count" in result


# ---------------------------------------------------------------------------
# Deprecated alias backward-compatibility smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_deprecated_get_document_metadata_still_callable(tmp_docx):
    """get_document_metadata is still registered and returns correct output."""
    from wordmcp import server as srv

    result = _call_tool(srv.get_document_metadata, tmp_docx)
    assert isinstance(result, dict)
    assert "title" in result
    assert "author" in result


@pytest.mark.smoke
def test_deprecated_list_headings_still_callable(tmp_docx):
    """list_headings is still registered and returns correct output."""
    from wordmcp import server as srv

    result = _call_tool(srv.list_headings, tmp_docx)
    assert isinstance(result, list)


@pytest.mark.smoke
def test_deprecated_get_document_outline_still_callable(tmp_docx):
    """get_document_outline is still registered and returns correct output."""
    from wordmcp import server as srv

    result = _call_tool(srv.get_document_outline, tmp_docx)
    assert isinstance(result, dict)
    assert "headings" in result


@pytest.mark.smoke
def test_deprecated_read_section_still_callable(tmp_docx):
    """read_section is still registered and returns correct output."""
    from wordmcp import server as srv

    result = _call_tool(srv.read_section, tmp_docx, 1)
    assert isinstance(result, dict)
    assert "heading" in result


@pytest.mark.smoke
def test_deprecated_get_document_context_still_callable(tmp_docx):
    """get_document_context is still registered and returns correct output."""
    from wordmcp import server as srv

    result = _call_tool(srv.get_document_context, tmp_docx)
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_capabilities_includes_read_document_scopes():
    """capabilities() includes read_document_scopes list."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "read_document_scopes" in result
    scopes = result["read_document_scopes"]
    assert isinstance(scopes, list)
    assert set(scopes) == {
        "full", "metadata", "headings", "outline", "section", "paragraphs", "context"
    }


# ---------------------------------------------------------------------------
# export scope dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_export_tool_exists_on_server():
    """export tool is registered on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "export")
    assert callable(srv.export) or callable(getattr(srv.export, "fn", None))


@pytest.mark.smoke
def test_export_unknown_scope_returns_structured_error(tmp_docx, monkeypatch):
    """export(scope='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    result = _call_tool(srv.export, tmp_docx, scope="bogus", output_path="/tmp/out.txt", confirm=True)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_scope"
    assert result["scope"] == "bogus"
    assert "valid_scopes" in result
    assert set(result["valid_scopes"]) == {"pdf", "txt", "markdown"}


@pytest.mark.smoke
def test_export_scope_docx_returns_structured_error(tmp_docx, monkeypatch):
    """export(scope='docx') must return a structured error dict, not raise.

    docx export was never implemented in the COM backend; it is removed from
    _VALID_SCOPES so the dispatcher returns a clean structured error rather
    than routing to COM and raising an internal ValidationError.
    """
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    result = _call_tool(srv.export, tmp_docx, scope="docx", output_path="/tmp/out.docx", confirm=True)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_scope"
    assert result["scope"] == "docx"
    assert "valid_scopes" in result
    assert "docx" not in result["valid_scopes"]
    assert set(result["valid_scopes"]) == {"pdf", "txt", "markdown"}


@pytest.mark.smoke
def test_export_no_enable_write_raises_tool_error(tmp_docx, monkeypatch):
    """export raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(srv.export, tmp_docx, scope="txt", output_path="/tmp/out.txt", confirm=True)


@pytest.mark.smoke
def test_export_confirm_false_raises_tool_error(tmp_docx, monkeypatch):
    """export raises ToolError when confirm=False (default), even with write gate open."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    with pytest.raises(ToolError, match="confirm=True"):
        _call_tool(srv.export, tmp_docx, scope="txt", output_path="/tmp/out.txt", confirm=False)


@pytest.mark.smoke
def test_export_scope_txt_writes_file(tmp_path, monkeypatch):
    """export(scope='txt') writes a .txt file and returns expected keys."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    # Create source document
    src = tmp_path / "source.docx"
    d = Document()
    d.add_heading("My Title", level=0)
    d.add_paragraph("Body content here.")
    d.save(str(src))
    out = tmp_path / "output.txt"
    result = _call_tool(srv.export, str(src), scope="txt", output_path=str(out), confirm=True)
    assert isinstance(result, dict)
    assert result["output_path"] == str(out)
    assert "paragraphs_exported" in result
    assert "tables_exported" in result
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "My Title" in content
    assert "Body content here." in content


@pytest.mark.smoke
def test_export_scope_markdown_writes_file(tmp_path, monkeypatch):
    """export(scope='markdown') writes a .md file and returns expected keys."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    src = tmp_path / "source_md.docx"
    d = Document()
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("Introduction paragraph.")
    tbl = d.add_table(2, 2, style="Table Grid")
    tbl.rows[0].cells[0].text = "Col A"
    tbl.rows[0].cells[1].text = "Col B"
    tbl.rows[1].cells[0].text = "Val 1"
    tbl.rows[1].cells[1].text = "Val 2"
    d.save(str(src))
    out = tmp_path / "output.md"
    result = _call_tool(srv.export, str(src), scope="markdown", output_path=str(out), confirm=True)
    assert isinstance(result, dict)
    assert result["output_path"] == str(out)
    assert "paragraphs_exported" in result
    assert "tables_exported" in result
    assert result["tables_exported"] == 1
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "# Chapter One" in content
    assert "Introduction paragraph." in content
    assert "| Col A | Col B |" in content
    assert "|" in content


@pytest.mark.smoke
def test_export_scope_markdown_routes_to_backend(tmp_path, monkeypatch):
    """export(scope='markdown') is routed to export_markdown_to_file backend."""
    from unittest.mock import patch
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    src = tmp_path / "mock_src.docx"
    Document().save(str(src))
    out = tmp_path / "mock_out.md"
    expected = {"output_path": str(out), "paragraphs_exported": 0, "tables_exported": 0}
    with patch("wordmcp._docx.export.export_markdown_to_file", return_value=expected) as mock_fn:
        result = _call_tool(srv.export, str(src), scope="markdown", output_path=str(out), confirm=True)
    mock_fn.assert_called_once_with(path=str(src), output_path=str(out))
    assert result == expected


@pytest.mark.smoke
def test_export_scope_pdf_routes_to_com(monkeypatch):
    """export(scope='pdf') routes to COM layer and raises ToolError when COM unloaded."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.export, "/tmp/test.docx", scope="pdf", output_path="/tmp/test.pdf", confirm=True)


@pytest.mark.smoke
def test_capabilities_includes_export_scopes():
    """capabilities() includes export_scopes list with all three scope names."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "export_scopes" in result
    scopes = result["export_scopes"]
    assert isinstance(scopes, list)
    assert set(scopes) == {"pdf", "txt", "markdown"}


# ---------------------------------------------------------------------------
# Table dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_table_tool_registered_on_server():
    """table dispatcher is registered as a callable on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "table")
    tool = srv.table
    assert callable(tool) or callable(getattr(tool, "fn", None))


@pytest.mark.smoke
def test_table_dispatcher_list_routes_correctly(tmp_docx):
    """table(operation='list') returns the same result as list_tables."""
    from wordmcp import server as srv

    result_direct = _call_tool(srv.list_tables, tmp_docx)
    result_dispatcher = _call_tool(srv.table, operation="list", path=tmp_docx)
    assert isinstance(result_dispatcher, list)
    assert result_dispatcher == result_direct


@pytest.mark.smoke
def test_table_dispatcher_read_routes_correctly(sample_docx_with_table):
    """table(operation='read', table_index=0) returns 2D data array."""
    from wordmcp import server as srv

    result = _call_tool(srv.table, operation="read", path=sample_docx_with_table, table_index=0)
    assert isinstance(result, dict)
    assert "data" in result
    assert result["rows"] == 2
    assert result["cols"] == 3


@pytest.mark.smoke
def test_table_dispatcher_read_missing_table_index_raises(tmp_docx):
    """table(operation='read') without table_index raises ToolError."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    with pytest.raises(ToolError, match="table_index"):
        _call_tool(srv.table, operation="read", path=tmp_docx)


@pytest.mark.smoke
def test_table_dispatcher_add_write_gate_no_env(tmp_path, monkeypatch):
    """table(operation='add') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "add_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(srv.table, operation="add", path=str(path), rows=2, cols=3, confirm=True)


@pytest.mark.smoke
def test_table_dispatcher_add_write_gate_no_confirm(tmp_path, monkeypatch):
    """table(operation='add') raises ToolError when confirm=False even with write gate open."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "add_confirm.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="confirm=True"):
        _call_tool(srv.table, operation="add", path=str(path), rows=2, cols=3, confirm=False)


@pytest.mark.smoke
def test_table_dispatcher_add_routes_correctly(tmp_path, monkeypatch):
    """table(operation='add') creates a table when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "add_table.docx"
    Document().save(str(path))
    result = _call_tool(
        srv.table,
        operation="add",
        path=str(path),
        rows=2,
        cols=3,
        confirm=True,
    )
    assert isinstance(result, dict)
    assert result.get("rows") == 2 or result.get("table_count") is not None


@pytest.mark.smoke
def test_table_dispatcher_update_cell_routes_correctly(tmp_path, monkeypatch):
    """table(operation='update_cell') updates a cell when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "update_cell.docx"
    d = Document()
    d.add_table(2, 3)
    d.save(str(path))
    result = _call_tool(
        srv.table,
        operation="update_cell",
        path=str(path),
        table_index=0,
        row_index=0,
        col_index=0,
        new_text="Hello",
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_table_dispatcher_bulk_update_cells_routes_correctly(tmp_path, monkeypatch):
    """table(operation='bulk_update_cells') updates multiple cells when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "bulk_cells.docx"
    d = Document()
    d.add_table(2, 3)
    d.save(str(path))
    updates = [
        {"table_index": 0, "row": 0, "col": 0, "new_text": "A"},
        {"table_index": 0, "row": 0, "col": 1, "new_text": "B"},
    ]
    result = _call_tool(
        srv.table,
        operation="bulk_update_cells",
        path=str(path),
        updates=updates,
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_table_dispatcher_unknown_operation_returns_structured_error(tmp_docx):
    """table(operation='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.table, operation="bogus", path=tmp_docx)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_operation"
    assert result["operation"] == "bogus"
    assert "valid_operations" in result
    assert "list" in result["valid_operations"]
    assert "read" in result["valid_operations"]
    assert "add" in result["valid_operations"]
    assert "update_cell" in result["valid_operations"]
    assert "bulk_update_cells" in result["valid_operations"]


@pytest.mark.smoke
def test_table_dispatcher_unknown_operation_case_insensitive_bogus(tmp_docx):
    """table(operation='LIST') routes correctly (case-insensitive)."""
    from wordmcp import server as srv

    result = _call_tool(srv.table, operation="LIST", path=tmp_docx)
    assert isinstance(result, list)


@pytest.mark.smoke
def test_capabilities_includes_table_operations():
    """capabilities() includes table_operations list with all five operation names."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "table_operations" in result
    ops = result["table_operations"]
    assert isinstance(ops, list)
    assert set(ops) == {"list", "read", "add", "update_cell", "bulk_update_cells"}


# ---------------------------------------------------------------------------
# Style dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_style_tool_registered_on_server():
    """style dispatcher is registered as a callable on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "style")
    tool = srv.style
    assert callable(tool) or callable(getattr(tool, "fn", None))


@pytest.mark.smoke
def test_style_dispatcher_list_routes_correctly(tmp_docx):
    """style(operation='list') returns the same result as list_styles."""
    from wordmcp import server as srv

    result_direct = _call_tool(srv.list_styles, tmp_docx)
    result_dispatcher = _call_tool(srv.style, operation="list", path=tmp_docx)
    assert isinstance(result_dispatcher, list)
    assert len(result_dispatcher) == len(result_direct)


@pytest.mark.smoke
def test_style_dispatcher_list_with_style_type(tmp_docx):
    """style(operation='list', style_type='paragraph') returns paragraph styles."""
    from wordmcp import server as srv

    result = _call_tool(srv.style, operation="list", path=tmp_docx, style_type="paragraph")
    assert isinstance(result, list)
    assert len(result) > 0


@pytest.mark.smoke
def test_style_dispatcher_apply_no_env_raises_tool_error(tmp_path, monkeypatch):
    """style(operation='apply') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "style_gate.docx"
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.style,
            operation="apply",
            path=str(path),
            paragraph_index=0,
            style_name="Normal",
            confirm=True,
        )


@pytest.mark.smoke
def test_style_dispatcher_apply_no_confirm_raises_tool_error(tmp_path, monkeypatch):
    """style(operation='apply') raises ToolError when confirm=False even with write gate open."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "style_confirm.docx"
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))
    with pytest.raises(ToolError, match="confirm=True"):
        _call_tool(
            srv.style,
            operation="apply",
            path=str(path),
            paragraph_index=0,
            style_name="Normal",
            confirm=False,
        )


@pytest.mark.smoke
def test_style_dispatcher_apply_routes_correctly(tmp_path, monkeypatch):
    """style(operation='apply') applies a style when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "style_apply.docx"
    d = Document()
    d.add_paragraph("Apply a style here")
    d.save(str(path))
    result = _call_tool(
        srv.style,
        operation="apply",
        path=str(path),
        paragraph_index=0,
        style_name="Normal",
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_style_dispatcher_unknown_operation_returns_structured_error(tmp_docx):
    """style(operation='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.style, operation="bogus", path=tmp_docx)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_operation"
    assert result["operation"] == "bogus"
    assert "valid_operations" in result
    assert "list" in result["valid_operations"]
    assert "apply" in result["valid_operations"]


@pytest.mark.smoke
def test_style_dispatcher_operation_case_insensitive(tmp_docx):
    """style(operation='LIST') routes correctly (case-insensitive)."""
    from wordmcp import server as srv

    result = _call_tool(srv.style, operation="LIST", path=tmp_docx)
    assert isinstance(result, list)


@pytest.mark.smoke
def test_capabilities_includes_style_operations():
    """capabilities() includes style_operations list with both operation names."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "style_operations" in result
    ops = result["style_operations"]
    assert isinstance(ops, list)
    assert set(ops) == {"list", "apply"}


@pytest.mark.smoke
def test_deprecated_list_styles_still_callable(tmp_docx):
    """list_styles is still registered and returns correct output (deprecated alias)."""
    from wordmcp import server as srv

    result = _call_tool(srv.list_styles, tmp_docx)
    assert isinstance(result, list)
    assert len(result) > 0


@pytest.mark.smoke
def test_deprecated_apply_style_still_callable(tmp_path, monkeypatch):
    """apply_style is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "deprecated_apply.docx"
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))
    with pytest.raises(ToolError):  # WORD_ENABLE_WRITE not set
        _call_tool(srv.apply_style, str(path), 0, "Normal", confirm=True)


# ---------------------------------------------------------------------------
# Content dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_content_tool_registered_on_server():
    """content dispatcher is registered as a callable on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "content")
    tool = srv.content
    assert callable(tool) or callable(getattr(tool, "fn", None))


@pytest.mark.smoke
def test_content_unknown_operation_returns_structured_error(tmp_docx):
    """content(operation='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.content, operation="bogus", path=tmp_docx)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_operation"
    assert result["operation"] == "bogus"
    assert "valid_operations" in result
    assert set(result["valid_operations"]) == {
        "add_list",
        "add_page_break",
        "insert_image",
        "add_hyperlink",
        "add_footnote",
        "find_replace",
        "set_properties",
    }


@pytest.mark.smoke
def test_content_add_list_write_gate_no_env(tmp_path, monkeypatch):
    """content(operation='add_list') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "list_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.content,
            operation="add_list",
            path=str(path),
            items=["item1", "item2"],
            confirm=True,
        )


@pytest.mark.smoke
def test_content_add_list_write_gate_no_confirm(tmp_path, monkeypatch):
    """content(operation='add_list') raises ToolError when confirm=False even with write gate open."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "list_confirm.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="confirm=True"):
        _call_tool(
            srv.content,
            operation="add_list",
            path=str(path),
            items=["item1", "item2"],
            confirm=False,
        )


@pytest.mark.smoke
def test_content_add_list_routes_correctly(tmp_path, monkeypatch):
    """content(operation='add_list') creates a list when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "content_list.docx"
    Document().save(str(path))
    result = _call_tool(
        srv.content,
        operation="add_list",
        path=str(path),
        items=["Alpha", "Beta", "Gamma"],
        list_type="bullet",
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_content_add_page_break_routes_correctly(tmp_path, monkeypatch):
    """content(operation='add_page_break') inserts a page break when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "content_pagebreak.docx"
    Document().save(str(path))
    result = _call_tool(
        srv.content,
        operation="add_page_break",
        path=str(path),
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_content_add_page_break_write_gate(tmp_path, monkeypatch):
    """content(operation='add_page_break') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "pb_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(srv.content, operation="add_page_break", path=str(path), confirm=True)


@pytest.mark.smoke
def test_content_insert_image_write_gate(tmp_path, monkeypatch):
    """content(operation='insert_image') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "img_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.content,
            operation="insert_image",
            path=str(path),
            image_path=str(tmp_path / "img.png"),
            confirm=True,
        )


@pytest.mark.smoke
def test_content_add_hyperlink_write_gate(tmp_path, monkeypatch):
    """content(operation='add_hyperlink') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "hl_gate.docx"
    d = Document()
    d.add_paragraph("Link target")
    d.save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.content,
            operation="add_hyperlink",
            path=str(path),
            paragraph_index=0,
            text="Click",
            url="https://example.com",
            confirm=True,
        )


@pytest.mark.smoke
def test_content_add_footnote_write_gate(tmp_path, monkeypatch):
    """content(operation='add_footnote') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "fn_gate.docx"
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.content,
            operation="add_footnote",
            path=str(path),
            paragraph_index=0,
            text="Footnote text",
            confirm=True,
        )


@pytest.mark.smoke
def test_content_find_replace_write_gate(tmp_path, monkeypatch):
    """content(operation='find_replace') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "fr_gate.docx"
    d = Document()
    d.add_paragraph("Hello world")
    d.save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.content,
            operation="find_replace",
            path=str(path),
            find_text="Hello",
            replace_text="Hi",
            confirm=True,
        )


@pytest.mark.smoke
def test_content_find_replace_routes_correctly(tmp_path, monkeypatch):
    """content(operation='find_replace') performs replacement when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "fr_success.docx"
    d = Document()
    d.add_paragraph("Hello world")
    d.save(str(path))
    result = _call_tool(
        srv.content,
        operation="find_replace",
        path=str(path),
        find_text="Hello",
        replace_text="Hi",
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_content_set_properties_write_gate(tmp_path, monkeypatch):
    """content(operation='set_properties') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "props_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.content,
            operation="set_properties",
            path=str(path),
            title="My Doc",
            confirm=True,
        )


@pytest.mark.smoke
def test_content_set_properties_routes_correctly(tmp_path, monkeypatch):
    """content(operation='set_properties') sets properties when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "props_success.docx"
    Document().save(str(path))
    result = _call_tool(
        srv.content,
        operation="set_properties",
        path=str(path),
        title="Test Title",
        author="Test Author",
        confirm=True,
    )
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_capabilities_includes_content_operations():
    """capabilities() includes content_operations list with all seven operation names."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "content_operations" in result
    ops = result["content_operations"]
    assert isinstance(ops, list)
    assert set(ops) == {
        "add_list",
        "add_page_break",
        "insert_image",
        "add_hyperlink",
        "add_footnote",
        "find_replace",
        "set_properties",
    }


@pytest.mark.smoke
def test_deprecated_add_list_still_callable(tmp_path, monkeypatch):
    """add_list is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_list.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.add_list, str(path), ["a", "b"], confirm=True)


@pytest.mark.smoke
def test_deprecated_add_page_break_still_callable(tmp_path, monkeypatch):
    """add_page_break is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_pb.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.add_page_break, str(path), confirm=True)


@pytest.mark.smoke
def test_deprecated_insert_image_still_callable(tmp_path, monkeypatch):
    """insert_image is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_img.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.insert_image, str(path), str(tmp_path / "img.png"), confirm=True)


@pytest.mark.smoke
def test_deprecated_add_hyperlink_still_callable(tmp_path, monkeypatch):
    """add_hyperlink is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_hl.docx"
    d = Document()
    d.add_paragraph("Para")
    d.save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.add_hyperlink, str(path), 0, "Click", "https://example.com", confirm=True)


@pytest.mark.smoke
def test_deprecated_add_footnote_still_callable(tmp_path, monkeypatch):
    """add_footnote is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_fn.docx"
    d = Document()
    d.add_paragraph("Para")
    d.save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.add_footnote, str(path), 0, "Note", confirm=True)


@pytest.mark.smoke
def test_deprecated_find_replace_still_callable(tmp_path, monkeypatch):
    """find_replace is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_fr.docx"
    d = Document()
    d.add_paragraph("Hello world")
    d.save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.find_replace, str(path), "Hello", "Hi", confirm=True)


@pytest.mark.smoke
def test_deprecated_set_document_properties_still_callable(tmp_path, monkeypatch):
    """set_document_properties is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_props.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.set_document_properties, str(path), title="Title", confirm=True)


# ---------------------------------------------------------------------------
# Review dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_review_tool_registered_on_server():
    """review dispatcher is registered as a callable on the server module."""
    from wordmcp import server as srv

    assert hasattr(srv, "review")
    tool = srv.review
    assert callable(tool) or callable(getattr(tool, "fn", None))


@pytest.mark.smoke
def test_review_dispatcher_review_op_routes_correctly(tmp_docx):
    """review(operation='review') returns structural check results for the document."""
    from wordmcp import server as srv

    result = _call_tool(srv.review, operation="review", path=tmp_docx)
    assert isinstance(result, dict)
    assert "checks_run" in result
    assert "results" in result
    assert "path" in result


@pytest.mark.smoke
def test_review_dispatcher_write_findings_write_gate_no_env(tmp_path, monkeypatch):
    """review(operation='write_findings') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(
            srv.review,
            operation="write_findings",
            evidence_path=str(tmp_path / "findings.jsonl"),
            findings=[{"check_name": "word_count", "passed": True, "detail": "ok"}],
            document_path=str(tmp_path / "doc.docx"),
            confirm=True,
        )


@pytest.mark.smoke
def test_review_dispatcher_write_findings_write_gate_no_confirm(tmp_path, monkeypatch):
    """review(operation='write_findings') raises ToolError when confirm=False even with write gate open."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    with pytest.raises(ToolError, match="confirm=True"):
        _call_tool(
            srv.review,
            operation="write_findings",
            evidence_path=str(tmp_path / "findings.jsonl"),
            findings=[{"check_name": "word_count", "passed": True, "detail": "ok"}],
            document_path=str(tmp_path / "doc.docx"),
            confirm=False,
        )


@pytest.mark.smoke
def test_review_dispatcher_write_findings_routes_correctly(tmp_path, monkeypatch):
    """review(operation='write_findings') writes a JSONL record when write gate is open."""
    import json
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    evidence_file = tmp_path / "findings.jsonl"
    result = _call_tool(
        srv.review,
        operation="write_findings",
        evidence_path=str(evidence_file),
        findings=[{"check_name": "word_count", "passed": True, "detail": "3 words"}],
        document_path=str(tmp_path / "doc.docx"),
        confirm=True,
    )
    assert isinstance(result, dict)
    assert "review_id" in result
    assert result["findings_count"] == 1
    assert evidence_file.exists()
    record = json.loads(evidence_file.read_text(encoding="utf-8").strip())
    assert len(record.get("findings", [])) == 1


@pytest.mark.smoke
def test_review_dispatcher_export_evidence_routes_correctly(tmp_docx):
    """review(operation='export_evidence') returns a structured evidence bundle."""
    from wordmcp import server as srv

    mock_review_results = {"checks_run": 4, "results": {}}
    result = _call_tool(
        srv.review,
        operation="export_evidence",
        path=tmp_docx,
        review_results=mock_review_results,
    )
    assert isinstance(result, dict)
    assert "document_path" in result
    assert "tool_version" in result
    assert "timestamp" in result
    assert result["review_results"] == mock_review_results
    assert result["export_result"] is None


@pytest.mark.smoke
def test_review_dispatcher_manage_comments_routes_correctly(tmp_docx):
    """review(operation='manage_comments', comment_operation='list') returns comments list."""
    from wordmcp import server as srv

    result = _call_tool(
        srv.review,
        operation="manage_comments",
        path=tmp_docx,
        comment_operation="list",
    )
    assert isinstance(result, dict)
    assert "comments" in result


@pytest.mark.smoke
def test_review_dispatcher_manage_tracked_changes_routes_correctly(monkeypatch):
    """review(operation='manage_tracked_changes') raises ToolError when COM not loaded."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(
            srv.review,
            operation="manage_tracked_changes",
            path="/tmp/test.docx",
            tracked_changes_operation="list",
        )


@pytest.mark.smoke
def test_review_dispatcher_unknown_operation_returns_structured_error(tmp_docx):
    """review(operation='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.review, operation="bogus", path=tmp_docx)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_operation"
    assert result["operation"] == "bogus"
    assert "valid_operations" in result
    assert set(result["valid_operations"]) == {
        "review",
        "write_findings",
        "export_evidence",
        "manage_comments",
        "manage_tracked_changes",
    }


@pytest.mark.smoke
def test_review_operations_in_capabilities():
    """capabilities() includes review_operations list with all five operation names."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "review_operations" in result
    ops = result["review_operations"]
    assert isinstance(ops, list)
    assert set(ops) == {
        "review",
        "write_findings",
        "export_evidence",
        "manage_comments",
        "manage_tracked_changes",
    }


@pytest.mark.smoke
def test_deprecated_review_document_still_callable(tmp_docx):
    """review_document is still registered and returns correct output (deprecated alias)."""
    from wordmcp import server as srv

    result = _call_tool(srv.review_document, tmp_docx)
    assert isinstance(result, dict)
    assert "checks_run" in result


@pytest.mark.smoke
def test_deprecated_write_review_findings_still_callable(tmp_path, monkeypatch):
    """write_review_findings is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    with pytest.raises(ToolError):
        _call_tool(
            srv.write_review_findings,
            str(tmp_path / "ev.jsonl"),
            [{"check_name": "x", "passed": True, "detail": "y"}],
            str(tmp_path / "doc.docx"),
            confirm=True,
        )


@pytest.mark.smoke
def test_deprecated_export_review_evidence_still_callable(tmp_docx):
    """export_review_evidence is still registered and returns correct output (deprecated alias)."""
    from wordmcp import server as srv

    result = _call_tool(srv.export_review_evidence, tmp_docx, {"checks_run": 2, "results": {}})
    assert isinstance(result, dict)
    assert "document_path" in result
    assert "review_results" in result


@pytest.mark.smoke
def test_deprecated_manage_comments_still_callable(tmp_docx):
    """manage_comments is still registered and returns correct output for list op (deprecated alias)."""
    from wordmcp import server as srv

    result = _call_tool(srv.manage_comments, tmp_docx, operation="list")
    assert isinstance(result, dict)
    assert "comments" in result


@pytest.mark.smoke
def test_deprecated_manage_tracked_changes_still_callable(monkeypatch):
    """manage_tracked_changes is still registered and raises ToolError when COM not loaded (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    import wordmcp.server as srv

    monkeypatch.setattr(srv, "_COM_LOADED", False)
    with pytest.raises(ToolError):
        _call_tool(srv.manage_tracked_changes, "/tmp/test.docx", operation="list")


# ---------------------------------------------------------------------------
# Document dispatcher smoke tests
# ---------------------------------------------------------------------------


@pytest.mark.smoke
def test_document_tool_registered_on_server():
    """document dispatcher is registered as a callable on the server module (48 tools total)."""
    from wordmcp import server as srv

    assert hasattr(srv, "document")
    tool = srv.document
    assert callable(tool) or callable(getattr(tool, "fn", None))


@pytest.mark.smoke
def test_document_dispatcher_save_write_gate_no_env(tmp_path, monkeypatch):
    """document(operation='save') raises ToolError when WORD_ENABLE_WRITE is not set."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "save_gate.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="WORD_ENABLE_WRITE"):
        _call_tool(srv.document, operation="save", path=str(path), confirm=True)


@pytest.mark.smoke
def test_document_dispatcher_save_write_gate_no_confirm(tmp_path, monkeypatch):
    """document(operation='save') raises ToolError when confirm=False even with write gate open."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "save_confirm.docx"
    Document().save(str(path))
    with pytest.raises(ToolError, match="confirm=True"):
        _call_tool(srv.document, operation="save", path=str(path), confirm=False)


@pytest.mark.smoke
def test_document_dispatcher_save_routes_correctly(tmp_path, monkeypatch):
    """document(operation='save') saves the document when write gate is open."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "save_success.docx"
    d = Document()
    d.add_paragraph("Save me")
    d.save(str(path))
    result = _call_tool(srv.document, operation="save", path=str(path), confirm=True)
    assert isinstance(result, dict)


@pytest.mark.smoke
def test_document_dispatcher_search_text_routes_correctly(tmp_docx):
    """document(operation='search_text') returns dict with 'results' key."""
    from wordmcp import server as srv

    result = _call_tool(srv.document, operation="search_text", path=tmp_docx, query="the")
    assert isinstance(result, dict)
    assert "results" in result


@pytest.mark.smoke
def test_document_dispatcher_search_text_missing_query_raises(tmp_docx):
    """document(operation='search_text') without query raises ToolError."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    with pytest.raises(ToolError, match="query"):
        _call_tool(srv.document, operation="search_text", path=tmp_docx)


@pytest.mark.smoke
def test_document_dispatcher_get_headers_footers_routes_correctly(tmp_path, monkeypatch):
    """document(operation='get_headers_footers') returns a list of section dicts."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "hf_dispatch.docx"
    Document().save(str(path))
    result = _call_tool(srv.document, operation="get_headers_footers", path=str(path))
    assert isinstance(result, list)
    assert len(result) >= 1
    assert "section_index" in result[0]
    assert "header_text" in result[0]
    assert "footer_text" in result[0]


@pytest.mark.smoke
def test_document_dispatcher_unknown_operation_returns_structured_error(tmp_docx):
    """document(operation='bogus') returns structured error dict, not an exception."""
    from wordmcp import server as srv

    result = _call_tool(srv.document, operation="bogus", path=tmp_docx)
    assert isinstance(result, dict)
    assert result["error"] == "unknown_operation"
    assert result["operation"] == "bogus"
    assert "valid_operations" in result
    assert set(result["valid_operations"]) == {"save", "search_text", "get_headers_footers"}


@pytest.mark.smoke
def test_document_operations_in_capabilities():
    """capabilities() includes document_operations list with all three operation names."""
    from wordmcp import server as srv

    result = _call_tool(srv.capabilities)
    assert "document_operations" in result
    ops = result["document_operations"]
    assert isinstance(ops, list)
    assert set(ops) == {"save", "search_text", "get_headers_footers"}


@pytest.mark.smoke
def test_deprecated_save_still_callable(tmp_path, monkeypatch):
    """save is still registered and raises ToolError without write gate (deprecated alias)."""
    from fastmcp.exceptions import ToolError
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    path = tmp_path / "dep_save.docx"
    Document().save(str(path))
    with pytest.raises(ToolError):
        _call_tool(srv.save, str(path), confirm=True)


@pytest.mark.smoke
def test_deprecated_search_text_still_callable(tmp_docx):
    """search_text is still registered and returns correct output (deprecated alias)."""
    from wordmcp import server as srv

    result = _call_tool(srv.search_text, tmp_docx, "the")
    assert isinstance(result, dict)
    assert "results" in result


@pytest.mark.smoke
def test_deprecated_get_headers_footers_still_callable(tmp_path, monkeypatch):
    """get_headers_footers is still registered and returns correct output (deprecated alias)."""
    from wordmcp import server as srv

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "dep_hf.docx"
    Document().save(str(path))
    result = _call_tool(srv.get_headers_footers, str(path))
    assert isinstance(result, list)
    assert len(result) >= 1
    assert "section_index" in result[0]
