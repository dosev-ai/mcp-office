"""
Unit tests — Batch C (add_list/bulk_add_paragraphs/export_to_markdown/get_headers_footers),
Batch D (manage_comments full oxml/add_hyperlink/add_footnote),
Batch E/COM tests (accept_one/reject_one/export_html/_check_output_path_html),
Bug Regression (BUG 1-4).
"""
from __future__ import annotations

import contextlib
import sys
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from wordmcp.document_docx import (
    NotAllowedError,
    ValidationError,
)
from wordmcp._com import _write as _document_com_write
from wordmcp._com import _export as _document_com_export


# ---------------------------------------------------------------------------
# Module-level helper for Batch E COM tests
# ---------------------------------------------------------------------------

def _make_mock_word_ctx(mock_app):
    """Return a context-manager mock that yields mock_app."""

    @contextlib.contextmanager
    def _ctx():
        yield mock_app

    return _ctx


# ===========================================================================
# Batch E: COM accept_one / reject_one / html-export (12 tests)
# ===========================================================================


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_manage_tracked_changes_accept_one_success(tmp_path, monkeypatch):
    """accept_one calls Revisions(1).Accept() and returns ok dict."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "rev.docx"
    Document().save(str(path))

    mock_revision = MagicMock()
    mock_doc = MagicMock()
    mock_doc.Revisions.Count = 3
    mock_doc.Revisions.return_value = mock_revision
    mock_app = MagicMock()
    mock_app.Documents.Open.return_value = mock_doc

    with patch.object(_document_com_write, "_word_app_context", _make_mock_word_ctx(mock_app)):
        result = document_com.manage_tracked_changes(
            path=str(path), operation="accept_one", revision_index=0, confirm=True
        )

    mock_doc.Revisions.assert_called_once_with(1)  # 1-based
    mock_revision.Accept.assert_called_once()
    mock_doc.Save.assert_called_once()
    assert result == {"status": "ok", "operation": "accept_one", "revision_index": 0}


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_manage_tracked_changes_reject_one_success(tmp_path, monkeypatch):
    """reject_one calls Revisions(1).Reject() and returns ok dict."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "rev2.docx"
    Document().save(str(path))

    mock_revision = MagicMock()
    mock_doc = MagicMock()
    mock_doc.Revisions.Count = 2
    mock_doc.Revisions.return_value = mock_revision
    mock_app = MagicMock()
    mock_app.Documents.Open.return_value = mock_doc

    with patch.object(_document_com_write, "_word_app_context", _make_mock_word_ctx(mock_app)):
        result = document_com.manage_tracked_changes(
            path=str(path), operation="reject_one", revision_index=1, confirm=True
        )

    mock_doc.Revisions.assert_called_once_with(2)  # 1-based: index 1 → Revisions(2)
    mock_revision.Reject.assert_called_once()
    mock_doc.Save.assert_called_once()
    assert result == {"status": "ok", "operation": "reject_one", "revision_index": 1}


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_manage_tracked_changes_accept_one_out_of_range(tmp_path, monkeypatch):
    """accept_one raises OfficeCOMError when revision_index >= count."""
    from wordmcp import document_com
    from wordmcp.document_com import OfficeCOMError

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "oor.docx"
    Document().save(str(path))

    mock_doc = MagicMock()
    mock_doc.Revisions.Count = 2
    mock_app = MagicMock()
    mock_app.Documents.Open.return_value = mock_doc

    with patch.object(_document_com_write, "_word_app_context", _make_mock_word_ctx(mock_app)):
        with pytest.raises(OfficeCOMError, match="out of range"):
            document_com.manage_tracked_changes(
                path=str(path), operation="accept_one", revision_index=5, confirm=True
            )


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_manage_tracked_changes_reject_one_none_revision(tmp_path, monkeypatch):
    """reject_one raises OfficeCOMError when revision_index is None."""
    from wordmcp import document_com
    from wordmcp.document_com import OfficeCOMError

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "none_rev.docx"
    Document().save(str(path))

    mock_doc = MagicMock()
    mock_doc.Revisions.Count = 1
    mock_app = MagicMock()
    mock_app.Documents.Open.return_value = mock_doc

    with patch.object(_document_com_write, "_word_app_context", _make_mock_word_ctx(mock_app)):
        with pytest.raises(OfficeCOMError, match="out of range"):
            document_com.manage_tracked_changes(
                path=str(path),
                operation="reject_one",
                revision_index=None,
                confirm=True,
            )


@pytest.mark.unit
def test_manage_tracked_changes_accept_one_no_com(monkeypatch):
    """accept_one raises NotAllowedError when WORD_ENABLE_COM is not set."""
    from wordmcp import document_com

    monkeypatch.delenv("WORD_ENABLE_COM", raising=False)
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_COM"):
        document_com.manage_tracked_changes(
            path="fake.docx", operation="accept_one", revision_index=0, confirm=True
        )


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_manage_tracked_changes_accept_one_no_confirm(tmp_path, monkeypatch):
    """accept_one raises ValidationError when confirm=False."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "noconfirm.docx"
    Document().save(str(path))
    with pytest.raises(ValidationError, match="confirm=True required"):
        document_com.manage_tracked_changes(
            path=str(path), operation="accept_one", revision_index=0, confirm=False
        )


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_manage_tracked_changes_accept_one_no_write(tmp_path, monkeypatch):
    """accept_one raises NotAllowedError when WORD_ENABLE_WRITE is not set."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "nowrite.docx"
    Document().save(str(path))
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        document_com.manage_tracked_changes(
            path=str(path), operation="accept_one", revision_index=0, confirm=True
        )


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")
@pytest.mark.unit
def test_export_document_html_success(tmp_path, monkeypatch):
    """HTML export calls SaveAs2 with FileFormat=10 and returns output_path."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ENABLE_COM", "true")
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    src = tmp_path / "source.docx"
    Document().save(str(src))
    out = tmp_path / "output.html"

    mock_doc = MagicMock()
    mock_app = MagicMock()
    mock_app.Documents.Open.return_value = mock_doc

    with patch.object(_document_com_export, "_word_app_context", _make_mock_word_ctx(mock_app)):
        result = document_com.export_document(
            path=str(src), output_path=str(out), format="html", confirm=True
        )

    mock_doc.SaveAs2.assert_called_once_with(
        FileName=str(out.resolve()), FileFormat=10
    )
    assert result["status"] == "ok"
    assert result["output_path"] == str(out.resolve())


@pytest.mark.unit
def test_check_output_path_html_valid(tmp_path, monkeypatch):
    """_check_output_path_html returns a resolved Path for a valid .html path."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    raw = str(tmp_path / "report.html")
    resolved = document_com._check_output_path_html(raw)
    assert resolved.suffix == ".html"
    assert resolved == (tmp_path / "report.html").resolve()


@pytest.mark.unit
def test_check_output_path_html_wrong_ext(tmp_path, monkeypatch):
    """_check_output_path_html raises ValidationError for non-.html extension."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match=".html extension"):
        document_com._check_output_path_html(str(tmp_path / "report.txt"))


@pytest.mark.unit
def test_check_output_path_html_null_byte(tmp_path, monkeypatch):
    """_check_output_path_html raises ValidationError when path contains null byte."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte"):
        document_com._check_output_path_html(str(tmp_path / "re\x00port.html"))


@pytest.mark.unit
def test_check_output_path_html_no_allowlist(monkeypatch):
    """_check_output_path_html raises ValidationError when WORD_ALLOWLIST_ROOTS is unset."""
    from wordmcp import document_com

    monkeypatch.delenv("WORD_ALLOWLIST_ROOTS", raising=False)
    with pytest.raises(ValidationError, match="WORD_ALLOWLIST_ROOTS"):
        document_com._check_output_path_html("C:/output/report.html")


@pytest.mark.unit
def test_check_output_path_html_unc(tmp_path, monkeypatch):
    """_check_output_path_html raises ValidationError for UNC paths."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError):
        document_com._check_output_path_html(r"\\server\share\out.html")


@pytest.mark.unit
def test_check_output_path_html_outside_allowlist(tmp_path, monkeypatch):
    """_check_output_path_html raises ValidationError for paths outside the allowlist."""
    from wordmcp import document_com

    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError):
        document_com._check_output_path_html(r"C:\Windows\test.html")


# ===========================================================================
