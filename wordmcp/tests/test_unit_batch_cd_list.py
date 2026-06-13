"""
Unit tests — Batch C (add_list, bulk_add_paragraphs, get_headers_footers)
and Batch D (manage_comments, add_hyperlink, add_footnote).
Split from test_unit_batch_cd.py.
"""
from __future__ import annotations

import contextlib
import sys
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from wordmcp.document_docx import (
    NotAllowedError,
    NotFoundError,
    ValidationError,
    WordMCPError,
    add_footnote,
    add_hyperlink,
    add_list,
    bulk_add_paragraphs,
    export_to_markdown,
    get_headers_footers,
    manage_comments,
)



# ---------------------------------------------------------------------------

def _make_mock_word_ctx(mock_app):
    """Return a context-manager mock that yields mock_app."""

    @contextlib.contextmanager
    def _ctx():
        yield mock_app

    return _ctx


# ===========================================================================
# Batch E: COM accept_one / reject_one / html-export (12 tests)
# ===========================================================================


@pytest.mark.skipif(sys.platform != "win32", reason="requires Windows COM")

# Batch C — add_list, bulk_add_paragraphs, export_to_markdown, get_headers_footers
# ===========================================================================


@pytest.mark.unit
def test_add_list_bullet_adds_paragraphs(write_enabled_docx):
    """add_list with list_type='bullet' appends paragraphs with List Bullet style."""
    result = add_list(write_enabled_docx, ["item A", "item B", "item C"], list_type="bullet", confirm=True)
    assert result["added"] == 3
    assert result["list_type"] == "bullet"
    d = Document(write_enabled_docx)
    # Last 3 paragraphs should have List Bullet style
    last_3 = d.paragraphs[-3:]
    for p in last_3:
        assert p.style.name == "List Bullet"


@pytest.mark.unit
def test_add_list_number_style(write_enabled_docx):
    """add_list with list_type='number' creates paragraphs with List Number style."""
    result = add_list(write_enabled_docx, ["first", "second"], list_type="number", confirm=True)
    assert result["added"] == 2
    assert result["list_type"] == "number"
    d = Document(write_enabled_docx)
    last_2 = d.paragraphs[-2:]
    for p in last_2:
        assert p.style.name == "List Number"


@pytest.mark.unit
def test_add_list_invalid_type_raises(write_enabled_docx):
    """add_list raises ValidationError for an unknown list_type."""
    with pytest.raises(ValidationError, match="Unknown list_type"):
        add_list(write_enabled_docx, ["x"], list_type="roman", confirm=True)


@pytest.mark.unit
def test_add_list_empty_raises(write_enabled_docx):
    """add_list raises ValidationError when items list is empty."""
    with pytest.raises(ValidationError, match="items must not be empty"):
        add_list(write_enabled_docx, [], confirm=True)


@pytest.mark.unit
def test_add_list_exceeds_500_raises(write_enabled_docx):
    """add_list raises ValidationError when more than 500 items are provided."""
    with pytest.raises(ValidationError, match="exceeds max 500"):
        add_list(write_enabled_docx, ["x"] * 501, confirm=True)


@pytest.mark.unit
def test_bulk_add_paragraphs_adds_all(tmp_path, monkeypatch):
    """bulk_add_paragraphs adds all 3 paragraphs in a single save."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "bulk.docx"
    Document().save(str(path))
    paras = [{"text": "Alpha"}, {"text": "Beta"}, {"text": "Gamma"}]
    result = bulk_add_paragraphs(str(path), paras, confirm=True)
    assert result["added"] == 3
    assert result["paragraph_count"] >= 3
    d = Document(str(path))
    texts = [p.text for p in d.paragraphs]
    assert "Alpha" in texts
    assert "Beta" in texts
    assert "Gamma" in texts


@pytest.mark.unit
def test_bulk_add_paragraphs_empty_raises(write_enabled_docx):
    """bulk_add_paragraphs raises ValidationError when paragraphs list is empty."""
    with pytest.raises(ValidationError, match="paragraphs must not be empty"):
        bulk_add_paragraphs(write_enabled_docx, [], confirm=True)


@pytest.mark.unit
def test_bulk_add_paragraphs_exceeds_500_raises(write_enabled_docx):
    """bulk_add_paragraphs raises ValidationError when more than 500 paragraphs are provided."""
    with pytest.raises(ValidationError, match="max 500 paragraphs"):
        bulk_add_paragraphs(write_enabled_docx, [{"text": "x"}] * 501, confirm=True)


@pytest.mark.unit
def test_bulk_add_paragraphs_unknown_style_raises(write_enabled_docx):
    """bulk_add_paragraphs raises ValidationError for an unknown style name."""
    with pytest.raises(ValidationError, match="Unknown style"):
        bulk_add_paragraphs(
            write_enabled_docx,
            [{"text": "hi", "style": "NonExistentStyle_XYZ999"}],
            confirm=True,
        )


@pytest.mark.unit
def test_bulk_add_paragraphs_validationerror_evicts_cache(tmp_path, monkeypatch):
    """ValidationError mid-loop must evict the partial doc so next read reloads from disk."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "evict.docx"
    Document().save(str(path))
    resolved_key = str(path.resolve())

    from wordmcp._docx.core import _doc_cache
    _doc_cache.pop(resolved_key, None)  # ensure clean state

    # First item valid (loads doc into cache), second has unknown style → ValidationError
    paras = [
        {"text": "Good paragraph"},
        {"text": "Bad paragraph", "style": "NonExistentStyle_XYZ999"},
    ]
    with pytest.raises(ValidationError, match="Unknown style"):
        bulk_add_paragraphs(str(path), paras, confirm=True)

    # Dirty doc must be evicted from cache — next read reloads unmodified doc from disk
    assert resolved_key not in _doc_cache

    # Disk is unmodified — atomic_save was never reached
    d = Document(str(path))
    texts = [p.text for p in d.paragraphs]
    assert "Good paragraph" not in texts


@pytest.mark.unit
def test_export_to_markdown_headings(tmp_path, monkeypatch):
    """export_to_markdown renders headings as ATX Markdown (## text)."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "md_test.docx"
    d = Document()
    d.add_heading("Chapter One", level=2)
    d.add_paragraph("Some body text")
    d.save(str(path))
    result = export_to_markdown(str(path))
    assert "## Chapter One" in result["markdown"]
    assert isinstance(result["truncated"], bool)


@pytest.mark.unit
def test_export_to_markdown_table(tmp_path, monkeypatch):
    """export_to_markdown renders tables as Markdown pipe tables."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "md_table.docx"
    d = Document()
    tbl = d.add_table(2, 3, style="Table Grid")
    tbl.rows[0].cells[0].text = "Name"
    tbl.rows[0].cells[1].text = "Age"
    tbl.rows[0].cells[2].text = "City"
    tbl.rows[1].cells[0].text = "Alice"
    tbl.rows[1].cells[1].text = "30"
    tbl.rows[1].cells[2].text = "London"
    d.save(str(path))
    result = export_to_markdown(str(path))
    assert "| Name" in result["markdown"]
    assert "| Alice" in result["markdown"]


@pytest.mark.unit
def test_get_headers_footers_returns_list(tmp_path, monkeypatch):
    """get_headers_footers returns a list with at least one section dict."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "hf_test.docx"
    Document().save(str(path))
    result = get_headers_footers(str(path))
    assert isinstance(result, list)
    assert len(result) >= 1
    section = result[0]
    assert "section_index" in section
    assert "header_text" in section
    assert "footer_text" in section
    assert "header_linked_to_previous" in section
    assert "footer_linked_to_previous" in section


# ===========================================================================
# Batch D — manage_comments (full oxml), add_hyperlink, add_footnote
# ===========================================================================


@pytest.mark.unit
def test_manage_comments_list_empty_doc(tmp_docx):
    """manage_comments 'list' returns {'comments': []} for a doc with no comments part."""
    result = manage_comments(path=tmp_docx, operation="list")
    assert result == {"comments": []}


@pytest.mark.unit
def test_manage_comments_list_returns_dict(tmp_docx):
    """manage_comments 'list' always returns a dict with a 'comments' key."""
    result = manage_comments(path=tmp_docx, operation="list")
    assert isinstance(result, dict)
    assert "comments" in result
    assert isinstance(result["comments"], list)


@pytest.mark.unit
def test_manage_comments_add_inserts_comment(write_enabled_docx):
    """manage_comments 'add' inserts a comment that 'list' can subsequently find."""
    add_result = manage_comments(
        path=write_enabled_docx,
        operation="add",
        text="Review this section",
        author="Tester",
        confirm=True,
    )
    assert add_result["added"] is True
    assert add_result["author"] == "Tester"
    new_id = add_result["comment_id"]

    list_result = manage_comments(path=write_enabled_docx, operation="list")
    ids = [c["id"] for c in list_result["comments"]]
    assert new_id in ids
    texts = [c["text"] for c in list_result["comments"]]
    assert "Review this section" in texts


@pytest.mark.unit
def test_manage_comments_add_requires_text(write_enabled_docx):
    """manage_comments 'add' with text=None (with write+confirm) raises ValidationError."""
    with pytest.raises(ValidationError, match="text is required"):
        manage_comments(
            path=write_enabled_docx,
            operation="add",
            text=None,
            confirm=True,
        )


# --- add_hyperlink ---


@pytest.mark.unit
def test_add_hyperlink_http_url_returns_dict(write_enabled_docx):
    """add_hyperlink with a valid https:// URL returns expected dict."""
    result = add_hyperlink(
        path=write_enabled_docx,
        paragraph_index=0,
        text="Visit us",
        url="https://example.com",
        confirm=True,
    )
    assert result["paragraph_index"] == 0
    assert result["url"] == "https://example.com"
    assert result["text"] == "Visit us"


@pytest.mark.unit
def test_add_hyperlink_javascript_blocked():
    """add_hyperlink raises ValidationError for javascript: URL before any file access."""
    with pytest.raises(ValidationError, match="http"):
        add_hyperlink(
            path="irrelevant.docx",
            paragraph_index=0,
            text="xss",
            url="javascript:alert(document.cookie)",
        )


@pytest.mark.unit
def test_add_hyperlink_file_protocol_blocked():
    """add_hyperlink raises ValidationError for file:// URL before any file access."""
    with pytest.raises(ValidationError, match="http"):
        add_hyperlink(
            path="irrelevant.docx",
            paragraph_index=0,
            text="local file",
            url="file:///etc/passwd",
        )


@pytest.mark.unit
def test_add_hyperlink_url_too_long():
    """add_hyperlink raises ValidationError for URL exceeding 2048 characters."""
    long_url = "https://example.com/" + "a" * 2030
    assert len(long_url) > 2048
    with pytest.raises(ValidationError, match="2048"):
        add_hyperlink(
            path="irrelevant.docx",
            paragraph_index=0,
            text="too long",
            url=long_url,
        )


@pytest.mark.unit
def test_add_hyperlink_paragraph_out_of_range(write_enabled_docx):
    """add_hyperlink raises NotFoundError for an out-of-range paragraph_index."""
    with pytest.raises(NotFoundError):
        add_hyperlink(
            path=write_enabled_docx,
            paragraph_index=9999,
            text="nowhere",
            url="https://example.com",
            confirm=True,
        )


@pytest.mark.unit
def test_add_hyperlink_requires_write_gate(tmp_path, monkeypatch):
    """add_hyperlink raises NotAllowedError when WORD_ENABLE_WRITE is not set."""
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    doc_path = str(tmp_path / "doc.docx")
    Document().save(doc_path)
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        add_hyperlink(path=doc_path, paragraph_index=0, text="x", url="http://x.com", confirm=True)


@pytest.mark.unit
def test_add_hyperlink_requires_confirm(write_enabled_docx, monkeypatch):
    """add_hyperlink raises ValidationError when confirm=False."""
    with pytest.raises(ValidationError, match="confirm"):
        add_hyperlink(path=write_enabled_docx, paragraph_index=0, text="x", url="http://x.com", confirm=False)


# --- add_footnote ---


@pytest.mark.unit
def test_add_footnote_no_footnotes_part_raises(tmp_path, monkeypatch):
    """add_footnote raises WordMCPError when document has no footnotes part."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "no_footnotes.docx"
    d = Document()
    d.add_paragraph("A paragraph")
    d.save(str(path))

    # Build a mock document with one paragraph and NO footnotes relationship
    mock_para = MagicMock()
    mock_doc = MagicMock()
    mock_doc.paragraphs = [mock_para]
    mock_doc.part.rels.values.return_value = []

    with patch("wordmcp.document_docx._load_doc", return_value=mock_doc):
        with pytest.raises(WordMCPError, match="no footnotes part"):
            add_footnote(str(path), 0, "My footnote", confirm=True)


@pytest.mark.unit
def test_add_footnote_requires_write_gate(tmp_path, monkeypatch):
    """add_footnote raises NotAllowedError when WORD_ENABLE_WRITE is not set."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "fn_gate.docx"
    d = Document()
    d.add_paragraph("Para")
    d.save(str(path))
    with pytest.raises(NotAllowedError):
        add_footnote(str(path), 0, "footnote text", confirm=True)


@pytest.mark.unit
def test_add_footnote_paragraph_out_of_range(tmp_path, monkeypatch):
    """add_footnote raises NotFoundError for an out-of-range paragraph_index."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    path = tmp_path / "fn_oor.docx"
    d = Document()
    d.add_paragraph("Single paragraph")
    d.save(str(path))
    with pytest.raises(NotFoundError):
        add_footnote(str(path), 99, "footnote", confirm=True)


