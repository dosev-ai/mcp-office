"""
Security tests for wordmcp — Sprint A hardening.
All tests import ONLY from wordmcp.document_docx (no server.py imports).
All tests use @pytest.mark.security.
"""
from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from wordmcp.document_docx import (
    FileError,
    NotAllowedError,
    ValidationError,
    _check_evidence_path,
    _check_image_path,
    _check_path,
    add_paragraph,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(path: Path) -> None:
    """Create a minimal valid .docx at path."""
    Document().save(str(path))


def _make_docx_with_text(path: Path, text: str) -> None:
    d = Document()
    d.add_paragraph(text)
    d.save(str(path))


def _make_png(path: Path) -> None:
    """Create a minimal valid 1x1 PNG at path."""
    def _chunk(name: bytes, data: bytes) -> bytes:
        c = name + data
        return (
            struct.pack(">I", len(data))
            + c
            + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        )
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\xff\xff")
    img = (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", idat)
        + _chunk(b"IEND", b"")
    )
    path.write_bytes(img)


# ===========================================================================
# GROUP A — Path traversal (5 tests)
# ===========================================================================


@pytest.mark.security
def test_a_path_traversal_dotdot(tmp_path, monkeypatch):
    """Path with ../escape is rejected by allowlist check."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    evil = str(tmp_path / ".." / "escape.docx")
    with pytest.raises(ValidationError):
        _check_path(evil)


@pytest.mark.security
def test_a_path_traversal_absolute_outside(tmp_path, monkeypatch):
    """Absolute path outside allowlist raises ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="not in allowlist"):
        _check_path("/absolute/outside/file.docx")


@pytest.mark.security
def test_a_path_traversal_empty_string(tmp_path, monkeypatch):
    """Empty string path raises ValidationError (extension check)."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises((ValidationError, FileError, OSError)):
        _check_path("")


@pytest.mark.security
def test_a_path_traversal_whitespace_only(tmp_path, monkeypatch):
    """Whitespace-only path raises ValidationError (extension check)."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises((ValidationError, FileError, OSError)):
        _check_path("   ")


@pytest.mark.security
def test_a_path_traversal_windows_drive_outside_allowlist(tmp_path, monkeypatch):
    """Windows drive letter path not in allowlist raises ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError):
        _check_path("Z:\\outside\\secret.docx")


# ===========================================================================
# GROUP B — Null-byte injection (4 tests)
# ===========================================================================


@pytest.mark.security
def test_b_null_byte_in_check_path(tmp_path, monkeypatch):
    """_check_path rejects path containing null byte before Path.resolve()."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte"):
        _check_path("safe\x00../evil.docx")


@pytest.mark.security
def test_b_null_byte_in_check_image_path(tmp_path, monkeypatch):
    """_check_image_path rejects image path containing null byte."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte"):
        _check_image_path("img\x00evil.png")


@pytest.mark.security
def test_b_null_byte_path_prefix_escape(tmp_path, monkeypatch):
    """Null byte in path rejects before any filesystem call is made."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte"):
        _check_path(str(tmp_path / "doc.docx") + "\x00../../../etc/passwd")


@pytest.mark.security
def test_b_null_byte_image_path_escape(tmp_path, monkeypatch):
    """Null byte in image_path rejects before any filesystem call."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte"):
        _check_image_path(str(tmp_path / "img.png") + "\x00evil")


@pytest.mark.security
def test_b_null_byte_in_check_evidence_path(tmp_path, monkeypatch):
    """B-005: _check_evidence_path rejects path containing null byte."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    with pytest.raises(ValidationError, match="Null byte in path"):
        _check_evidence_path("\x00evil.jsonl")


# ===========================================================================
# GROUP C — Allowlist enforcement (6 tests)
# ===========================================================================


@pytest.mark.security
def test_c_allowlist_unset_raises(tmp_path, monkeypatch):
    """Unset WORD_ALLOWLIST_ROOTS raises ValidationError."""
    monkeypatch.delenv("WORD_ALLOWLIST_ROOTS", raising=False)
    with pytest.raises(ValidationError, match="WORD_ALLOWLIST_ROOTS"):
        _check_path(str(tmp_path / "test.docx"))


@pytest.mark.security
def test_c_allowlist_one_valid_root_accepted(tmp_path, monkeypatch):
    """Path inside single allowlist root is accepted."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    resolved = _check_path(str(path))
    assert resolved == path.resolve()


@pytest.mark.security
def test_c_allowlist_two_roots_accept_both(tmp_path, monkeypatch, tmp_path_factory):
    """Path in either of two allowlist roots is accepted."""
    second = tmp_path_factory.mktemp("second_root")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", f"{tmp_path},{second}")
    p1 = tmp_path / "a.docx"
    p2 = second / "b.docx"
    _make_docx(p1)
    _make_docx(p2)
    assert _check_path(str(p1)) == p1.resolve()
    assert _check_path(str(p2)) == p2.resolve()


@pytest.mark.security
def test_c_allowlist_path_resolves_within_root(tmp_path, monkeypatch):
    """Subdirectory within allowlist root is accepted."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    sub = tmp_path / "sub"
    sub.mkdir()
    path = sub / "doc.docx"
    _make_docx(path)
    resolved = _check_path(str(path))
    assert resolved.is_relative_to(tmp_path)


@pytest.mark.security
def test_c_allowlist_path_outside_root_rejected(tmp_path, monkeypatch, tmp_path_factory):
    """Path in a different tempdir not in allowlist is rejected."""
    other = tmp_path_factory.mktemp("other_root")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    outside = other / "file.docx"
    _make_docx(outside)
    with pytest.raises(ValidationError, match="not in allowlist"):
        _check_path(str(outside))


@pytest.mark.security
def test_c_allowlist_whitespace_strip_in_roots(tmp_path, monkeypatch):
    """Allowlist root with surrounding whitespace is accepted after strip."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", f"  {tmp_path}  ")
    path = tmp_path / "test.docx"
    _make_docx(path)
    resolved = _check_path(str(path))
    assert resolved == path.resolve()


# ===========================================================================
# GROUP D — Extension validation (5 tests)
# ===========================================================================


@pytest.mark.security
def test_d_extension_docx_accepted(tmp_path, monkeypatch):
    """.docx extension is accepted by _check_path."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "valid.docx"
    _make_docx(path)
    resolved = _check_path(str(path))
    assert resolved.suffix == ".docx"


@pytest.mark.security
def test_d_extension_xlsx_rejected(tmp_path, monkeypatch):
    """.xlsx extension raises ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "spreadsheet.xlsx"
    path.write_bytes(b"PK\x03\x04")
    with pytest.raises(ValidationError, match="Unsupported extension"):
        _check_path(str(path))


@pytest.mark.security
def test_d_extension_py_rejected(tmp_path, monkeypatch):
    """.py extension raises ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "script.py"
    path.write_text("print('evil')")
    with pytest.raises(ValidationError, match="Unsupported extension"):
        _check_path(str(path))


@pytest.mark.security
def test_d_extension_double_docx_accepted(tmp_path, monkeypatch):
    """file.docx.docx has suffix .docx and is accepted by extension check."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "file.docx.docx"
    _make_docx(path)
    resolved = _check_path(str(path))
    assert resolved.suffix == ".docx"


@pytest.mark.security
def test_d_extension_no_extension_rejected(tmp_path, monkeypatch):
    """File with no extension raises ValidationError."""
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "nodotfile"
    path.write_bytes(b"data")
    with pytest.raises(ValidationError, match="Unsupported extension"):
        _check_path(str(path))


# ===========================================================================
# GROUP E — WORD_ENABLE_WRITE gate (5 tests)
# ===========================================================================


@pytest.mark.security
def test_e_write_gate_unset_rejects_add_paragraph(tmp_path, monkeypatch):
    """add_paragraph raises NotAllowedError when WORD_ENABLE_WRITE is unset."""
    monkeypatch.delenv("WORD_ENABLE_WRITE", raising=False)
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(NotAllowedError, match="WORD_ENABLE_WRITE"):
        add_paragraph(str(path), "hello", confirm=True)


@pytest.mark.security
def test_e_write_gate_empty_string_rejects(tmp_path, monkeypatch):
    """WORD_ENABLE_WRITE='' raises NotAllowedError."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(NotAllowedError):
        add_paragraph(str(path), "hello", confirm=True)


@pytest.mark.security
def test_e_write_gate_false_string_rejects(tmp_path, monkeypatch):
    """WORD_ENABLE_WRITE='false' raises NotAllowedError."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "false")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    with pytest.raises(NotAllowedError):
        add_paragraph(str(path), "hello", confirm=True)


@pytest.mark.security
def test_e_write_gate_true_accepts_add_paragraph(tmp_path, monkeypatch):
    """WORD_ENABLE_WRITE='true' allows add_paragraph to succeed."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "true")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    result = add_paragraph(str(path), "hello", confirm=True)
    assert result["text"] == "hello"


@pytest.mark.security
def test_e_write_gate_true_case_insensitive(tmp_path, monkeypatch):
    """WORD_ENABLE_WRITE='True' (capital T) also allows add_paragraph."""
    monkeypatch.setenv("WORD_ENABLE_WRITE", "True")
    monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
    path = tmp_path / "test.docx"
    _make_docx(path)
    result = add_paragraph(str(path), "world", confirm=True)
    assert result["text"] == "world"


# ===========================================================================
# GROUP F — Confirm gate (4 tests)
# ===========================================================================


# ===========================================================================
# NT-005: UNC path rejection in all 3 word path-check functions
# ===========================================================================

@pytest.mark.security
class TestWordUncPathRejection:
    """NT-005: UNC path guards in _check_path, _check_image_path, _check_evidence_path."""

    def test_check_path_rejects_unc(self, monkeypatch, tmp_path):
        """NT-005: _check_path must raise ValidationError for \\\\server\\share UNC path.
        WORD_ALLOWLIST_ROOTS must be set so the allowlist pre-check doesn't fire first.
        """
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        with pytest.raises(ValidationError, match="UNC"):
            _check_path("\\\\server\\share\\evil.docx")

    def test_check_image_path_rejects_unc(self, monkeypatch, tmp_path):
        """NT-005: _check_image_path must raise ValidationError for \\\\server\\share UNC path.
        WORD_ALLOWLIST_ROOTS must be set so the allowlist pre-check doesn't fire first.
        """
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        with pytest.raises(ValidationError, match="UNC"):
            _check_image_path("\\\\server\\share\\evil.png")

    def test_check_evidence_path_rejects_unc(self, monkeypatch, tmp_path):
        """NT-005: _check_evidence_path must raise ValidationError for \\\\server\\share UNC path.
        WORD_ALLOWLIST_ROOTS must be set so the allowlist pre-check doesn't fire first.
        """
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", str(tmp_path))
        with pytest.raises(ValidationError, match="UNC"):
            _check_evidence_path("\\\\server\\share\\evil.json")

    def test_check_path_rejects_forward_slash_unc(self, monkeypatch):
        """W2-005: forward-slash UNC paths must be rejected by _check_path."""
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", "C:\\Temp")
        with pytest.raises(ValidationError, match="UNC"):
            _check_path("//server/share/evil.docx")

    def test_check_image_path_rejects_forward_slash_unc(self, monkeypatch):
        """W2-005: forward-slash UNC paths must be rejected by _check_image_path."""
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", "C:\\Temp")
        with pytest.raises(ValidationError, match="UNC"):
            _check_image_path("//server/share/evil.png")

    def test_check_evidence_path_rejects_forward_slash_unc(self, monkeypatch):
        """W2-005: forward-slash UNC paths must be rejected by _check_evidence_path."""
        monkeypatch.setenv("WORD_ALLOWLIST_ROOTS", "C:\\Temp")
        with pytest.raises(ValidationError, match="UNC"):
            _check_evidence_path("//server/share/evil.jsonl")

